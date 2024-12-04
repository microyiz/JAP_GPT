import re
import os
import glob
import time
import warnings
import docx
import mysql.connector
from docx import Document
from jap_paper_revise import read_docx_to_string
from jap_paper_revise import read_docx_to_string_with_format
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate  
from langchain.chains import LLMChain   
from typing import Any


def split_into_sentences(text):
        """
        Splits text into sentences based on common Japanese sentence endings.
        
        :param text: Text to split.
        :return: List of sentences.
        """
        sentence_endings = re.compile(r'(?<=[。！？])\s*')
        sentences = sentence_endings.split(text)
        return sentences

def knowledge_point_analysis(paper, question_type):
    """
    Revises a list using a language model and saves the results.

    :param rows: Loaded list data to revise.
    :param filename: Path to save the revised document.
    """
    llm = ChatOpenAI(
        temperature=0.8,
        model="gpt-4o"
    ) 
    prompt_two = ChatPromptTemplate.from_template(
        "I have provided a sample of questions Japanese Language Proficiency Test N4 level: {question_paper}.\
        Now base on these questions, detailed analyze the types of questions"
    )
    
    chain_one = LLMChain(llm=llm, prompt=prompt_two)


    inputs_two = {'question_paper': paper}
    revise_result = chain_one.run(inputs_two)
    output_doc = Document()
    sentences = split_into_sentences(revise_result)
    for sentence in sentences:
        output_doc.add_paragraph(sentence)

    output_path = os.path.join(question_type, f"question_type.docx")
    output_doc.save(output_path)
# 删除包含某个关键词的行（“意味”，“例文”）
def delete_keyword_line(path, key_words, save_path):
    doc = docx.Document(path)

    full_text = []
    count = 0
    for paragraph in doc.paragraphs:
        flag = False
        paragraph.text = paragraph.text.replace('\n', '')
        for key_word in key_words:
            matchPattern = re.compile(key_word)
            if matchPattern.search(paragraph.text):
                flag = True
            elif len(paragraph.text) == 0:
                flag = True
            
        if (flag == False):
            full_text.append(str(count)+ '. '+paragraph.text)
            count += 1
        else:
            pass
                
    text = "\n".join(full_text)
    outputdoc = Document()
    outputdoc.add_paragraph(text)
    outputdoc.save(save_path)
    
def process_material(meterial_folder):
    material_doc = []
    for filepath in glob.glob(os.path.join(meterial_folder, "*.docx")):
        material_doc.append(read_docx_to_string_with_format(filepath))
    return material_doc

def knowledge_points_questions_generate(path,sample_questions, meterial_paper):
        llm = ChatOpenAI(
            temperature=0.8,
            model="gpt-4o"
        )
        prompt_four = ChatPromptTemplate.from_template(
            "Now here are the knowledge points of N4 level Japanese language test :{material}, it includes vocabulary and grammar knowledge of Japanese.\
            Please use this material as a reference, imitate the format of these questions:{sample_questions}, give a corresponding question for each knowledge point, all the knowledge points should have one corresponding question. \
            All provided questions should meet the following criteria:\
            1.No duplicate questions. All the questions should be unique. Delete any repeated questions and replace them with new ones.\
            2.No duplicate options. All options should be unique and meaningful within the context of the question.\
            3.No duplicate correct answers. The answer to the question should be unique in the context of the exam. Please not have two or more than two suitable answer to choose the most suitable one, make sure it has only one suitable answer that is absolutely correct, you can add specific condition in the question stem or change the options.\
            4.Grammatical correctness. The title and stem of the question should be grammatically correct.You can put back the correct option to the question stem, if there is a grammar issue, please revise the question stem and the options.\
            5.Relevance of options. \
            One modification idea is that the correct option should more clearly point to a suitable answer which is reasonable and fits the context of the stem, while ensuring that the other options are clearly inappropriate or incorrect. \
            For example, in the question 'わたしは、毎朝（ 　　　　　 ）を飲みます。', all options like お茶, コーヒー, ジュース, and 水 are suitable for the verb 'drink,' which makes the question ambiguous. A better example would be 'わたしは、毎朝（ 　　　　　 ）を食べます。1. お茶 2. コーヒー 3. パン 4. 花', where only パン is an appropriate option for 'eat,' and the other options (お茶, コーヒー, 花) are clearly unsuitable for eating, which makes it a good question because it has only one clear answer “パン”.\
            Another modification idea is that the question should clearly indicate what cannot be chosen. The stem must specify the context in which one option is clearly inappropriate, while all other options are suitable.\
            For example, in the stem 'その 映画は ( 　　　　　 ) ではありません', options like “つまらない”, “面白い”, and “怖い” are appropriate descriptors for a film, but “おいしい” is not, making it the correct answer. If the question asks an obvious 'no' (choose the most inappropriate one), make sure the question stem itself is in negative form “ません”.\
            So in these options, ignore the culture backgroud and avoid subjective consciousness questions and options.\
            6. If the question is about the pronunciation of a word or how a particular word is used or its katakana, hiragana, use the brackets to emphasize the Japanese words. Do not have any underline in the questions. Do not show the right answer in the question stem.\
            If the question is ask a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\
            If the question is ask a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question. "
            "The new questions should be in a multiple-choice format and appropriate for the Japanese Language Proficiency Test N3 level.\n"
            "Please create the new questions for each knowledge point, each with four different options. Ensure that only one of these options is correct and should be evenly distributed among 1, 2, 3, and 4.\n"
            "The instruction of the questions should be attached in front of each question. The corresponding specific knowledge point (for example: ～くれる【N4】Meaning: Expresses the accepting of ownership of something from someone else. 「あげる・もらう・くれる」each express giving and receiving from different standpoints. 用法：（だれか）に（なに）をあげる・もらう・くれる) should be attached after each question.\n"
            "Finally, all the answers will be attached at the end. Do not attach the answer after each question.")
        
        chain_four = LLMChain(llm=llm, prompt = prompt_four)
        input_four = {'material': meterial_paper,
                      'sample_questions':sample_questions}
        
        matching_result = chain_four.run(input_four)

        output_doc = Document()
        sentences = split_into_sentences(matching_result)
        
        for sentence in sentences:
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(path, f"All_questions_corresponding_knowledge_points.docx")
        output_doc.save(output_path)

        return matching_result

def knowledge_points_match(path, exam_questions, meterial_paper):
        llm = ChatOpenAI(
            temperature=0.8,
            model="gpt-4o"
        )
        prompt_four = ChatPromptTemplate.from_template(
            "Now here are the knowledge points list of Japanese language test :{material}, it includes vocabulary and grammar knowledge of Japanese.\
            Please use it as a reference to matching this list of questions which belongs to Japanese test exam: {exam_questions}. You need to match every question !\
            You need to find the corresponding specific knowledge points that the question stem and options contain for each question, there may be multiple knowledge points for each question, you need to find all of them. For example if the question is asked about grammar or vocabulary of several combination words, specifically give the list numbers of these words, and which part they belong to: vocabulary or grammar. \
            You should give the whole content, include the stem and all the four original options of the original questions, attach the corresponding knowledge points with them. ")
        
        chain_four = LLMChain(llm=llm, prompt = prompt_four)
        input_four = {'material': meterial_paper,
                      'exam_questions': exam_questions}
        
        matching_result = chain_four.run(input_four)

        output_doc = Document()
        sentences = split_into_sentences(matching_result)
        
        for sentence in sentences:
            output_doc.add_paragraph(sentence)
            
        # 路径修改
        output_path = os.path.join(path, f"exam_numbered_knowledge_points.docx")
        output_doc.save(output_path)

        return matching_result

def main():
    paper_1 = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\TDLEG learning materials\\N4 Notes 文法_numbered.docx"
    paper_2 = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\TDLEG learning materials\\N4 Notes 語彙_processed.docx"
    save_path = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\TDLEG learning materials"
    exam_questions = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\test 1 paper\\Test 1 Question Paper.docx"

    document_1 = Document(paper_1)
    document_2 = Document(paper_2)
    text = ""
    for paragraph in document_1.paragraphs:
        text += paragraph.text
    for paragraph in document_2.paragraphs:
        text += paragraph.text

    text_1 = read_docx_to_string(paper_1)
    text_2 = read_docx_to_string(paper_2)
    text = text_1 + text_2
    question_text = read_docx_to_string_with_format(exam_questions)
    result = knowledge_points_match(save_path, question_text, text)
    
    # 去除意思和例文
    # key_words = ["意味","例"]
    # grammar_paper = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\TDLEG learning materials\\N4 Notes 語彙.docx"
    # save_path = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\TDLEG learning materials\\N4 Notes 語彙_processed.docx"
    # delete_keyword_line(grammar_paper, key_words, save_path)
    


if __name__ == "__main__":
    main()