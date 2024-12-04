import re
import os
import glob
import time
import warnings
import docx
import mysql.connector
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate  
from langchain.chains import LLMChain   
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from typing import Any

from jap_paper_revise import return_revised_result
from jap_paper_revise import extract_student_id
from jap_paper_revise import read_docx_to_string_with_format
from jap_paper_revise import produce_split_new_question_list
from db_util import drop_table_query ,create_table_query ,insert_query,show_fiverows_query,select_mistake_query,db


class AnswerChecker:
    """
    A class to check student answers against correct answers, generate a mistakes report, and save it to a specified output folder.
    """
    def __init__(self, correct_answers_path, input_folder, output_folder):
        """
        Initializes the AnswerChecker with paths to the correct answers, input folder, and output folder.

        :param correct_answers_path: Path to the document containing the correct answers.
        :param input_folder: Path to the folder containing student answer documents.
        :param output_folder: Path to the folder where the mistakes reports will be saved.
        """
        self.correct_answers_path = correct_answers_path
        self.input_folder = input_folder
        self.output_folder = output_folder
        os.makedirs(self.output_folder, exist_ok=True)
    
    def read_answers(self, doc_path):
        """
        Reads and extracts answers from a given document.

        :param doc_path: Path to the document containing answers.
        :return: Two dictionaries containing answers from Part 1 and Part 2 of the test.
        """
        doc = docx.Document(doc_path)
        answers_part1 = {}
        answers_part2 = {}
        current_part = None

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "日本語実力テスト1 (第1部" in text:
                current_part = "part1"
            elif "日本語実力テスト1 (第2部" in text:
                current_part = "part2"
            elif text.startswith("問題"):
                question, answer = text.split("：")
                question_num = question.split("　")[1]
                if current_part == "part1":
                    answers_part1[question_num] = answer
                elif current_part == "part2":
                    answers_part2[question_num] = answer

        return answers_part1, answers_part2

    def save_mistakes_to_docx(self, mistakes, output_filename, filename):
        """
        Saves identified mistakes to a DOCX file.

        :param mistakes: A dictionary of mistakes identified in the student's answers.
        :param output_filename: Path where the DOCX file will be saved.
        :param filename: The base name of the student's document (used for report title).
        """
        doc = docx.Document()
        doc.add_heading(f"{filename}の誤答レポート", level=1)

        k=1
        
        for part, mistakes_part in mistakes.items():
            doc.add_heading(part, level=2)
            for question, (student_answer, correct_answer) in mistakes_part.items():
                doc.add_paragraph(f"{k}." f"問題 {question}:")
                doc.add_paragraph(f"学生の回答: {student_answer}", style='List Bullet')
                doc.add_paragraph(f"正解: {correct_answer}", style='List Bullet')
                doc.add_paragraph("\n")
                k=k+1

        doc.save(output_filename)

    def compare_answers_and_generate_report(self, student_answers_path):
        """
        Compares student answers with correct answers, identifies mistakes, and generates a report.

        :param student_answers_path: Path to the student's answers document.
        """
        # Read answers from both documents
        student_answers_part1, student_answers_part2 = self.read_answers(student_answers_path)
        correct_answers_part1, correct_answers_part2 = self.read_answers(self.correct_answers_path)

        # Compare the answers and identify mistakes
        mistakes_part1 = {}
        for question, correct_answer in correct_answers_part1.items():
            student_answer = student_answers_part1.get(question, None)
            if student_answer != correct_answer:
                mistakes_part1[question] = (student_answer, correct_answer)

        mistakes_part2 = {}
        for question, correct_answer in correct_answers_part2.items():
            student_answer = student_answers_part2.get(question, None)
            if student_answer != correct_answer:
                mistakes_part2[question] = (student_answer, correct_answer)

        # Get the filename without the extension
        filename = os.path.splitext(os.path.basename(student_answers_path))[0]

        # Save the mistakes to a DOCX file
        mistakes = {"Part 1": mistakes_part1, "Part 2": mistakes_part2}
        mistakes_output_filename = os.path.join(self.output_folder, f"{filename}_mistakes.docx")
        self.save_mistakes_to_docx(mistakes, mistakes_output_filename, filename)
        print(f"Mistakes report saved to {mistakes_output_filename}")

    def process_all_files(self):
        """
        Processes all .docx files in the input folder, compares the answers, and saves the mistakes reports in the output folder.
        """
        # Iterate over all .docx files in the input folder
        for filepath in glob.glob(os.path.join(self.input_folder, "*.docx")):
            self.compare_answers_and_generate_report(filepath)

    


class DocumentProcessor:
    def __init__(self, input_folder, output_folder, output_analysis_folder, revised_newpaper_folder, mistake_database, material_folder, matched_knowledge_points_folder):
        """
        Initializes the DocumentProcessor with input and output folders.
        
        :param input_folder: Folder containing the input Word documents.
        :param output_folder: Folder to save the new Word documents.
        :param output_analysis_folder: Folder to save the processed knowledge point Word documents.
        """
        self.input_folder = input_folder
        self.output_folder = output_folder
        self.output_analysis_folder = output_analysis_folder
        self.revised_newpaper_folder = revised_newpaper_folder
        self.mistake_database = mistake_database
        self.material_folder = material_folder
        self.matched_knowledge_points_folder = matched_knowledge_points_folder

    def load_document(self, filepath):
        """
        Loads a Word document from the specified filepath.
        
        :param filepath: Path to the Word document.
        :return: Loaded document data.
        """
        loader = UnstructuredWordDocumentLoader(filepath)
        data = loader.load()
        return data
    
    def append_splits(self, splits, output_path):
        """
        Appends the split document chunks into a new Word document.
        
        :param splits: List of document chunks.
        :param output_path: Path to save the new Word document.
        """
        doc = Document()
        for split in splits:
            doc.add_paragraph(split.page_content)
        doc.save(output_path)

    def split_into_sentences(self, text):
        """
        Splits text into sentences based on common Japanese sentence endings.
        
        :param text: Text to split.
        :return: List of sentences.
        """
        sentence_endings = re.compile(r'(?<=[。！？])\s*')
        sentences = sentence_endings.split(text)
        return sentences

    
    def paper_revise(self, rows, matching, material, mistake_count, filename):
        """
        Revises a list using a language model and saves the results.

        :param rows: Loaded list data to revise.
        :param filename: Path to save the revised document.
        """
        llm = ChatOpenAI(
            temperature=0.8,
            model="gpt-4o"
        )
        if mistake_count > 5:
            prompt_one = ChatPromptTemplate.from_template(
                "Below is a list of incorrect answers provided by Japanese language students: {error_report}\n"
                "Each question includes the student's incorrect choice and the correct answer.\n"
                "Based on these errors and the corresponding knowledge points {matching_knowledge_points}, generate new practice questions targeting similar grammar or vocabulary points to help students strengthen their understanding.\n"
                "1.No duplicate questions. All the questions should be unique. Delete any repeated questions and replace them with new ones.\n"
                "2.No duplicate options. All options should be unique and meaningful within the context of the question.\n"
                "3.No duplicate answers. The answer to the question should be unique in the context of the exam. Please not have two or more than two suitable answer to choose the most suitable one, make sure it has only one suitable answer, you can add specific condition in the question stem or change the options.\n"
                "4.Grammatical correctness. The title and stem of the question should be grammatically correct. You can put back the correct option to the question stem, if there is a grammar issue such as the object word of the sentence is incorrect, please revise the question stem and the options. If all the answer can't fit the question, recreate a question with same knowledge point to replace it.\n"
                "5.Relevance of options. \n"
                "One modification idea is that the correct option should more clearly point to a suitable answer which is reasonable and fits the context of the stem, while ensuring that the other options are clearly inappropriate or incorrect. \n"
                "For example, in the question 'わたしは、毎朝（ 　　　　　 ）を飲みます。', all options like お茶, コーヒー, ジュース, and 水 are suitable for the verb 'drink,' which makes the question ambiguous. A better example would be 'わたしは、毎朝（ 　　　　　 ）を食べます。1. お茶 2. コーヒー 3. パン 4. 花', where only パン is an appropriate option for 'eat,' and the other options (お茶, コーヒー, 花) are clearly unsuitable for eating, which makes it a good question because it has only one clear answer “パン”.\n"
                "Another modification idea is that the question should clearly indicate what cannot be chosen. The stem must specify the context in which one option is clearly inappropriate, while all other options are suitable.\n"
                "For example, in the stem 'その 映画は ( 　　　　　 ) ではありません', options like “つまらない”, “面白い”, and “怖い” are appropriate descriptors for a film, but “おいしい” is not, making it the correct answer. If the question asks an obvious 'no' (choose the most inappropriate one), make sure the question stem itself is in negative form “ません”.\n"
                "So in these options, ignore the culture background and avoid subjective consciousness questions and options. There is a good example: この料理は 塩が （ ）。1. 入っています 2. 入れておきます 3. 入れています 4. 入ってあります Correct answer: 1. 入っています Wrong answer: 2. が àを 入れておきます 3. が àを 入れていますThe Information is: 1. 自動詞 Intransitive verb 不及物動詞 が 入る はいる 2. 他動詞Transitive verb 及物動詞 を 入れる    いれる. You can imitate this to make questions\n"
                "6. If the question is about the pronunciation of a word or how a particular word is used or its katakana, hiragana, use the brackets to emphasize the Japanese words. Do not have any underline in the questions. Do not show the right answer in the question stem.\n"
                "If the question is ask a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\n"
                "If the question is ask a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question.\n"
                "If any of the above problems occur, please modify the questions to eliminate these issues. Ensure that the structure remains the same as the original questions, and all answers should be attached at the end. Do not attach the answer after each question. \n"
                "The new questions should be in a multiple-choice format and appropriate for the Japanese Language Proficiency Test N3 level.\n"
                "Please create {num_questions} new questions, each with four different options. Ensure that only one of these options is correct and should be evenly distributed among 1, 2, 3, and 4.\n"
                "The instruction of the questions should be attached in front of each question. Do not have hints on the questions. The numbers of the questions should be in this format: **1**. The intruction of answer list should be **Answers**.\n"
                "Finally, all the answers will be attached at the end. Do not attach the answer after each question."
            )
            
            inputs_one = {
                'matching_knowledge_points':matching,
                'error_report': rows,
                'num_questions': 20
            }
        elif mistake_count <= 5:
            prompt_one = ChatPromptTemplate.from_template(
                "Below is a list of incorrect answers provided by Japanese language students: {error_report}\n"
                "Each question includes the student's incorrect choice and the correct answer.\n"
                "Based on these errors and the corresponding knowledge points: {matching_knowledge_points} and use these materials as supplementary knowledge points: {material}, generate new practice questions targeting similar grammar or vocabulary points to help students strengthen their understanding.\n"
                "The question should meet following standards:\n"
                "1.No duplicate questions. All the questions should be unique. Delete any repeated questions and replace them with new ones.\n"
                "2.No duplicate options. All options should be unique and meaningful within the context of the question.\n"
                "3.No duplicate answers. The answer to the question should be unique in the context of the exam. Please not have two or more than two suitable answer to choose the most suitable one, make sure it has only one suitable answer, you can add specific condition in the question stem or change the options.\n"
                "4.Grammatical correctness. The title and stem of the question should be grammatically correct. You can put back the correct option to the question stem, if there is a grammar issue such as the object word of the sentence is incorrect, please revise the question stem and the options. If all the answer can't fit the question, recreate a question with same knowledge point to replace it.\n"
                "5.Relevance of options. \n"
                "One modification idea is that the correct option should more clearly point to a suitable answer which is reasonable and fits the context of the stem, while ensuring that the other options are clearly inappropriate or incorrect. \n"
                "For example, in the question 'わたしは、毎朝（ 　　　　　 ）を飲みます。', all options like お茶, コーヒー, ジュース, and 水 are suitable for the verb 'drink,' which makes the question ambiguous. A better example would be 'わたしは、毎朝（ 　　　　　 ）を食べます。1. お茶 2. コーヒー 3. パン 4. 花', where only パン is an appropriate option for 'eat,' and the other options (お茶, コーヒー, 花) are clearly unsuitable for eating, which makes it a good question because it has only one clear answer “パン”.\n"
                "Another modification idea is that the question should clearly indicate what cannot be chosen. The stem must specify the context in which one option is clearly inappropriate, while all other options are suitable.\n"
                "For example, in the stem 'その 映画は ( 　　　　　 ) ではありません', options like “つまらない”, “面白い”, and “怖い” are appropriate descriptors for a film, but “おいしい” is not, making it the correct answer. If the question asks an obvious 'no' (choose the most inappropriate one), make sure the question stem itself is in negative form “ません”.\n"
                "So in these options, ignore the culture background and avoid subjective consciousness questions and options\n"
                "6. If the question is about the pronunciation of a word or how a particular word is used or its katakana, hiragana, use the brackets to emphasize the Japanese words. Do not have any underline in the questions. Do not show the right answer in the question stem.\n"
                "If the question is ask a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\n"
                "If the question is ask a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question.\n"
                "If any of the above problems occur, please modify the questions to eliminate these issues. Ensure that the structure remains the same as the original questions, and all answers should be attached at the end. Do not attach the answer after each question. \n"
                "The new questions should be in a multiple-choice format and appropriate for the Japanese Language Proficiency Test N3 level.\n"
                "Please create {num_questions} new questions, each with four different options. Ensure that only one of these options is correct and should be evenly distributed among 1, 2, 3, and 4.\n"
                "The instruction of the questions should be attached in front of each question. Do not have hints on the questions. The numbers of the questions should be in this format: **1**. The intruction of answer list should be **Answers**.\n"
                "Finally, all the answers will be attached at the end. Do not attach the answer after each question."
            )
            
            inputs_one = {
                'material':material,
                'matching_knowledge_points':matching,
                'error_report': rows,
                'num_questions': 20
            }

        chain_one = LLMChain(llm=llm, prompt=prompt_one)

        revise_result = chain_one.run(inputs_one)

        output_doc = Document()
        sentences = self.split_into_sentences(revise_result)
    
        for sentence in sentences:
            # 纠正格式
            sentence.replace("**Answers:**", "**Answers**")
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(self.output_folder, f"{filename}_new_report.docx")
        output_doc.save(output_path)

    
    def knowledge_point_analysis(self, rows, filename, sample_analysis):
        """
        Revises a list using a language model and saves the results.

        :param rows: Loaded list data to revise.
        :param filename: Path to save the revised document.
        """
        llm = ChatOpenAI(
            temperature=0.8,
            model="gpt-4o"
        ) 
        prompt_two = ChatPromptTemplate.from_template(
            "I have provided a sample analysis of a student's mistakes in a Japanese practice test below, labeled as {sample}. \
            The analysis is organized into two main sections: 1.1 Kanji/Vocabulary related mistakes and 1.2 Grammar mistakes. \
            Each section is further divided into smaller sub-sections, such as Pronunciation mistake, Long vowel and short vowel pronunciation mistake, etc. \
            Every sub-section summarizes the specific knowledge points where the student made errors. This format is crucial."

            "Now, I will provide you with another student's error report, labeled as {error_report}. \
            Please analyze the mistakes made by this student using the same structure and detail as in the {sample}. \
            Ensure that the question numbers in {error_report} are retained in the generated analysis.\
            The analysis should be comprehensive and organized into appropriate sections and sub-sections based on the knowledge points involved. \
            Use specific question numbers and describe the errors in a similar manner. Please attach every student's mistakes to the related specific knowledge points."
        )
        
        chain_one = LLMChain(llm=llm, prompt=prompt_two)

        inputs_two = {
            'error_report': rows,
            'sample': sample_analysis
        }
        revise_result = chain_one.run(inputs_two)
        output_doc = Document()
        sentences = self.split_into_sentences(revise_result)
        for sentence in sentences:
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(self.output_analysis_folder, f"{filename}_mistakes_analysis.docx")
        output_doc.save(output_path)


    def question_revise(self, rows, filename, max_iterations=5):
        """
        Revise the generated questions using llm.
        """
        llm = ChatOpenAI(
            temperature=0.8,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt_three = ChatPromptTemplate.from_template(
            "Now these are the new generated Japanese practice questions: {new_paper} \
            You are an excellent Japanese N3 examiner and provide students with appropriate multiple-choice test questions. All provided questions should meet the following criteria:\
            1.No duplicate questions. All the questions should be unique. Delete any repeated questions and replace them with new ones.\
            2.No duplicate options. All options should be unique and meaningful within the context of the question.\
            3.No duplicate correct answers. The answer to the question should be unique in the context of the exam. Please not have two or more than two suitable answer to choose the most suitable one, make sure it has only one suitable answer that is absolutely correct, you can add specific condition in the question stem or change the options.\
            4.Grammatical correctness. The title and stem of the question should be grammatically correct.You can put back the correct option to the question stem, if there is a grammar issue, please revise the question stem and the options.\
            5.Relevance of options. \
            One modification idea is that the correct option should more clearly point to a suitable answer which is reasonable and fits the context of the stem, while ensuring that the other options are clearly inappropriate or incorrect. \
            For example, in the question 'わたしは、毎朝（ 　　　　　 ）を飲みます。', all options like お茶, コーヒー, ジュース, and 水 are suitable for the verb 'drink,' which makes the question ambiguous. A better example would be 'わたしは、毎朝（ 　　　　　 ）を食べます。1. お茶 2. コーヒー 3. パン 4. 花', where only パン is an appropriate option for 'eat,' and the other options (お茶, コーヒー, 花) are clearly unsuitable for eating, which makes it a good question because it has only one clear answer “パン”.\
            Another modification idea is that the question should clearly indicate what cannot be chosen. The stem must specify the context in which one option is clearly inappropriate, while all other options are suitable.\
            For example, in the stem 'その 映画は ( 　　　　　 ) ではありません', options like “つまらない”, “面白い”, and “怖い” are appropriate descriptors for a film, but “おいしい” is not, making it the correct answer. If the question asks an obvious 'no' (choose the most inappropriate one), make sure the question stem itself is in negative form “ません”.\
            So in these options, ignore the culture backgroud and avoid subjective consciousness questions and options.\
            6. If the question is about the pronunciation of a word or how a particular word is used or its katakana, hiragana, use the brackets to emphasize the Japanese words. Do not have any underline in the questions. Do not show the right answer in the question stem.\
            If the question is ask a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\
            If the question is ask a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question.\
            If any of the above problems occur, please modify the questions to eliminate these issues. Ensure that the structure remains the same as the original questions, and all answers should be attached at the end. Do not attach the answer after each question. Do not add questions. \
            Report the changes made at last of the file."
        )
        rows, answer_list = self.loop_each_question(rows)
        final_answer_list = []
        for i in range(20):
            question_number = str(i+1) + ". "
            final_answer_list.append(question_number + answer_list[i])
        answers = "**Answers**"+ "\n"+'\n'.join(final_answer_list)
        rows = rows + answers

        chain_three = LLMChain(llm=llm, prompt = prompt_three)
        input_three = {'new_paper': rows}
        
        for iteration in range(max_iterations):
            revise_result = chain_three.run(input_three)
            if self.check_for_errors(revise_result):  # Implement this method to validate output
                input_three['new_paper'] = revise_result
            else:
                break  # Exit loop if no errors are found

        output_doc = Document()
        sentences = self.split_into_sentences(revise_result)
        
        for sentence in sentences:
            sentence.replace("＿＿＿", "[ ]")
            output_doc.add_paragraph(sentence)

        # 路径修改
        output_path = os.path.join(self.revised_newpaper_folder, f"{filename}_revised.docx")
        output_doc.save(output_path)

    def knowledge_points_match(self, rows, meterial_paper, filename):
        llm = ChatOpenAI(
            temperature=0.8,
            model="gpt-4o"
        )
        prompt_four = ChatPromptTemplate.from_template(
            "Now here are the knowledge points list of Japanese language test :{material}, it includes vocabulary and grammar knowledge of Japanese.\
            Please use it as a reference to matching this list of questions which belongs to Japanese test exam: {error_report}. You need to match every question !\
            You need to find the corresponding specific knowledge points that the question stem and options contain for each question, there may be multiple knowledge points for each question, you need to find all of them. For example if the question is asked about grammar or vocabulary of several combination words, specifically give the list numbers of these words, and which part they belong to: vocabulary or grammar. \
            You should give the whole content, include the stem and all the four original options of the original questions, attach the corresponding knowledge points with them.")
        
        chain_four = LLMChain(llm=llm, prompt = prompt_four)
        input_four = {'material': meterial_paper,
                      'error_report': rows}
        
        matching_result = chain_four.run(input_four)

        output_doc = Document()
        sentences = self.split_into_sentences(matching_result)
        
        for sentence in sentences:
            output_doc.add_paragraph(sentence)
            
        # 路径修改
        output_path = os.path.join(self.matched_knowledge_points_folder, f"{filename}_knowledge_points.docx")
        output_doc.save(output_path)

        return matching_result



    def check_for_errors(self, revised_text):
        """
        Check for errors in the revised question set, including:
        - Multiple correct answers
        - Duplicate questions
        - Errors in the question stem
        - Duplicate options
        
        :param revised_text: The revised text output from GPT.
        :return: True if errors are found, False otherwise.
        """

        # Error check 1: Multiple correct answers
        if self.has_multiple_correct_answers(revised_text):
            return True

        # Error check 2: Duplicate questions
        if self.has_duplicate_questions(revised_text):
            return True

        # Error check 3: Errors in the question stem
        if self.has_stem_errors(revised_text):
            return True

        # Error check 4: Duplicate options
        if self.has_duplicate_options(revised_text):
            return True

        return False  # No errors detected
    

    def has_multiple_correct_answers(self, text):
        llm = ChatOpenAI(
            temperature=0.6,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now check the new generated Japanese practice questions: {new_paper} \
            Please revise are there multiple correct answers for the question options? \
            If so, just output True, otherwise output False"
        )

        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'new_paper': text}

        result = chain.run(input)
        bool_dict = {"True": True, "False": False, "true": True, "false":False}

        return bool_dict[result]

    def has_duplicate_questions(self, text):
        """
        Check if any questions are duplicated.
        
        :param text: The text to check.
        :return: True if duplicate questions are detected.
        """
        questions = re.findall(r'\d+\.\s*(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n', 
                               text, re.DOTALL
                            )
        seen_questions = set()
        for question in questions:
            question_text = question[0].strip()  # Get the question part only
            if question_text in seen_questions:
                return True  # Duplicate found
            seen_questions.add(question)
        return False

    def has_stem_errors(self, text):
        """
        Check for errors in the question stem.
        
        :param text: The text to check.
        :return: True if errors in the stem are detected.
        """
        # # 只是一个最简单的检查是否有标点符号的检查，需要修改为检查题目内容？
        # # Example: Detect missing or malformed question stems.
        # questions = re.findall(r'\d+\.\s*(.*?)\n', text)
        # for question in questions:
        #     if not question.strip() or len(question.strip().split()) < 5:  # Example check
        #         return True  # Detected a malformed stem
        # return False
    
        llm = ChatOpenAI(
            temperature=0.6,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now check the new generated Japanese practice questions: {new_paper} \
            Please revise are there some errors in the questions? \
            If so, just output True, otherwise output False"
        )

        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'new_paper': text}

        result = chain.run(input)
        bool_dict = {"True": True, "False": False, "true": True, "false":False}
        return bool_dict[result]

    def has_duplicate_options(self, text):
        """
        Check for duplicate options in the questions.
        
        :param text: The text to check.
        :return: True if - options are found within a question.
        """
        # Example: Detect duplicate options for a given question.
        questions_with_options = re.findall(
            r'(\d+)\.\s*(.*?)\n(1\.\s*(.*?)\n)(2\.\s*(.*?)\n)(3\.\s*(.*?)\n)(4\.\s*(.*?)\n)',
            text, re.DOTALL
        )

        for question, _, opt1, _, opt2, _, opt3, _, opt4, _ in questions_with_options:
            options = {opt1.strip(), opt2.strip(), opt3.strip(), opt4.strip()}
            if len(options) < 4:  # If any options are duplicates
                print(f"Duplicate options detected in question {question}: {opt1.strip()}, {opt2.strip()}, {opt3.strip()}, {opt4.strip()}")
                return True
        return False

    # check each question
    def loop_question(self, question, max_iterations):
        llm = ChatOpenAI(
            temperature= 0.6,
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now here is a Japanese practice question:{question}, please revise it to make sure that the question stem and the options have no grammarly error.\
            You should put the options back to the blank in the stem one by one, then check if the sentence has any grammar error, if so, modify it.\
            Then you need to check if there are multiple correct options for this question, if so, please revise the question stem or the options to make sure this question has only one correct answer. \
            Please not have two or more than two suitable answers to choose the most suitable one, make sure it has only one suitable answer that is absolutely correct, you can add specific condition in the question stem or change the options.\
            Only output the modified question and a single number which is the correct option of this question at last. Do not include the analysis."
        )
        # 目前问题，还是有答案错误的
        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'question':question}

        for iteration in range(max_iterations):
            revise_result = chain.run(input)
            if self.question_check(revise_result):  # Implement this method to validate output
                input['question'] = revise_result
            else:
                break  # Exit loop if no errors are found
        return revise_result




    def question_check(self, question):
        llm = ChatOpenAI(
            temperature=0.6,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now here is a Japanese practice question:{question}, each one has four options. You need to put back each options to the question stem to check if there is any grammar issue or just check if the stem has any grammar issue.\
            If there are grammar issues, only output a single word: True, otherwise just output a single word: False. Do not output the analysis."
        )

        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'question': question}

        result = chain.run(input)
        for i in range(50):
            if(result == 'True' or result == 'true' or result == 'False' or result == 'False'):
                break
            else:
                result = chain.run(input)

        bool_dict = {"True": True, "False": False, "true": True, "false":False}

        return bool_dict[result]
    
    def get_answer(self,question):
        return question[len(question)-1]
    
    def loop_each_question(self, questions):
        # questions 是从gpt直接生成的result
        question_list = produce_split_new_question_list(questions)
        question_number = 1
        new_question_list = []
        answer_list = []
        for question in question_list:
            revised_question = self.loop_question(question, 10)
            answer = self.get_answer(revised_question)
            answer_list.append(answer)
            revised_question = revised_question[:-1]

            new_question_list.append( "**"+ str(question_number)+"**" +revised_question)

            print(revised_question)
            print("end of question:",question_number)
            question_number += 1

        return '\n'.join(new_question_list), answer_list

    def process(self, input_paper, correct_answers_path, sample_analysis):
        """
        Processes all .docx files in the input folder and saves the results in the output folder.

        :param input_paper: Question paper.
        :param correct_answers_path: Answer to the question paper.
        """
        warnings.filterwarnings("ignore", category=DeprecationWarning, module='langchain_core._api')
        # Ensure the output folder exists
        os.makedirs(self.output_folder, exist_ok=True)
        os.makedirs(self.output_analysis_folder, exist_ok=True)
        os.makedirs(self.mistake_database, exist_ok=True)


        # process the material and match the knowledge points of the exam paper
        material = process_material(self.material_folder)
        # paper = read_docx_to_string_with_format(input_paper)
        # match_result = self.knowledge_points_match(paper, material, "exam paper")

        # Iterate over all .docx files in the input folder
        # for each student, match knowledge points
        # for filepath in glob.glob(os.path.join(self.input_folder, "*.docx")):
        #     # Get the filename without the extension
        #     filename = os.path.splitext(os.path.basename(filepath))[0]
            
        #     start_time = time.time()

        #     rows_ = process_paper_and_store_results(input_paper, correct_answers_path, filepath)
        #     rows = rows_[0]
        #     mistake_count = rows_[1]
        #     ### 这里打印出来看看
        #     output_doc = Document()
        #     # Iterate through each row and add it to the document
        #     for row in rows:
        #         # Assuming each row is a tuple; format it as desired (e.g., join elements with a separator)
        #         mistake_entry = ', '.join(str(item) for item in row)  # Convert each item to string and join
        #         output_doc.add_paragraph(mistake_entry)

        #     # 路径修改
        #     output_path = os.path.join(self.mistake_database, f"{filename}_mistake_database.docx")
        #     output_doc.save(output_path)
        #     print(f"Completed storing {filename} mistake database.")
            

        #     problem_list =[]
        #     for item in rows:
        #         problem_list.append(item[2])
            
        #     matching_result = self.knowledge_points_match(' '.join(problem_list), material, filename)
        #     end_time = time.time()
        #     print(f"Completed knowledge points match of {filename} in: {end_time - start_time:.2f} seconds")

        #     # 还未把match好的知识点放到knowledge point analysis 里
        #     self.knowledge_point_analysis(' '.join(problem_list), filename, sample_analysis)
        #     # paper_revise还要修改
        #     self.paper_revise(' '.join(problem_list), matching_result, material, mistake_count, filename)

        #     end_time = time.time()

        #     print(f"Completed revising {filename} in: {end_time - start_time:.2f} seconds")

        # Iterate over all the new question files and fix them
        for filepath in glob.glob(os.path.join(self.output_folder, "*.docx")):

            filename = os.path.splitext(os.path.basename(filepath))[0]
            
            start_time = time.time()

            new_que = read_docx_to_string_with_format(filepath)
            self.question_revise(new_que, filename)

            end_time = time.time()
            print(f"Completed revising new questions {filename} in: {end_time - start_time:.2f} seconds")



# def process_newquestion_and_store_results(new_question_path):
#     filename = os.path.splitext(os.path.basename(new_question_path))[0]
def process_material(meterial_folder):
    material_doc = []
    for filepath in glob.glob(os.path.join(meterial_folder, "*.docx")):
        material_doc.append(read_docx_to_string_with_format(filepath))
    return material_doc

def clear_folder(filepath):
    listdir = os.listdir(filepath)  # 获取文件和子文件夹
    for dirname in listdir:
        dirname = filepath + "//" + dirname
        if os.path.isfile(dirname): # 是文件
            os.remove(dirname)
def process_paper_and_store_results(question_path, right_answer_path, wrong_answer_path):
    """
    Processes a student's answer paper, compares it with the correct answers, stores the results in a database, and retrieves the student's mistakes.

    :param question_path: Path to the question paper file.
    :param right_answer_path: Path to the correct answers file.
    :param wrong_answer_path: Path to the student's wrong answers file.
    :return: A list of the student's mistakes.
    """
    # 获取文件名和学生ID
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_id = extract_student_id(wrong_answer_path)

    # 获取修正后的结果
    answer = return_revised_result(question_path, right_answer_path, wrong_answer_path, filename)
    revised_problem_answer_list = answer[0]
    right_or_not = answer[1]
    mistake_count = answer[2]

    cursor = db.cursor()
    
    # 删除已经存在的表并创建新表(后面正式开始运行时注释掉下面第一行)
    cursor.execute(drop_table_query)
    cursor.execute(create_table_query)
    
    # 插入数据到数据库
    for i in range(len(right_or_not)):
        cursor.execute(insert_query, (student_id, revised_problem_answer_list[i], right_or_not[i]))
    db.commit()
    
    # 查询学生的错误
    cursor.execute(select_mistake_query, (student_id,))
    rows = cursor.fetchall()

    return rows,mistake_count




def main():

    material_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\TDLEG learning materials\\materials"
    # input_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Test1\\student paper"
    # output_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Test1\\New Paper"
    # output_mistakes_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Test1\\Student Mistakes"
    # output_analysis_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Test1\Knowledge Point Analysis"
    # correct_answers_path = "C:\\Users\\30998\Desktop\\JAP_GPT\\template paper from CUHK\Test1\\test 1 paper\\Test 1 Model Answer.docx"
    # input_paper = "C:\\Users\\30998\Desktop\\JAP_GPT\\template paper from CUHK\\Test1\\test 1 paper\\Test 1 Question Paper.docx"
    # sample_mistake_analysis = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Test1\\1155159595 Test 1_sample_mistakes_analysis.doc"
    # Revised_newpaper_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Test1\\Revised_newpaper_folder"
    # Mistake_database = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Test1\\Mistake_database"

    # sample, just for check function
    input_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\sample\\input"
    output_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\sample\\new paper"
    output_mistakes_folder = "C:\\Users\\30998\\Desktop\\template paper from CUHK\\JAP_GPT\\Jap_GPT_hk\\sample\\mistake"
    output_analysis_folder = "C:\\Users\\30998\\Desktop\\template paper from CUHK\\JAP_GPT\\Jap_GPT_hk\\sample\\analysis"
    correct_answers_path = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\test 1 paper\\Test 1 Model Answer.docx"
    input_paper = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\test 1 paper\\Test 1 Question Paper.docx"
    sample_mistake_analysis = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\1155159595 Test 1_sample_mistakes_analysis.doc"
    Revised_newpaper_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\sample\\revised new paper"
    Mistake_database = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\sample\\Mistake_database"
    matched_knowledge_points_folder = "C:\\Users\\30998\\Desktop\\JAP_GPT\\template paper from CUHK\\Jap_GPT_hk\\sample\\matched knowledge points"

    # clear_folder(output_folder)
    # clear_folder(Revised_newpaper_folder)

    checker = AnswerChecker(correct_answers_path, input_folder, output_mistakes_folder)
    checker.process_all_files()
    processor = DocumentProcessor(input_folder, output_folder, output_analysis_folder,Revised_newpaper_folder, Mistake_database, material_folder, matched_knowledge_points_folder)
    processor.process(input_paper, correct_answers_path, sample_mistake_analysis)


    



if __name__ == "__main__":
    main()
