import re
import os
import glob
import time
import warnings
import docx
import mysql.connector
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate  
from langchain.chains import LLMChain   
from typing import Any


def split_into_sentences(text):
        """
        Splits text into sentences based on common Japanese sentence endings.
        
        :param text: Text to split.
        :return: List of sentences.
        """
        sentence_endings = re.compile(r'(?<=[。！？])\s*')
        sentences = sentence_endings.split(text)
        return sentences

def knowledge_point_analysis(paper, question_type):
    """
    Revises a list using a language model and saves the results.

    :param rows: Loaded list data to revise.
    :param filename: Path to save the revised document.
    """
    llm = ChatOpenAI(
        temperature=0.8,
        model="gpt-4o"
    ) 
    prompt_two = ChatPromptTemplate.from_template(
        "I have provided a sample of questions Japanese Language Proficiency Test N4 level: {question_paper}.\
        Now base on these questions, detailed analyze the types of questions"
    )
    
    chain_one = LLMChain(llm=llm, prompt=prompt_two)


    inputs_two = {'question_paper': paper}
    revise_result = chain_one.run(inputs_two)
    output_doc = Document()
    sentences = split_into_sentences(revise_result)
    for sentence in sentences:
        output_doc.add_paragraph(sentence)

    output_path = os.path.join(question_type, f"question_type.docx")
    output_doc.save(output_path)

def main():
    paper = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\JAP_GPT\\Test1_new\\test 1 paper\\Test 1 Question Paper.docx"
    question_type = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\JAP_GPT\\Test1_new\\test 1 paper"
    document = Document(paper)
    text = ""
    for paragraph in document.paragraphs:
        text += paragraph.text
    knowledge_point_analysis(text,question_type)


if __name__ == "__main__":
    main()