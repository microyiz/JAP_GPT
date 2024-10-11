import re
import os
import glob
import time
import warnings
import docx
import mysql.connector
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate  
from langchain.chains import LLMChain   
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from typing import Any

from jap_paper_revise import return_revised_result
from jap_paper_revise import extract_student_id
from jap_paper_revise import read_docx_to_string
from db_util import drop_table_query ,create_table_query ,insert_query,show_fiverows_query,select_mistake_query,db


class AnswerChecker:
    """
    A class to check student answers against correct answers, generate a mistakes report, and save it to a specified output folder.
    """
    def __init__(self, correct_answers_path, input_folder, output_folder):
        """
        Initializes the AnswerChecker with paths to the correct answers, input folder, and output folder.

        :param correct_answers_path: Path to the document containing the correct answers.
        :param input_folder: Path to the folder containing student answer documents.
        :param output_folder: Path to the folder where the mistakes reports will be saved.
        """
        self.correct_answers_path = correct_answers_path
        self.input_folder = input_folder
        self.output_folder = output_folder
        os.makedirs(self.output_folder, exist_ok=True)
    
    def read_answers(self, doc_path):
        """
        Reads and extracts answers from a given document.

        :param doc_path: Path to the document containing answers.
        :return: Two dictionaries containing answers from Part 1 and Part 2 of the test.
        """
        doc = docx.Document(doc_path)
        answers_part1 = {}
        answers_part2 = {}
        current_part = None

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "日本語実力テスト1 (第1部" in text:
                current_part = "part1"
            elif "日本語実力テスト1 (第2部" in text:
                current_part = "part2"
            elif text.startswith("問題"):
                question, answer = text.split("：")
                question_num = question.split("　")[1]
                if current_part == "part1":
                    answers_part1[question_num] = answer
                elif current_part == "part2":
                    answers_part2[question_num] = answer

        return answers_part1, answers_part2

    def save_mistakes_to_docx(self, mistakes, output_filename, filename):
        """
        Saves identified mistakes to a DOCX file.

        :param mistakes: A dictionary of mistakes identified in the student's answers.
        :param output_filename: Path where the DOCX file will be saved.
        :param filename: The base name of the student's document (used for report title).
        """
        doc = docx.Document()
        doc.add_heading(f"{filename}の誤答レポート", level=1)

        k=1
        
        for part, mistakes_part in mistakes.items():
            doc.add_heading(part, level=2)
            for question, (student_answer, correct_answer) in mistakes_part.items():
                doc.add_paragraph(f"{k}." f"問題 {question}:")
                doc.add_paragraph(f"学生の回答: {student_answer}", style='List Bullet')
                doc.add_paragraph(f"正解: {correct_answer}", style='List Bullet')
                doc.add_paragraph("\n")
                k=k+1

        doc.save(output_filename)

    def compare_answers_and_generate_report(self, student_answers_path):
        """
        Compares student answers with correct answers, identifies mistakes, and generates a report.

        :param student_answers_path: Path to the student's answers document.
        """
        # Read answers from both documents
        student_answers_part1, student_answers_part2 = self.read_answers(student_answers_path)
        correct_answers_part1, correct_answers_part2 = self.read_answers(self.correct_answers_path)

        # Compare the answers and identify mistakes
        mistakes_part1 = {}
        for question, correct_answer in correct_answers_part1.items():
            student_answer = student_answers_part1.get(question, None)
            if student_answer != correct_answer:
                mistakes_part1[question] = (student_answer, correct_answer)

        mistakes_part2 = {}
        for question, correct_answer in correct_answers_part2.items():
            student_answer = student_answers_part2.get(question, None)
            if student_answer != correct_answer:
                mistakes_part2[question] = (student_answer, correct_answer)

        # Get the filename without the extension
        filename = os.path.splitext(os.path.basename(student_answers_path))[0]

        # Save the mistakes to a DOCX file
        mistakes = {"Part 1": mistakes_part1, "Part 2": mistakes_part2}
        mistakes_output_filename = os.path.join(self.output_folder, f"{filename}_mistakes.docx")
        self.save_mistakes_to_docx(mistakes, mistakes_output_filename, filename)
        print(f"Mistakes report saved to {mistakes_output_filename}")

    def process_all_files(self):
        """
        Processes all .docx files in the input folder, compares the answers, and saves the mistakes reports in the output folder.
        """
        # Iterate over all .docx files in the input folder
        for filepath in glob.glob(os.path.join(self.input_folder, "*.docx")):
            self.compare_answers_and_generate_report(filepath)

    


class DocumentProcessor:
    def __init__(self, input_folder, output_folder, output_analysis_folder, revised_newpaper_folder, mistake_database):
        """
        Initializes the DocumentProcessor with input and output folders.
        
        :param input_folder: Folder containing the input Word documents.
        :param output_folder: Folder to save the new Word documents.
        :param output_analysis_folder: Folder to save the processed knowledge point Word documents.
        """
        self.input_folder = input_folder
        self.output_folder = output_folder
        self.output_analysis_folder = output_analysis_folder
        self.revised_newpaper_folder = revised_newpaper_folder
        self.mistake_database = mistake_database

    def load_document(self, filepath):
        """
        Loads a Word document from the specified filepath.
        
        :param filepath: Path to the Word document.
        :return: Loaded document data.
        """
        loader = UnstructuredWordDocumentLoader(filepath)
        data = loader.load()
        return data
    
    def append_splits(self, splits, output_path):
        """
        Appends the split document chunks into a new Word document.
        
        :param splits: List of document chunks.
        :param output_path: Path to save the new Word document.
        """
        doc = Document()
        for split in splits:
            doc.add_paragraph(split.page_content)
        doc.save(output_path)

    def split_into_sentences(self, text):
        """
        Splits text into sentences based on common Japanese sentence endings.
        
        :param text: Text to split.
        :return: List of sentences.
        """
        sentence_endings = re.compile(r'(?<=[。！？])\s*')
        sentences = sentence_endings.split(text)
        return sentences


    def paper_revise(self, rows, filename):
        """
        Revises a list using a language model and saves the results.

        :param rows: Loaded list data to revise.
        :param filename: Path to save the revised document.
        """
        llm = ChatOpenAI(
            temperature=0.8,
            model="gpt-4o"
        )

        prompt_one = ChatPromptTemplate.from_template(
            "Below is a list of incorrect answers provided by Japanese language students: {error_report}\n"
            "Each question includes the student's incorrect choice and the correct answer.\n"
            "Based on these errors, generate new practice questions targeting similar grammar or vocabulary points to help students strengthen their understanding.\n"
            "The new questions should be in a multiple-choice format and appropriate for the Japanese Language Proficiency Test N4 level.\n"
            "Please create {num_questions} new questions, each with four different options. Ensure that only one of these options is correct and should be evenly distributed among 1, 2, 3, and 4.\n"
            "Finally, all the answers will be attached at the end. Do not attach the answer after each question."
        )
        
        chain_one = LLMChain(llm=llm, prompt=prompt_one)

        inputs_one = {
            'error_report': rows,
            'num_questions': 20
        }
        revise_result = chain_one.run(inputs_one)
        output_doc = Document()
        sentences = self.split_into_sentences(revise_result)
        for sentence in sentences:
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(self.output_folder, f"{filename}_new_report.docx")
        output_doc.save(output_path)

    
    def knowledge_point_analysis(self, rows, filename, sample_analysis):
        """
        Revises a list using a language model and saves the results.

        :param rows: Loaded list data to revise.
        :param filename: Path to save the revised document.
        """
        llm = ChatOpenAI(
            temperature=0.8,
            model="gpt-4o"
        ) 
        prompt_two = ChatPromptTemplate.from_template(
            "I have provided a sample analysis of a student's mistakes in a Japanese practice test below, labeled as {sample}. \
            The analysis is organized into two main sections: 1.1 Kanji/Vocabulary related mistakes and 1.2 Grammar mistakes. \
            Each section is further divided into smaller sub-sections, such as Pronunciation mistake, Long vowel and short vowel pronunciation mistake, etc. \
            Every sub-section summarizes the specific knowledge points where the student made errors. This format is crucial."

            "Now, I will provide you with another student's error report, labeled as {error_report}. \
            Please analyze the mistakes made by this student using the same structure and detail as in the {sample}. \
            Ensure that the question numbers in {error_report} are retained in the generated analysis.\
            The analysis should be comprehensive and organized into appropriate sections and sub-sections based on the knowledge points involved. \
            Use specific question numbers and describe the errors in a similar manner. Please attach every student's mistakes to the related specific knowledge points."
        )
        
        chain_one = LLMChain(llm=llm, prompt=prompt_two)

        inputs_two = {
            'error_report': rows,
            'sample': sample_analysis
        }
        revise_result = chain_one.run(inputs_two)
        output_doc = Document()
        sentences = self.split_into_sentences(revise_result)
        for sentence in sentences:
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(self.output_analysis_folder, f"{filename}_mistakes_analysis.docx")
        output_doc.save(output_path)


    def question_revise(self, rows, filename, max_iterations=5):
        """
        Revise the generated questions using llm.
        """
        llm = ChatOpenAI(
            temperature=0.6,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt_three = ChatPromptTemplate.from_template(
            "Now these are the new generated Japanese practice questions: {new_paper} \
            Please revise these questions to check: \
            1. Are there multiple correct answers for the question options? \
            2. Are there any duplicate questions? If so, replace the duplicate one with a new question.\
            3. Are there any errors in the question stem? \
            4. Are there any duplicate options for one question? If so, change the options to avoid that.\
            If the above problems occur, please modify the questions so that they do not have the above problems. \
            The structure should be same with original questions, all the answers will be attached at the end. Do not attach the answer after each question. \
            Report the changes made at last of the file."
        )

        chain_three = LLMChain(llm=llm, prompt = prompt_three)
        input_three = {'new_paper': rows}
        
        for iteration in range(max_iterations):
            revise_result = chain_three.run(input_three)
            if self.check_for_errors(revise_result):  # Implement this method to validate output
                input_three['new_paper'] = revise_result
            else:
                break  # Exit loop if no errors are found


        output_doc = Document()
        sentences = self.split_into_sentences(revise_result)
        for sentence in sentences:
            output_doc.add_paragraph(sentence)

        # 路径修改
        output_path = os.path.join(self.revised_newpaper_folder, f"{filename}_revised_new_paper.docx")
        output_doc.save(output_path)


    # def check_for_errors(self, revise_result):
    #     questions = revise_result.split('\n')

    #     for question_text in questions:
    #         question = {
    #             'text': question_text,
    #             'options': ['1', '2', '3', '4'],
    #             'correct_answer': 'correct_answer'
    #         }
    #         # Iterate through each question and check for errors
    #         if self.check_for_multiple_correct_answers(question):
    #             return True  # Error found
    #         if self.check_for_duplicate_questions(revise_result['questions']):
    #             return True  # Error found
    #         # if self.check_for_stem_errors(question):
    #         #     return True  # Error found
    #         if self.check_for_duplicate_options(question):
    #             return True  # Error found

    #     return False  # No errors found
    
    # def check_for_multiple_correct_answers(self, question):
    #     correct_answer = question['correct_answer']
    #     answer_options = question['options']
    
    #     # Count how many times the correct answer appears in the options
    #     correct_count = sum([1 for option in answer_options if option == correct_answer])
    
    #     return correct_count > 1  # Returns True if multiple correct answers are found
    
    # def check_for_duplicate_questions(self, questions):
    #     seen_questions = set()
    #     for question in questions:
    #         question_text = question['text']
    #         if question_text in seen_questions:
    #             return True  # Duplicate found
    #         seen_questions.add(question_text)
    #     return False
    
    # def check_for_stem_errors(self, question):
    #     question_text = question['text']
    #     # Example of basic error check: check for missing punctuation at the end
    #     if not re.search(r'[。！？?]$', question_text):  # Must end with proper punctuation
    #         return True  # Stem error found
    
    #     # You can add other rules to check for incomplete sentences or grammatical issues
    #     return False

    # def check_for_duplicate_options(self, question):
    #     answer_options = question['options']
    
    #     # Convert list to a set and compare lengths to detect duplicates
    #     if len(answer_options) != len(set(answer_options)):
    #         return True  # Duplicate options found
    
    #     return False




    def check_for_errors(self, revised_text):
        """
        Check for errors in the revised question set, including:
        - Multiple correct answers
        - Duplicate questions
        - Errors in the question stem
        - Duplicate options
        
        :param revised_text: The revised text output from GPT.
        :return: True if errors are found, False otherwise.
        """

        # Error check 1: Multiple correct answers
        if self.has_multiple_correct_answers(revised_text):
            return True

        # Error check 2: Duplicate questions
        if self.has_duplicate_questions(revised_text):
            return True

        # Error check 3: Errors in the question stem
        if self.has_stem_errors(revised_text):
            return True

        # Error check 4: Duplicate options
        if self.has_duplicate_options(revised_text):
            return True

        return False  # No errors detected
    
    # def has_multiple_correct_answers(self, text):
    #     """
    #     Check for multiple correct answers in the text.
    #     This could be adapted by parsing the question and answers, or using regex patterns.
        
    #     :param text: The text to check.
    #     :return: True if multiple correct answers are detected.
    #     """
    #     # This is an example and may need further refinement based on the format.
    #     # Assuming correct answers are marked or follow a specific format:
    #     correct_answers = re.findall(r'\[correct answer: (\w+)\]', text, re.IGNORECASE)
    #     if len(set(correct_answers)) > 1:  # More than one unique correct answer found
    #         return True
    #     return False
    
    def has_multiple_correct_answers(self, text):
        # 还没想好怎么改
        pass

    def has_duplicate_questions(self, text):
        """
        Check if any questions are duplicated.
        
        :param text: The text to check.
        :return: True if duplicate questions are detected.
        """
        questions = re.findall(r'\d+\.\s*(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n', 
                               text, re.DOTALL
                            )
        seen_questions = set()
        for question in questions:
            question_text = question[0].strip()  # Get the question part only
            if question_text in seen_questions:
                return True  # Duplicate found
            seen_questions.add(question)
        return False

    def has_stem_errors(self, text):
        """
        Check for errors in the question stem.
        
        :param text: The text to check.
        :return: True if errors in the stem are detected.
        """
        # 只是一个最简单的检查是否有标点符号的检查，需要修改为检查题目内容？
        # Example: Detect missing or malformed question stems.
        questions = re.findall(r'\d+\.\s*(.*?)\n', text)
        for question in questions:
            if not question.strip() or len(question.strip().split()) < 5:  # Example check
                return True  # Detected a malformed stem
        return False

    def has_duplicate_options(self, text):
        """
        Check for duplicate options in the questions.
        
        :param text: The text to check.
        :return: True if duplicate options are found within a question.
        """
        # Example: Detect duplicate options for a given question.
        questions_with_options = re.findall(
            r'(\d+)\.\s*(.*?)\n(1\.\s*(.*?)\n)(2\.\s*(.*?)\n)(3\.\s*(.*?)\n)(4\.\s*(.*?)\n)',
            text, re.DOTALL
        )

        for question, opt1, _, opt2, _, opt3, _, opt4, _ in questions_with_options:
            options = {opt1.strip(), opt2.strip(), opt3.strip(), opt4.strip()}
            if len(options) < 4:  # If any options are duplicates
                print(f"Duplicate options detected in question {question}: {opt1.strip()}, {opt2.strip()}, {opt3.strip()}, {opt4.strip()}")
                return True
        return False



    def process(self, input_paper, correct_answers_path, sample_analysis):
        """
        Processes all .docx files in the input folder and saves the results in the output folder.

        :param input_paper: Question paper.
        :param correct_answers_path: Answer to the question paper.
        """
        warnings.filterwarnings("ignore", category=DeprecationWarning, module='langchain_core._api')
        # Ensure the output folder exists
        os.makedirs(self.output_folder, exist_ok=True)
        os.makedirs(self.output_analysis_folder, exist_ok=True)
        os.makedirs(self.mistake_database, exist_ok=True)


        # Iterate over all .docx files in the input folder
        for filepath in glob.glob(os.path.join(self.input_folder, "*.docx")):
            # Get the filename without the extension
            filename = os.path.splitext(os.path.basename(filepath))[0]
            
            start_time = time.time()

            rows = process_paper_and_store_results(input_paper, correct_answers_path, filepath)
            ### 这里打印出来看看
            output_doc = Document()
            # Iterate through each row and add it to the document
            for row in rows:
                # Assuming each row is a tuple; format it as desired (e.g., join elements with a separator)
                mistake_entry = ', '.join(str(item) for item in row)  # Convert each item to string and join
                output_doc.add_paragraph(mistake_entry)

            # 路径修改
            output_path = os.path.join(self.mistake_database, f"{filename}_mistake_database.docx")
            output_doc.save(output_path)
            print(f"Completed storing {filename} mistake database.")
            

            problem_list =[]
            for item in rows:
                problem_list.append(item[2])
            
            self.knowledge_point_analysis(' '.join(problem_list), filename, sample_analysis)
            # paper_revise还要修改
            self.paper_revise(' '.join(problem_list), filename)

            end_time = time.time()

            print(f"Completed revising {filename} in: {end_time - start_time:.2f} seconds")

        # Iterate over all the new question files and fix them
        for filepath in glob.glob(os.path.join(self.output_folder, "*.docx")):

            filename = os.path.splitext(os.path.basename(filepath))[0]
            
            start_time = time.time()

            new_que = read_docx_to_string(filepath)
            self.question_revise(new_que, filename)

            end_time = time.time()
            print(f"Completed revising new questions {filename} in: {end_time - start_time:.2f} seconds")



# def process_newquestion_and_store_results(new_question_path):
#     filename = os.path.splitext(os.path.basename(new_question_path))[0]



def process_paper_and_store_results(question_path, right_answer_path, wrong_answer_path):
    """
    Processes a student's answer paper, compares it with the correct answers, stores the results in a database, and retrieves the student's mistakes.

    :param question_path: Path to the question paper file.
    :param right_answer_path: Path to the correct answers file.
    :param wrong_answer_path: Path to the student's wrong answers file.
    :return: A list of the student's mistakes.
    """
    # 获取文件名和学生ID
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_id = extract_student_id(wrong_answer_path)

    # 获取修正后的结果
    answer = return_revised_result(question_path, right_answer_path, wrong_answer_path, filename)
    revised_problem_answer_list = answer[0]
    right_or_not = answer[1]

    cursor = db.cursor()
    
    # 删除已经存在的表并创建新表(后面正式开始运行时注释掉下面第一行)
    cursor.execute(drop_table_query)
    cursor.execute(create_table_query)
    
    # 插入数据到数据库
    for i in range(len(right_or_not)):
        cursor.execute(insert_query, (student_id, revised_problem_answer_list[i], right_or_not[i]))
    db.commit()
    
    # 查询学生的错误
    cursor.execute(select_mistake_query, (student_id,))
    rows = cursor.fetchall()

    return rows




def main():
    input_folder = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\student paper"
    output_folder = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\New Paper"
    output_mistakes_folder = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\Student Mistakes"
    output_analysis_folder = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\Knowledge Point Analysis"
    correct_answers_path = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\test 1 paper\\Test 1 Model Answer.docx"
    input_paper = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\test 1 paper\\Test 1 Question Paper.docx"
    sample_mistake_analysis = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\1155159595 Test 1_sample_mistakes_analysis.doc"
    Revised_newpaper_folder = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\Revised_newpaper_folder"
    Mistake_database = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\Test1_new\\Mistake_database"
    

    checker = AnswerChecker(correct_answers_path, input_folder, output_mistakes_folder)
    checker.process_all_files()
    processor = DocumentProcessor(input_folder, output_folder, output_analysis_folder, Revised_newpaper_folder, Mistake_database)
    processor.process(input_paper, correct_answers_path, sample_mistake_analysis)
    



if __name__ == "__main__":
    main()
