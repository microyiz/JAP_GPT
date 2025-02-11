import mysql.connector
from docx import Document
from docx.shared import Inches
import re
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
import os

def get_latest_paper_id(test_paper):
    files = [f for f in os.listdir("D:\\JAP_GPT\\JAP_GPT\\newly_generated_papers") if f.startswith(f"{test_paper}")]
    if not files:
        return None
    full_paths = [os.path.join("D:\\JAP_GPT\\JAP_GPT\\newly_generated_papers", f) for f in files]
    return max(full_paths, key=os.path.getmtime)


def sort_key(item):
    #知识点排序（语法知识点按数字排序，词汇知识点按字母排序）
    level, knowledge_point = item[0]
    try:
        num = int(knowledge_point.split('.')[0])
        return (0, level, num)
    except:
        return (1, level, knowledge_point)


def Word_Document(test_paper, question_vocabulary_type1, question_vocabulary_type2, question_vocabulary_type3, question_vocabulary_type4, question_vocabulary_type5, question_grammar_type1, save_path):

    def title_Paragraph(doc, text, size, color):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.text = text
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(color[0], color[1], color[2])
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS PGothic')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.style = 'Normal'

    def set_run_black(run):
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS PGothic')
        run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def set_run_purple(run):
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128, 0, 128)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'MS PGothic')
        run.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    doc = Document()

    def Part_one_Starting_Sentence():
        title_Paragraph(doc, "日本語実力テスト1 (第1部：文字と語彙)", 13.5, (128, 0, 128))
        title_Paragraph(doc, "Japanese Language Level Checking Test 1", 13.5, (128, 0, 128))
        title_Paragraph(doc, "(Part 1: Characters and Vocabulary)", 13.5, (128, 0, 128))
        doc.add_paragraph("")
        title_Paragraph(doc, "Name:", 13.5, (128, 0, 128))
        doc.add_paragraph("")
        title_Paragraph(doc, "Student I.D. Number:", 13.5, (128, 0, 128))
        doc.add_paragraph("")
        title_Paragraph(doc, "Select the most appropriate answer for the underlined part in each question.", 12, (128, 0, 128))
        title_Paragraph(doc, "There is an example in each section for your reference.", 12, (128, 0, 128))
        doc.add_paragraph("")

    def Part_two_Starting_Sentence():
        title_Paragraph(doc, "日本語実力テスト1 (第2部：文法)", 13.5, (128, 0, 128))
        title_Paragraph(doc, "Japanese Language Level Checking Test 1 (Part 2: Grammar)", 13.5, (128, 0, 128))
        doc.add_paragraph("")
        title_Paragraph(doc, "Name:", 13.5, (128, 0, 128))
        doc.add_paragraph("")
        title_Paragraph(doc, "Student I.D. Number:", 13.5, (128, 0, 128))
        doc.add_paragraph("")
        title_Paragraph(doc, "Select the most appropriate answer for the underlined part in each question.", 12, (128, 0, 128))
        doc.add_paragraph("")

    def example(stem, options, answer):
        title_Paragraph(doc, "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿", 12, (0, 0, 0))
        paragraph = doc.add_paragraph()
        run1 = paragraph.add_run("()")
        run2 = paragraph.add_run("(Example)")
        run3 = paragraph.add_run(f"　{stem}")
        set_run_black(run1)
        set_run_purple(run2)
        set_run_black(run3)
        title_Paragraph(doc, f"{options}" , 12, (0, 0, 0))
        paragraph = doc.add_paragraph()
        run1 = paragraph.add_run("答え ")
        run2 = paragraph.add_run("(Answer)")
        run3 = paragraph.add_run(f"：　{answer}")
        set_run_black(run1)
        set_run_purple(run2)
        set_run_black(run3)
        title_Paragraph(doc, "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿", 12, (0, 0, 0))
        doc.add_paragraph("")
    
    list_index = ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩', '⑪', '⑫', '⑬', '⑭', '⑮', '⑯', '⑰', '⑱', '⑲', '⑳','㉑', '㉒', '㉓', '㉔', '㉕', '㉖', '㉗', '㉘', '㉙', '㉚', '㉛', '㉜', '㉝', '㉞', '㉟', '㊱', '㊲', '㊳', '㊴', '㊵', '㊶', '㊷', '㊸', '㊹', '㊺', '㊻', '㊼', '㊽', '㊾', '㊿']

    if (question_vocabulary_type1 != [] or question_vocabulary_type2 != [] or question_vocabulary_type3 != [] or question_vocabulary_type4 != [] or question_vocabulary_type5 != []):
        Part_one_Starting_Sentence()
        
        number = 1
        type = 1
        if question_vocabulary_type1 != []:
            paragraph = doc.add_paragraph()
            paragraph.style = 'Normal'
            bold_run = paragraph.add_run(f"もんだい{type}")
            bold_run.bold = True
            set_run_black(bold_run)
            unbold_run = paragraph.add_run("　＿＿＿の　ことばは　ひらがなで　どう　かきますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。")
            set_run_black(unbold_run)
            doc.add_paragraph("")
            type += 1
            
            example("あそこに　かわいい　鳥が　います。", "1　いぬ		2　とり			3　ねこ			4　むし", "2　とり")

            for question in question_vocabulary_type1:
                circle_number = list_index[number - 1]
                paragraph = doc.add_paragraph(f"{circle_number} {question}")
                run = paragraph.runs[0]
                set_run_black(run)
                doc.add_paragraph("")
                number += 1
        
            doc.add_paragraph("")
            doc.add_paragraph("")
            
        if question_vocabulary_type2 != []:
            paragraph = doc.add_paragraph()
            paragraph.style = 'Normal'
            bold_run = paragraph.add_run(f"もんだい{type}")
            bold_run.bold = True
            set_run_black(bold_run)
            unbold_run = paragraph.add_run("　＿＿＿の　ことばは　どう　かきますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。")
            set_run_black(unbold_run)
            doc.add_paragraph("")
            type += 1
    
            example("おっとは　今、出かけています。", "1　大			2　犬			3　太            4　夫", "4　夫")

            for question in question_vocabulary_type2:
                circle_number = list_index[number - 1]
                paragraph = doc.add_paragraph(f"{circle_number} {question}")
                run = paragraph.runs[0]
                set_run_black(run)
                doc.add_paragraph("")
                number += 1

            doc.add_paragraph("")
            doc.add_paragraph("")
            
        if question_vocabulary_type3 != []:
            paragraph = doc.add_paragraph()
            paragraph.style = 'Normal'
            bold_run = paragraph.add_run(f"もんだい{type}")
            bold_run.bold = True
            set_run_black(bold_run)
            unbold_run = paragraph.add_run("　(   　  ) に　　なにを　いれますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。")
            set_run_black(unbold_run)
            doc.add_paragraph("")
            type += 1

            example("これから　ひこうきに　（  　　　　　 ）。", "1　おります		2　のります		3　あがります		4　のぼります", "2　のります")

            for question in question_vocabulary_type3:
                circle_number = list_index[number - 1]
                paragraph = doc.add_paragraph(f"{circle_number} {question}")
                run = paragraph.runs[0]
                set_run_black(run)
                doc.add_paragraph("")
                number += 1

            doc.add_paragraph("")
            doc.add_paragraph("")
            
        if question_vocabulary_type4 != []:
            paragraph = doc.add_paragraph()
            paragraph.style = 'Normal'
            bold_run = paragraph.add_run(f"もんだい{type}")
            bold_run.bold = True
            set_run_black(bold_run)
            unbold_run = paragraph.add_run("　＿＿＿の　ぶんと　だいたい　おなじ　いみの　ぶんが　あります。　1・2・3・4から　ひとつ　えらんで　ください。")
            set_run_black(unbold_run)
            doc.add_paragraph("")
            type += 1

            example("　ギターは　ちちに　ならいました。", "1　ギターは　ちちに　もらいました。\n2　ギターは　ちちに　えらんでもらいました。\n3　ギターは　ちちに　おしえてもらいました。\n4　ギターは　ちちに　かってもらいました。", "3　ギターは　ちちに　おしえてもらいました。")

            for question in question_vocabulary_type4:
                circle_number = list_index[number - 1]
                paragraph = doc.add_paragraph(f"{circle_number} {question}")
                run = paragraph.runs[0]
                set_run_black(run)
                doc.add_paragraph("")
                number += 1
            
            doc.add_paragraph("")
            doc.add_paragraph("")
            

        if question_vocabulary_type5 != []:
            paragraph = doc.add_paragraph()
            paragraph.style = 'Normal'
            bold_run = paragraph.add_run(f"もんだい{type}")
            bold_run.bold = True
            set_run_black(bold_run)
            unbold_run = paragraph.add_run("　つぎの　ことばの　つかいかたで　いちばん　いい　ものを　1・2・3・4から　ひとつ　えらんで　ください。")
            set_run_black(unbold_run)
            doc.add_paragraph("")
            type += 1

            example("　わる", "1　おさらを　わって　母に　おこられました。\n2　おさらを　わって　へやに　かざりました。\n3　おさらを　わって　りょうりを　つくりました。\n4　おさらを　わって　コーヒーを　のみました。", "1　おさらを　わって　母に　おこられました。")

            for question in question_vocabulary_type5:
                circle_number = list_index[number - 1]
                paragraph = doc.add_paragraph(f"{circle_number} {question}")
                run = paragraph.runs[0]
                set_run_black(run)
                doc.add_paragraph("")
                number += 1

            doc.add_paragraph("")
            doc.add_paragraph("")
            
        
        title_Paragraph(doc,"End of Part 1.  Thank you!  Please continue to complete Part 2.  ", 13.5, (128, 0, 128))
        doc.add_page_break()

    if question_grammar_type1 != []:

        Part_two_Starting_Sentence()

        paragraph = doc.add_paragraph()
        paragraph.style = 'Normal'
        bold_run = paragraph.add_run("もんだい1")
        bold_run.bold = True
        set_run_black(bold_run)
        unbold_run = paragraph.add_run("　（  　　　　　 ）に　何を　入れますか。　1・2・3・4から　いちばん　いい　ものを　一つ　えらんで　ください。")
        set_run_black(unbold_run)
        doc.add_paragraph("")

        number = 1
                        
        for question in question_grammar_type1:
            circle_number = list_index[number - 1]
            paragraph = doc.add_paragraph(f"{circle_number} {question}")
            run = paragraph.runs[0]
            set_run_black(run)
            doc.add_paragraph("")
            number += 1

        title_Paragraph(doc,"End of Part 2.  Thank you very much for your participation!   ", 12, (128, 0, 128))

    for paragraph in doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(8)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(f"D:\\JAP_GPT\\JAP_GPT\\newly_generated_papers\\{save_path}.docx")


def questions_preprocess(test_paper, questions, save_path):

    question_grammar_type1 = []
    question_vocabulary_type1 = []
    question_vocabulary_type2 = []
    question_vocabulary_type3 = []
    question_vocabulary_type4 = []
    question_vocabulary_type5 = []

    for question in questions:
        # 处理题干重复问题
        if question[2].startswith("1　（  　　　　　 ）に　何を　入れますか。　1・2・3・4から　いちばん　いい　ものを　一つ　えらんで　ください。"):
            stem = re.sub(r'^1　（  　　　　　 ）に　何を　入れますか。　1・2・3・4から　いちばん　いい　ものを　一つ　えらんで　ください。\n', '', question[2])
            if re.search(r'\s*\n+End of Part', stem):
                stem = re.split(r'\s*End of Part\s*', stem)[0]
                question_grammar_type1.append(stem)
            else:
                question_grammar_type1.append(stem)

        if question[2].startswith("1　＿＿＿の　ことばは　ひらがなで　どう　かきますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。"):
            stem = re.sub(r'^1　＿＿＿の　ことばは　ひらがなで　どう　かきますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。\n','', question[2])
            question_vocabulary_type1.append(stem)

        if question[2].startswith("２　＿＿＿の　ことばは　どう　かきますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。"):
            stem = re.sub(r'^２　＿＿＿の　ことばは　どう　かきますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。\n', '', question[2])
            question_vocabulary_type2.append(stem)
        if question[2].startswith("3　(   　  ) に　　なにを　いれますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。"):
            stem = question[2].replace("3　(   　  ) に　　なにを　いれますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。\n", "")
            question_vocabulary_type3.append(stem)

        if question[2].startswith("４　＿＿＿の　ぶんと　だいたい　おなじ　いみの　ぶんが　あります。　1・2・3・4から　ひとつ　えらんで　ください。"):
            stem = re.sub(r'^４　＿＿＿の　ぶんと　だいたい　おなじ　いみの　ぶんが　あります。　1・2・3・4から　ひとつ　えらんで　ください。\n', '', question[2])
            question_vocabulary_type4.append(stem)

        if question[2].startswith("5　つぎの　ことばの　つかいかたで　いちばん　いい　ものを　1・2・3・4から　ひとつ　えらんで　ください。"):
            stem = re.sub(r'^5　つぎの　ことばの　つかいかたで　いちばん　いい　ものを　1・2・3・4から　ひとつ　えらんで　ください。\n', '', question[2])
            question_vocabulary_type5.append(stem)
        
    Word_Document(test_paper, question_vocabulary_type1, question_vocabulary_type2, question_vocabulary_type3, question_vocabulary_type4, question_vocabulary_type5, question_grammar_type1, save_path)
    

def analyze(id, test_paper, result, save_path):
    knowledge_points = {}
    mistakes_sum = 0
    print(end='\n\n\n\n\n')
    print("student info: ", "student_no:",result[0][0]," ", "name:",result[0][1]," ", "email:",result[0][2],end = '\n')
    for row in result:
        if row[10] == 1:
            mistakes_sum += 1
            types = re.split(r'[:,]', row[4])

            types = [type.strip() for type in types[1:]]

            for type in types:
                knowledge_points[(row[5], type)] = knowledge_points.get((row[5],type), 0) + 1

            #knowledge_points[(row[5], row[4])] = knowledge_points.get((row[5],row[4]), 0) + 1
            print("question info: ", "question_index:",row[3]," ", "type:",row[4]," ", "level:",row[5]," ", "is_gpt:",row[6]," ",end = '\n')

            if re.search(r'\s*\n+End of Part', row[7]):
                row[7] = re.split(r'\s*End of Part\s*', row[7])[0]

            print("question content: ",row[7],end = '\n')
            print("correct answer: ",row[8], ";" ,"student answer: ",row[9], end = '\n')
            print("------------------------------------------------------")
    print("mistakes_sum: ", mistakes_sum, end = '\n')
    knowledge_points = sorted(knowledge_points.items(), key=sort_key, reverse=False)
    for key in knowledge_points:
        if re.match(r'\d+\.', key[0][1]):
            type = 'Grammar: '
        else:
            type = 'Vocabulary: '
        print("level:", key[0][0])
        print("knowledge_point:", type+key[0][1])
        print(end='\n')
            

def generate(id, test_paper, result, save_path):
    knowledge_points = {}
    mistakes_sum = 0
    for row in result:
        if row[10] == 1:
            mistakes_sum += 1
            types = re.split(r'[:,]', row[4])

            types = [type.strip() for type in types[1:]]

            for type in types:
                knowledge_points[(row[5], type)] = knowledge_points.get((row[5],type), 0) + 1
            #knowledge_points[(row[5], row[4])] = knowledge_points.get((row[5],row[4]), 0) + 1
    for key, value in knowledge_points.items():
        if round(value * 40 / mistakes_sum) == 0:
            value = 1
        else:
            value = round(value * 40 / mistakes_sum)
        knowledge_points[key] = value
    new_questions = set()
    knowledge_points = sorted(knowledge_points.items(), key=sort_key, reverse=False)
    for key in knowledge_points:
        query = '''
                SELECT questions.* FROM questions 
                WHERE questions.type LIKE %s AND questions.level = %s
                ORDER BY RAND()
                LIMIT %s;
                '''
        cursor.execute(query, (f'%{key[0][1]}%', key[0][0], key[1]))
        new_questions.update(cursor.fetchall())
    
    questions_preprocess(test_paper, new_questions, save_path)


def general_analysis(id , test_paper, operation, save_path):
    #查询学生该试卷的所有题目
    inner_search_query ="""
    SELECT students.student_no, students.name, students.email, 
    questions.question_index, questions.type, questions.level, questions.is_gpt, questions.content, questions.correct_answer, 
    exam_results.student_answer, exam_results.is_correct 
    FROM students, questions, exam_results
    WHERE students.student_id = %s AND questions.question_index like %s AND students.student_id = exam_results.student_id AND questions.question_id = exam_results.question_id;
    """

    cursor.execute(inner_search_query, (id, test_paper + '%'))
    result = cursor.fetchall()
    
    result = [list(row) for row in result]

    if operation == 'ANALYZE':
        analyze(id, test_paper, result, save_path)
    elif operation == 'GENERATE':
        generate(id, test_paper, result, save_path)


def query_papers(id):
    #查询该学生所有做过的试卷
    query = '''
            SELECT questions.question_index 
            FROM questions, exam_results
            WHERE exam_results.student_id = %s AND exam_results.question_id = questions.question_id;
            '''
    cursor.execute(query, (id,))
    paper_list = cursor.fetchall()

    paper_set = set()
    
    #把试卷+问题编号的形式转换为试卷编号的形式
    for paper in paper_list:
        question_index = paper[0]
        paper_id = re.sub(r'\d+$', '', question_index)
        paper_set.add(paper_id)
    for paper in paper_set:
        print(paper)
    test_paper = input("select the specific test paper:")
    if test_paper not in paper_set:
        raise ValueError("The test paper does not exist.")
    else:
        last_file = get_latest_paper_id(test_paper)
        if last_file:
            name = os.path.splitext(os.path.basename(last_file))[0]
            version = name.split("\\")[-1].split(" ")[-1]
            save_path = test_paper + ' ' + str(int(version) + 1)
        else:
            save_path = test_paper+' 1'
    operation = input("select the operation: ANALYZE or GENERATE:").upper()
    general_analysis(id, test_paper, operation, save_path)
    

def main():
    #通过姓名、student_no查询学生主键student_id
    info = input("enter name or student no of student: ")
    try:
        info = int(info)
    except:
        pass
    
    if isinstance(info, int):
        cursor.execute("SELECT students.student_id FROM students WHERE students.student_no = %s", (info,))
        id = cursor.fetchall()
    else:
        cursor.execute("SELECT students.student_id FROM students WHERE students.name = %s", (info,))
        id = cursor.fetchall()
    id = id[0][0]
    query_papers(id)

if __name__ == '__main__':
    db = mysql.connector.connect(
        host="localhost",
        user="root",
        password="123",
        database="JAPGPT"
    )
    cursor = db.cursor()
    
    main()

    cursor.close()
    db.close()