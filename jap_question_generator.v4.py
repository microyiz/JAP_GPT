'''Question_Generator'''
import re
import os
import glob
import time
import string
import warnings
import docx
import mysql.connector
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate  
from langchain.chains import LLMChain   
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from typing import Any
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.worksheet.datavalidation import DataValidation

from jap_paper_revise import produce_new_question_list
from jap_paper_revise import read_docx_to_string_with_format

def split_into_sentences(text):
    sentence_endings = re.compile(r'(?<=[。！？])\s*')
    sentences = sentence_endings.split(text)
    return sentences

def extract_numbered_content(file_path, start_number, end_number):
    """
    Extracts numbered content (e.g., 1-8) from a Word document.

    Args:
        file_path (str): The path to the Word document.
        start_number (int): The starting number of the range to extract.
        end_number (int): The ending number of the range to extract.

    Returns:
        list: A list of tuples where each tuple contains the number and the corresponding text.
    """
    # Load the Word document
    doc = Document(file_path)
    
    # Compile a regex pattern to match the numbered entries
    pattern = re.compile(rf"^\s*(\d+)\.\s*(.*)")
    
    # List to store the extracted content
    content_list = []

    # Iterate through all paragraphs in the document
    for para in doc.paragraphs:
        match = pattern.match(para.text)
        if match:
            number = int(match.group(1))
            text = match.group(2).strip()

            # Check if the number is within the specified range
            if start_number <= number <= end_number:
                content_list.append((number, text))

    return content_list



def extract_grammar_points(docx_file):
    # 读取docx文件
    doc = docx.Document(docx_file)
    
    knowledge_points = []
    
    # 遍历文档的所有段落
    for para in doc.paragraphs:
        print(f"读取的段落内容: {para.text}")
        # 使用正则表达式匹配类似 "1.～あいだ（間）" 的格式
        match = re.match(r'(\d+)\.([^\n]+)', para.text.strip())
        if match:
            # 提取编号和短语部分
            number = match.group(1)
            phrase = match.group(2).strip()
            knowledge_points.append({
                'number': number,
                'phrase': phrase
            })
    
    return knowledge_points


def extract_vocabulary(docx_path):
    doc = Document(docx_path)
    vocabulary_list = []
    number_pattern = r'^\d+\. '  # 匹配以数字和点号开头的行
    
    # 遍历每个段落
    for para in doc.paragraphs:
        # 将每个段落按行分割
        for line in para.text.splitlines():
            # 如果行包含带序号的内容
            if re.match(number_pattern, line.strip()):
                vocabulary_list.append(line.strip())
    
    return vocabulary_list



# 示例：根据生成文本解析出题目、选项和答案
def parse_questions(text):
    """
    解析生成的文本，返回一个包含元组 (question_index, content, options, answer) 的列表。
    
    适用于格式：
    
    Q1 かれが　手伝って　（  　　　　　 ）　宿題 (しゅくだい) が　終わらなっかった。
    1　もらったから
    2　くれなかったから
    3　ほしいから
    4　ほしかったから
    Answer: 2

    Q2 うちの　子どもは　勉強 (べんきょう) しないで　（  　　　　　 ）　ばかりいる。
    1　あそび
    2　あそぶ
    3　あそばない
    4　あそんで
    Answer: 4
    """

    # 用正则提取每个完整的题目（Q+编号开头，直到下一个Q编号或文本结束）
    # question_pattern = re.compile(r'(Q\d+.*?)\nAnswer:\s*(\d)', re.DOTALL)
    question_pattern = re.compile(r'(Q\d+.*?)\n(?:Answer:|$)\s*(\d)?', re.DOTALL)

    question_matches = question_pattern.findall(text)

    qa_list = []
    for q_text, answer in question_matches:
        lines = q_text.strip().splitlines()
        if not lines:
            continue

        # 解析题号和题目内容
        m = re.match(r'(Q\d+)[\.\s]*(.*)', lines[0].strip())
        if m:
            q_index = m.group(1)  # Q1, Q2, ...
            content = m.group(2).strip()  # 题干部分
        else:
            q_index = ""
            content = lines[0].strip()

        # 解析选项
        options = []
        for line in lines[1:]:
            line = line.strip()
            m_opt = re.match(r'^(\d+)[\.\、\s]+(.*)', line)
            if m_opt:
                options.append(f"{m_opt.group(1)}. {m_opt.group(2).strip()}")
            else:
                # 处理题目可能换行的情况
                content += " " + line

        # 存入解析结果
        qa_list.append((q_index, content, "\n".join(options), answer))

    return qa_list



def store_questions_to_excel(qa_list, output, filename):
    """
    将题目信息存储到 Excel 文件中。

    参数:
      qa_list: 包含元组 (question_index, content, options, answer) 的列表
      output: 文件存储的目录
      filename: 文件名（应包含扩展名，例如 "test_1.xlsx"）
    """
    # 创建 Excel 工作簿和工作表
    wb = Workbook()
    ws = wb.active
    ws.title = "Questions"

    # 设置表头，共6列：前4列存题目信息，后2列留空供人工填写审阅意见
    headers = ["Question Index", "Content", "Options", "Answer", "Suggestions", "Modifications (if any)"]
    ws.append(headers)

    # 将题目信息写入表格
    for qa in qa_list:
        q_index, content, options, answer = qa
        ws.append([q_index, content, options, answer, "", ""])

    # 设置各列宽度
    column_widths = {
        "A": 15,
        "B": 50,
        "C": 30,
        "D": 20,
        "E": 20,
        "F": 40
    }
    for col_letter, width in column_widths.items():
        ws.column_dimensions[col_letter].width = width

    # # 设置行高：除了第一行（标题行）以外的所有行高度为 30
    # for row in range(2, ws.max_row + 1):
    #     ws.row_dimensions[row].height = 30

    # 设置所有单元格自动换行，确保内容显示完整
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True)

    # 创建下拉选择的数据验证
    dv = DataValidation(type="list", formula1='"OK,Drop,Minor changes"', allow_blank=True)
    ws.add_data_validation(dv)
    dv_range = "E2:E10000"  # 将数据验证应用到 E 列，从第2行到第10000行
    dv.add(dv_range)

    # 构造完整的文件保存路径，并保存 Excel 文件
    output_path = os.path.join(output, filename)
    wb.save(output_path)
    print(f"Successfully stored in {output_path}")




def grammar_points_revise(vocabulary_list, output, filepath):
    filename = os.path.splitext(os.path.basename(filepath))[0]
    llm = ChatOpenAI(
        temperature=0.6,
        model='gpt-4o'
    )

    i = 1
    # 遍历每个知识点（从vocabulary_list提取的知识点）
    for knowledge_point in vocabulary_list:
        print(f'Processing knowledge point: {knowledge_point}...')

        # 动态生成 prompt，要求 GPT 依次为每个词汇出题
        prompt_grammar = ChatPromptTemplate.from_template(
            f'''
    You are an experienced Japanese examiner, well-versed in the N4 and N5 levels of the Japanese Language Proficiency Test (JLPT). Your task is to create 10 multiple-choice questions based on the following grammar knowledge point: **{knowledge_point}**.

    ### Requirements:
    1. Each question should have **4 options**, with only **one correct answer**.
    2. The correct answer must strictly align with the logic of the question stem.
    3. The incorrect options should be reasonably close to the correct answer but clearly wrong when examined in context.
    4. Ensure that the question is unambiguous by adding necessary contextual constraints (e.g., verb tense, sentence structure, or meaning) to eliminate multiple correct answers.
    5. Ensure that the correct option fits naturally in the sentence context and reflects the meaning of the grammar point.
    6. The incorrect options should follow a clear pattern of errors, such as common learner mistakes in conjugation, word choice, or syntax. Avoid random incorrect options that do not relate to actual language learning challenges.
    7. After generating the question, apply a check to ensure that the correct answer is unique, and the other options are incorrect and clearly unsuitable.
    8. Avoid misleading phrasing or unnatural sentence constructions that may confuse test-takers. The question should resemble real exam-level language usage.
    9. Prepend the knowledge point you used to generate the question to the front of all questions.
    10. Attach all your answers at the end of the 10 questions.

            ### Question Formats:

            （  　　　　　 ）に　何を　入れますか。　1・2・3・4から　いちばん　いい　ものを　一つ　えらんで　ください。  
                Q1 かれが　手伝って　（  　　　　　 ）　宿題 (しゅくだい) が　終わらなっかった。  
                1　もらったから			2　くれなかったから		
                3　ほしいから				4　ほしかったから
                Answer: 2

                Q2 うちの　子どもは　勉強 (べんきょう) しないで　（  　　　　　 ）　ばかりいる。
                1　あそび		2　あそぶ		3　あそばない		4　あそんで
                Answer: 4

                Q3  A　「田中さんは　かのじょが　いますか。」
	                B　「いいえ、田中さんは　前の　かのじょと　別れてから、人を好き　（  　　　　　 ）。」
                1　ではありませんでした		    2　にならなくなりました		
                3　でもよくなりました			4　にしなくなりました
                Answer: 2

    ### Additional Notes:
    - The generated questions must maintain high linguistic and contextual accuracy.
    - Avoid using cultural or subjective biases that could confuse learners.
    - Ensure that the sentence structure follows standard Japanese grammar rules. Avoid artificial sentence constructions that do not resemble natural spoken or written Japanese.
    - Limit the use of uncommon words or phrases that may be beyond N4/N5 level unless necessary for testing a specific grammar structure.

    ### Before finalizing, check your output against these rules:
    1. Each question must start with `Qx` (e.g., `Q1`, `Q2`...).
    2. Each question must have exactly 4 options (`1` to `4`).
    3. Each question must have an `Answer: x` at the end.
    4. If any question is missing `Answer: x`, fix it before outputting.
    5. Do not include other comments such as "**"
    '''
        )

        
        # 创建链条来运行LLM
        chain_one = LLMChain(llm=llm, prompt=prompt_grammar)

        # 输入数据
        input_data = {
            'knowledge_point': knowledge_point,  # 使用当前知识点
            'num_questions': 10
        }

        # 获取生成的题目
        revise_result = chain_one.run(input_data)

        # 处理生成的题目并保存
        output_doc = Document()
        sentences = split_into_sentences(revise_result)
        
        # 将生成的题目逐句添加到文档中
        for sentence in sentences:
            sentence.replace("**Answers:**", "**Answers**")
            sentence = sentence.replace("＿＿＿", "[ ]")  # 替换空格部分
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(output, f"{filename}_new{i}.docx")
        output_doc.save(output_path)

        # # 将生成文本解析成题目列表
        # qa_list = parse_questions(revise_result)
        # excel_filename = f"{filename}_new{i}.xlsx"
        # store_questions_to_excel(qa_list, output, excel_filename)
        # # print(f'Generated questions for {knowledge_point} saved to {output_path}')
        i += 1



def vocabulary_points_revise(vocabulary_list, output, filepath):
    filename = os.path.splitext(os.path.basename(filepath))[0]
    llm = ChatOpenAI(
        temperature=0.6,
        model='gpt-4o'
    )

    i = 1
    # 遍历每个知识点（从vocabulary_list提取的知识点）
    for knowledge_point in vocabulary_list:
        print(f'Processing knowledge point: {knowledge_point}...')

        # 动态生成 prompt，要求 GPT 依次为每个词汇出题
        prompt_vocabulary = ChatPromptTemplate.from_template(
    f'''
    You are an experienced Japanese examiner, well-versed in the N4 and N5 levels of the Japanese Language Proficiency Test (JLPT). Your task is to create 10 multiple-choice questions based on the following vocabulary knowledge point: **{knowledge_point}**.

    ### Requirements:
    1. Each question should have **4 options**, with only **one correct answer**.
    2. The correct answer must strictly match the meaning or usage of the vocabulary word, ensuring no ambiguity.
    3. The incorrect options should deviate from the correct one but be contextually plausible.
    4. Introduce logical conditions or context into the stem that restricts the possible answers, ensuring only one option fits.
    5. The options should be meaningful, with the incorrect ones being close to the correct answer, yet still clearly wrong in context.
    6. After generating the question, apply a check to ensure that only one answer is clearly correct, and the other options are incorrect.
    7. Prepend the vocabulary knowledge point to the front of all questions.
    8. Attach all your answers at the end of the 10 questions.
    9. If the question is asking how to write a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\
        If the question is asking how to write a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question.\

    ### Question Formats:
            1. **How to write in hiragana:**
                Q1: ＿＿＿の　ことばは　ひらがなで　どう　かきますか。   
                あそこに　かわいい　[鳥]が　います。  
                1. いぬ  2. とり  3. ねこ  4. むし 
                Answer: 2 

            2. **Kanji recognition:**  
                Q2: ＿＿＿の　ことばは　どう　かきますか。  
                [おっと]は　今、出かけています。  
                1. 大  2. 犬  3. 太  4. 夫  
                Answer: 4

            3. **Filling in the blanks:**  
                Q3: (   　  ) に　なにを　いれますか。  
                これから　ひこうきに　（  　　　　　 ）。  
                1. おります  2. のります  3. あがります  4. のぼります  
                Answer: 2

            4. **Sentence meaning comparison:**  
                Q4: ＿＿＿の　ぶんと　だいたい　おなじ　いみの　ぶんが　あります。  
                ギターは　ちちに　ならいました。  
                1. ギターは　ちちに　もらいました。  
                2. ギターは　ちちに　えらんでもらいました。  
                3. ギターは　ちちに　おしえてもらいました。  
                4. ギターは　ちちに　かってもらいました。  
                Answer: 3

            5. **Usage of vocabulary:**  
                Q5: つぎの　ことばの　つかいかたで　いちばん　いい　ものを　1・2・3・4から　ひとつ　えらんで　ください。  
                わる  
                1. おさらを　[わって]　母に　おこられました。  
                2. おさらを　[わって]　へやに　かざりました。  
                3. おさらを　[わって]　りょうりを　つくりました。  
                4. おさらを　[わって]　コーヒーを　のみました。  
                Answer: 1


    ### Additional Notes:
    - The generated questions must maintain high linguistic and contextual accuracy.
    - Avoid using cultural or subjective biases that could confuse learners.

    ### Before finalizing, check your output against these rules:
    1. Each question must start with `Qx` (e.g., `Q1`, `Q2`...).
    2. Each question must have exactly 4 options (`1` to `4`).
    3. Each question must have an `Answer: x` at the end.
    4. If any question is missing `Answer: x`, fix it before outputting.
    5. Do not include other comments such as "**"
    '''
)

        
        # 创建链条来运行LLM
        chain_one = LLMChain(llm=llm, prompt=prompt_vocabulary)

        # 输入数据
        input_data = {
            'knowledge_point': knowledge_point,  # 使用当前知识点
            'num_questions': 10
        }

        # 获取生成的题目
        revise_result = chain_one.run(input_data)

        # 处理生成的题目并保存
        output_doc = Document()
        sentences = split_into_sentences(revise_result)
        
        # 将生成的题目逐句添加到文档中
        for sentence in sentences:
            sentence.replace("**Answers:**", "**Answers**")
            sentence = sentence.replace("＿＿＿", "[ ]")  # 替换空格部分
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(output, f"{filename}_new{i}.docx")
        output_doc.save(output_path)

        # # 将生成文本解析成题目列表
        # qa_list = parse_questions(revise_result)
        # excel_filename = f"{filename}_new{i}.xlsx"
        # store_questions_to_excel(qa_list, output, excel_filename)
        # # print(f'Generated questions for {knowledge_point} saved to {output_path}')
        i += 1



def question_revise_simple(rows, filename, revised_newpaper_folder, max_iterations=5):
    llm = ChatOpenAI(
        temperature=0.6,
        model='gpt-4o'
    )
    
    prompt_revise = ChatPromptTemplate.from_template(
    "Here are the new generated Japanese practice questions: {new_paper} \
    You are an experienced Japanese N4/N5 examiner tasked with reviewing and ensuring that all multiple-choice test questions meet the following criteria:\n\n\
    1. **No duplicate questions**: Ensure that all questions are unique. If a question is repeated or too similar to another, please revise it to create a new question with a distinct structure or context.\n\n\
    2. **No duplicate options**: All options within a question should be unique, contextually meaningful, and grammatically correct. Avoid options that are too similar to each other, and ensure they are relevant to the question.\n\n\
    3. **No duplicate correct answers**: The question should have only one correct answer. Do not allow multiple options to seem correct; ensure only one answer is suitable and unambiguous. **If necessary, modify the question stem or incorrect options to ensure clear distinctions.**\n\n\
    4. **Grammatical correctness**: The title and stem of each question must be grammatically correct. **Review for unnatural sentence structures and revise them to ensure fluency and correctness.**\n\n\
    5. **Relevance of options**: The stem should clearly indicate what cannot be chosen. Ensure that one option is inappropriate or clearly wrong in context, while all other options are suitable. Avoid subjective or culturally biased content. **Check that incorrect options reflect common mistakes learners make rather than being randomly incorrect.**\n\n\
    6. **Pronunciation and Word Usage**: If the question involves pronunciation, katakana, or hiragana forms, the Japanese word should be enclosed in brackets for clarity. If asking about hiragana or katakana conversion, ensure that the word is written in the appropriate form (katakana for hiragana conversion, and vice versa) and that the correct answer is not shown in the question stem.\n\n\
    - If the question asks for a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show the right answer in the question.\n\
    - If the question asks for a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show the right answer in the question.\n\
    - Check for spelling inconsistencies or errors in transcriptions.\n\n\
    7. **General guidance**: Eliminate any ambiguity, revise unclear options, and avoid subjective or culturally biased phrasing. **Ensure all questions are at an appropriate difficulty level for the target JLPT level (N4/N5).**\n\n\
    8. **Output Format**: Each question must keep the original format, meaning that each question should have its answer immediately following it. Do NOT move all answers to the end.\n\n\
    ### Before finalizing, check your output against these rules:\n\
    8.1. Each question must start with `Qx` (e.g., `Q1`, `Q2`...).\n\
    8.2. Each question must have exactly 4 options (`1` to `4`).\n\
    8.3. Each question must have an `Answer: x` at the end of it.\n\
    8.4. If any question is missing `Answer: x`, fix it before outputting.\n\
    8.5. Do not include other comments such as '**'. "
    )

    
    chain = LLMChain(llm=llm, prompt=prompt_revise)
    input_data = {'new_paper': rows}
    
    for iteration in range(max_iterations):
        revised_result = chain.run(input_data)
        errors = check_for_error(revised_result)

        if not errors:
            print(f"No issues found after {iteration + 1} iterations.")
            break

        # Ensure errors is iterable
        print(f"Iteration {iteration + 1}: Detected errors - {', '.join(errors)}")
        input_data['new_paper'] = revised_result

        
        # 保存中间修订结果
        intermediate_path = os.path.join(revised_newpaper_folder, f"{filename}_iteration_{iteration + 1}.docx")
        output_doc = Document()
        sentences = split_into_sentences(revised_result)
        for sentence in sentences:
            output_doc.add_paragraph(sentence)
        output_doc.save(intermediate_path)
        
        # 保存错误日志
        log_path = os.path.join(revised_newpaper_folder, f"{filename}_error_log.txt")
        with open(log_path, 'a', encoding='utf-8') as log_file:
            log_file.write(f"Iteration {iteration + 1} Errors: {', '.join(errors)}\n")
    else:
        print(f"Maximum iterations ({max_iterations}) reached. Errors may still exist.")
    
    # 保存最终修订结果
    output_path = os.path.join(revised_newpaper_folder, f"{filename}_revised.docx")
    output_doc = Document()
    sentences = split_into_sentences(revised_result)
    for sentence in sentences:
        output_doc.add_paragraph(sentence)
    output_doc.save(output_path)

    # 存储到excel
    qa_list = parse_questions(revised_result)
    excel_filename = f"{filename}_revised.xlsx"
    store_questions_to_excel(qa_list, revised_newpaper_folder, excel_filename)



# def check_for_error(revised_text):
#     """
#     Check for errors in the revised question set, including:
#     - Multiple correct answers
#     - Duplicate questions
#     - Errors in the question stem
#     - Duplicate options
    
#     :param revised_text: The revised text output from GPT.
#     :return: True if errors are found, False otherwise.
#     """
#     try:
#         if has_multiple_correct_answers(revised_text):
#             print("Error detected: Multiple correct answers")
#             return True

#         if has_duplicate_questions(revised_text):
#             print("Error detected: Duplicate questions")
#             return True

#         if has_stem_errors(revised_text):
#             print("Error detected: Stem errors")
#             return True

#         if has_duplicate_options(revised_text):
#             print("Error detected: Duplicate options")
#             return True

#         return False  # No errors detected
#     except Exception as e:
#         print(f"Error in check_for_error: {e}")
#         return True  # Treat as error if an exception occurs

def check_for_error(revised_text):
    """
    Check for errors in the revised question set, including:
    - Multiple correct answers
    - Duplicate questions
    - Errors in the question stem
    - Duplicate options
    
    :param revised_text: The revised text output from GPT.
    :return: List of errors if any are found, empty list otherwise.
    """
    errors = []  # Initialize an empty list to store error messages
    
    try:
        if has_multiple_correct_answers(revised_text):
            errors.append("Multiple correct answers")
        
        if has_duplicate_questions(revised_text):
            errors.append("Duplicate questions")
        
        if has_stem_errors(revised_text):
            errors.append("Stem errors")
        
        if has_duplicate_options(revised_text):
            errors.append("Duplicate options")
        
        return errors  # Return the list of errors (can be empty if no errors)
    
    except Exception as e:
        print(f"Error in check_for_error: {e}")
        return ["Unexpected error in check_for_error"]  # Return a list with an error message if an exception occurs


def has_multiple_correct_answers(text):
    """
    Checks if a Japanese multiple-choice question has more than one possible correct answer.
    
    :param text: The text containing the multiple-choice questions.
    :return: True if multiple correct answers exist, False otherwise.
    """
    llm = ChatOpenAI(
        temperature=0.3,  # Lower temperature for more deterministic output
        model="gpt-4o"
    )
    
    prompt = ChatPromptTemplate.from_template(
        "You are an experienced Japanese N4/N5 examiner reviewing the following multiple-choice questions:\n\n"
        "{new_paper}\n\n"
        "Check if any question has **more than one correct answer**. This means that multiple options are valid for the question given its context.\n"
        "If at least one question has multiple valid correct answers, respond with 'True'. Otherwise, respond with 'False'.\n"
        "Your output must be exactly 'True' or 'False', nothing else."
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    input_data = {'new_paper': text}

    try:
        result = chain.run(input_data).strip().lower()
        return result == "true"
    except Exception as e:
        print(f"Error processing has_multiple_correct_answers: {e}")
        return False  # Default to False if an error occurs


def has_stem_errors(text):
    """
    Checks if there are grammatical errors or ambiguities in the question stems.
    
    :param text: The text containing the multiple-choice questions.
    :return: True if errors exist, False otherwise.
    """
    llm = ChatOpenAI(
        temperature=0.3,  # Lower temperature to improve reliability
        model="gpt-4o"
    )
    
    prompt = ChatPromptTemplate.from_template(
        "You are an experienced Japanese N4/N5 examiner reviewing the following multiple-choice questions:\n\n"
        "{new_paper}\n\n"
        "Check if any **question stem** (the main question part before the options) has errors, such as:\n"
        "- Grammatical mistakes\n"
        "- Unnatural sentence structures\n"
        "- Ambiguous wording\n"
        "If there is at least one issue in the stems, respond with 'True'. Otherwise, respond with 'False'.\n"
        "Your output must be exactly 'True' or 'False', nothing else."
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    input_data = {'new_paper': text}

    try:
        result = chain.run(input_data).strip().lower()
        return result == "true"
    except Exception as e:
        print(f"Error processing has_stem_errors: {e}")
        return False  # Default to False if an error occurs

def has_duplicate_options(text):
        """
        Check for duplicate options in the questions.
        
        :param text: The text to check.
        :return: True if - options are found within a question.
        """
        # Example: Detect duplicate options for a given question.
        questions_with_options = re.findall(
            r'(\d+)\.\s*(.*?)\n(1\.\s*(.*?)\n)(2\.\s*(.*?)\n)(3\.\s*(.*?)\n)(4\.\s*(.*?)\n)',
            text, re.DOTALL
        )

        for question, _, opt1, _, opt2, _, opt3, _, opt4, _ in questions_with_options:
            options = {opt1.strip(), opt2.strip(), opt3.strip(), opt4.strip()}
            if len(options) < 4:  # If any options are duplicates
                print(f"Duplicate options detected in question {question}: {opt1.strip()}, {opt2.strip()}, {opt3.strip()}, {opt4.strip()}")
                return True
        return False




def normalize_text(text):
    """
    Normalize text by:
    1. Converting to lowercase.
    2. Removing non-essential characters such as punctuation and extra spaces.
    
    :param text: The text to normalize.
    :return: Normalized text.
    """
    # Convert to lowercase
    text = text.lower()
    
    # Remove punctuation and extra spaces
    text = text.translate(str.maketrans('', '', string.punctuation))
    text = re.sub(r'\s+', ' ', text).strip()  # Remove extra spaces
    
    return text

def has_duplicate_questions(text):
    """
    Check if any questions are duplicated, considering both the question text and options,
    while ignoring case and non-key characters like spaces and punctuation.
    
    :param text: The text to check.
    :return: True if duplicate questions are detected.
    """
    questions = re.findall(
        r'(\d+)\.\s*(.*?)\n(1\.\s*(.*?)\n)(2\.\s*(.*?)\n)(3\.\s*(.*?)\n)(4\.\s*(.*?)\n)',  # Capture question text and options
        text, re.DOTALL
    )
    
    seen_questions = set()
    
    for question, _, opt1, _, opt2, _, opt3, _, opt4, _ in questions:
        # Normalize question text and options
        question_text = normalize_text(question.strip())
        options = {normalize_text(opt1.strip()), normalize_text(opt2.strip()), 
                   normalize_text(opt3.strip()), normalize_text(opt4.strip())}
        
        # Create a normalized string for comparison: question + sorted options
        normalized_question = f"{question_text} - {', '.join(sorted(options))}"
        
        if normalized_question in seen_questions:
            print(f"Duplicate question detected: {question_text} with options {options}")
            return True  # Duplicate found
        seen_questions.add(normalized_question)
    
    return False




# 测试
docx_file_grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\N4N5 material\\N4 Notes 文法_numbered.docx"
docx_file_vocabulary = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\N4N5 material\\N4 Notes 語彙_numbered.docx"
test_knowledge_points = "C:\\Users\\刘宇\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\test_knowledge_points.docx"
test_grammar = "C:\\Users\\刘宇\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\test_grammar.docx"
test_vocabulary = "C:\\Users\\刘宇\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\test_vocabulary.docx"

output_grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\N4 grammar"
output_vocabulary = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\N4 vocabulary"
revised_output_vocabualry = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\revised_vocabulary"
revised_output_grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\revised_grammar"
# #grammar_points = extract_grammar_points(docx_file_grammar)
# #grammar_points_revise(grammar_points, output_grammar, docx_file_grammar)

# #rows = extract_numbered_content(test_knowledge_points, 1, 8)
vocabulary = extract_numbered_content(test_vocabulary, 1, 6)
grammar = extract_numbered_content(test_grammar, 1, 4)

# print(vocabulary)
# print(grammar)

vocabulary_points_revise(vocabulary, revised_output_vocabualry, test_vocabulary)
grammar_points_revise(grammar, revised_output_grammar, test_grammar)

# Iterate over all the new question files and fix them
for filepath in glob.glob(os.path.join(revised_output_vocabualry, "*.docx")):

    filename = os.path.splitext(os.path.basename(filepath))[0]
            
    start_time = time.time()

    new_que = read_docx_to_string_with_format(filepath)
    #question_revise(new_que, filename, revised_output)
    question_revise_simple(new_que, filename, revised_output_vocabualry)

    end_time = time.time()
    print(f"Completed revising new questions {filename} in: {end_time - start_time:.2f} seconds")


# Iterate over all the new question files and fix them
for filepath in glob.glob(os.path.join(revised_output_grammar, "*.docx")):

    filename = os.path.splitext(os.path.basename(filepath))[0]
            
    start_time = time.time()

    new_que = read_docx_to_string_with_format(filepath)
    #question_revise(new_que, filename, revised_output)
    question_revise_simple(new_que, filename, revised_output_grammar)

    end_time = time.time()
    print(f"Completed revising new questions {filename} in: {end_time - start_time:.2f} seconds")




# test_duplicate_stem = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\AI Questions\\test_vocabulary_new1_revised.docx"
# filename = os.path.splitext(os.path.basename(test_duplicate_stem))[0]
# new_que = read_docx_to_string_with_format(test_duplicate_stem)
# question_revise_simple(new_que, filename, revised_output_vocabualry)