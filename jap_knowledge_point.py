import re
import os
import glob
import win32com.client
from typing import Any
from docx import Document
from lxml import etree
from zipfile import ZipFile
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate  
from langchain.chains import LLMChain 
from langchain.chains import ConversationalRetrievalChain
from langchain.text_splitter import RecursiveCharacterTextSplitter  # 文本分割器
from langchain_community.embeddings import OpenAIEmbeddings  # OpenAI嵌入
from langchain_community.vectorstores import FAISS  # 向量数据库
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain.agents import initialize_agent, Tool, AgentType  # 智能代理

from jap_paper_revise import clean_document, read_docx_to_string, generate_question_separators, split_text_with_separators


class JapaneseCharacterTextSplitter(RecursiveCharacterTextSplitter):
    def __init__(self, **kwargs: Any):
        #separators = ["\n\n", "\n", "。", "、", " ", ""]
        separators = [
        "\n\n",
        "\n",
        "。",
        "、",
        " ",
        ".",
        ",",
        "?",
        "\u200B",  # Zero-width space
        "\uff0c",  # Fullwidth comma
        "\u3001",  # Ideographic comma
        "\uff0e",  # Fullwidth full stop
        "\u3002",  # Ideographic full stop
        "",
        ]

        question_separators = [f"Q.{i}" for i in range(1, 500)]
        separators.extend(question_separators)
        number_space_separators = [f"{str(i).zfill(2)}"for i in range(1, 100)]
        separators.extend(number_space_separators)
        super().__init__(separators=separators, **kwargs)




# 1. 读取普通文本并标记 ruby 的函数
def extract_ruby_and_modify_text(docx_path):
    with ZipFile(docx_path, 'r') as zip_ref:
        xml_content = zip_ref.read('word/document.xml')
    
    # 使用 lxml 解析 XML
    root = etree.XML(xml_content)
    namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    ruby_texts = []
    for ruby in root.xpath('//w:ruby', namespaces=namespace):
        # 提取 rubyBase 和 rt 信息
        base_text = ''.join([t.text for t in ruby.xpath('.//w:rubyBase//w:t', namespaces=namespace) if t.text])
        ruby_text = ''.join([t.text for t in ruby.xpath('.//w:rt//w:t', namespaces=namespace) if t.text])
        
        # 构造带标签的 ruby 标记
        ruby_tag = f'<ruby base="{base_text}" rt="{ruby_text}" />'
        ruby_texts.append({
            'main_text': base_text,
            'ruby_text': ruby_text,
            'ruby_tag': ruby_tag,
        })

    # 替换 XML 中的 ruby 标签为自定义标记
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for ruby in root.xpath('//w:ruby', namespaces=namespace):
        # 找到 rubyBase 的文字
        base_text = ''.join([t.text for t in ruby.xpath('.//w:rubyBase//w:t', namespaces=namespace) if t.text])
        ruby_tag = next((item['ruby_tag'] for item in ruby_texts if item['main_text'] == base_text), None)
        if ruby_tag:
            # 替换 ruby 的 XML 结构为带标签的内容
            ruby_element = etree.Element(f"{{{ns}}}t")  # 使用完整命名空间定义标签
            ruby_element.text = ruby_tag  # 设置文本为自定义 ruby 标记
            ruby.getparent().replace(ruby, ruby_element)

    # 提取标记替换后的完整文本
    modified_text = etree.tostring(root, encoding="unicode", method="text")
    return modified_text, ruby_texts


# 2. 分块处理文本
def split_document(text):
    splitter = JapaneseCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
    )
    return splitter.split_text(text)


# 3. 保存处理后的文档
def save_to_docx(splits, output_folder, filename):
    os.makedirs(output_folder, exist_ok=True)  # 确保文件夹存在
    output_path = os.path.join(output_folder, f"{filename}_processed.docx")
    doc = Document()
    for split in splits:
        doc.add_paragraph(split)
    doc.save(output_path)
    print(f"Saved processed document: {output_path}")





def split_into_sentences(text):
    sentence_endings = re.compile(r'(?<=[。！？])\s*')
    sentences = sentence_endings.split(text)
    return sentences


def split_test_paper(input_file, filename):
    file_name = f"{filename}.docx"
    output_file_path = os.path.join("C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed test paper with knowledge points", 
                                    file_name)
    clean_document(input_file,output_file_path)
    d = read_docx_to_string(output_file_path)
    question_sep = generate_question_separators(100)
    split_result = split_text_with_separators(d, question_sep)

    pattern1 = r'もんだい\d+'
    pattern2 = r'問題\d+'

    split_result_copy = []
    ques_type =" "
    for element in split_result:
        if re.search(pattern1, element):
            test = element.split("もんだい")
            ques_type = test[1]
            split_result_copy.append(test[1]+"\n"+test[0])
          
        elif re.search(pattern2,element):
            test = element.split("問題")
            ques_type = test[1]
            split_result_copy.append(test[1]+"\n"+test[0])

        else:
            split_result_copy.append(ques_type+"\n"+element)
    
    return split_result_copy

# 这里GPT无法直接读取到grammar和vocabulary的内容，需要切分一下再输入
def add_knowledge_points(rows, grammar_vocabulary, output_folder, filename):
    llm = ChatOpenAI(
        temperature=0.6,
        model="gpt-4o"
    )

    # with open("prompt_combine_knowledgepoint.txt", "r", encoding="utf-8") as file:
    #     prompt_content = file.read()
    # prompt_one = ChatPromptTemplate.from_template(prompt_content)
    prompt_one = ChatPromptTemplate.from_template(
        "I have provided you with two documents, among which {paper} is a Japanese N4 test paper, {knowlege_points} are Japanese N4 knowledge points.\n"
        "Now please read the above two documents carefully and complete the following three tasks in order:\n"
        "1. Carefully analyze the knowledge points in {knowlege_points};\n"
        "2. Analyze the test questions in the {paper}, focusing on the knowledge points examined by each question;\n"
        "3. Correspond to the knowledge points analyzed in step 1 and the knowledge points in step 2;\n"
        "4. Generate a document, including the test questions in the test1 question paper, and attach the knowledge points examined after each question. \n"
        "Please directly quote the Knowledge points and keep the document format the same as {paper}."
        "Please only keep the content of step 4 and provide the whole paper."
    )
    
    chain_one = LLMChain(llm=llm, prompt=prompt_one)

    inputs_one = {
        'paper': rows,
        'knowlege_points': grammar_vocabulary
    }
    revise_result = chain_one.run(inputs_one)
    output_doc = Document()
    sentences = split_into_sentences(revise_result)
    for sentence in sentences:
        sentence.replace("＿＿＿", "[ ]")
        output_doc.add_paragraph(sentence)

    output_path = os.path.join(output_folder, f"{filename}_modified.docx")
    output_doc.save(output_path)


# def main():
#     test_paper = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\Test1_new\\test 1 paper\\Test 1 Question Paper.docx"
#     output_folder = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed test paper with knowledge points"
#     processed_material = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed material"
#     grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed material\\N4 Notes 文法.docx"
#     vocabulary = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\N4N5 material\\N4 Notes 語彙.docx"
#     # filename = "Test 1 Question Paper"
#     for filepath in glob.glob(os.path.join(processed_material, "*.docx")):
#         filename = os.path.splitext(os.path.basename(filepath))[0]
#         docs = load_docx(filepath)
#         splits = split_document(docs)
#         append_splits_to_docx(splits, output_folder)
#     # d = split_test_paper(test_paper, filename)
#     # add_knowledge_points(d, grammar, vocabulary, output_folder, filename)

def main():
    test_paper = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\Test1_new\\test 1 paper\\Test 1 Question Paper.docx"
    output_folder = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed test paper with knowledge points"
    processed_material = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed material"
    grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed material\\N4 Notes 文法.docx"
    vocabulary = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed material\\N4 Notes 語彙.docx"

    for filepath in [grammar, vocabulary]:
        filename = os.path.splitext(os.path.basename(filepath))[0]
        print(f"Processing file: {filepath}")
        
        # 提取文本并标记 ruby 信息
        modified_text, ruby_texts = extract_ruby_and_modify_text(filepath)
        
        # 分块处理带标记的文本
        splits = split_document(modified_text)
        
        # # 保存分块后的文档
        # save_to_docx(splits, processed_material, filename)

        knowledge_points_list = []
        for split in splits:
            knowledge_points_list.append(split)

    
    file = os.path.splitext(os.path.basename(test_paper))[0]
    d = split_test_paper(test_paper, file)
    add_knowledge_points(d, ' '.join(knowledge_points_list), output_folder, file)
    print("Add knowledge points successfully.")


if __name__ == "__main__":
    main()

