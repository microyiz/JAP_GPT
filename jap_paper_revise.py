import re
from docx import Document
import datetime
import os

"""
 readin : the path to answer.docx
 output: a list containing all the roptions（按照题目的顺序）
"""
def read_answers_from_docx(file_path):
    doc = Document(file_path)
    answers = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("問題　"):
            temp = text.split('：')[-1].strip()
            answers.append(temp)
    return answers

"""
 readin : list1(right answer) list2(students' answer)
 output: a list containing tuple
        for example (0,4,3) means 第一题是错的，正确答案是4，学生选了3,
"""
def read_list_difference(list1,list2):
    if len(list1)!=len(list2):
        raise ValueError("The two lists must have the same length.")
    differences = []
    for index, (item1, item2) in enumerate(zip(list1, list2)):
        differences.append((index,item1,item2))
    return differences

# a = read_answers_from_docx("C:\\Users\\chen\\Desktop\\paper\\Test 1 Model Answer.docx")
# b = read_answers_from_docx("C:\\Users\\chen\\Desktop\\paper\\1155142665 Test 1.docx")
# c = read_list_difference(a,b)
# print(a)
# print(b)
# print(c)
"""
 readin : answer sheet
 output: student id who submitted the answer sheet
"""
def extract_student_id(file_path):
    # 定义正则表达式模式来匹配学号
    pattern = r'\\(\d{10})\s'
    match = re.search(pattern, file_path)
    
    if match:
        # 提取匹配的学号
        student_id = match.group(1)
        return student_id
    else:
        return None


#remove all the contents between 两个长分隔符之间的内容，也就是 删除所有的example问题
def remove_delimiters(text):
    delimiter_pattern = r'＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿.*?＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿'
    text = re.sub(delimiter_pattern, '', text, flags=re.DOTALL)
    
    return text

def remove_specific_sentence(text, target_sentence):
    # Find the target sentence and its preceding sentence
    lines = text.split('\n')
    new_lines = [line for line in lines if target_sentence not in line]
    return '\n'.join(new_lines)

#清理试卷
def clean_document(filepath, outputfilepath):
    doc = Document(filepath)
    new_doc = Document()
    full_text =[]
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    text = '\n'.join(full_text)
    cleaned_text = remove_delimiters(text)
    # target_sentence = "End of Part 1.  Thank you!  Please continue to complete Part 2."
    # cleaned_text = remove_specific_sentence(cleaned_text, target_sentence)
    new_doc = Document()
    new_doc.add_paragraph(cleaned_text)
    # for line in cleaned_text.split('\n'):
    #     new_doc.add_paragraph(line)
            
    new_doc.save(outputfilepath)

# 生成题号
def generate_question_separators(max_questions: int):
        """
        Generate a list of question separators up to a specified number of questions.
        Combines circled numbers with other numeric patterns to cover up to max_questions.
        """
        circled_numbers = []
        for i in range(1, max_questions + 1):
            if 1 <= i <= 20:
                circled_numbers.append(chr(0x2460 + i - 1))  # ① to ⑳
            elif 21 <= i <= 35:
                circled_numbers.append(chr(0x3251 + i - 21))  # ㉑ to ㉟
            elif 36 <= i <= 50:
                circled_numbers.append(chr(0x32B1 + i - 36))  # ㊱ to ㊿
        return circled_numbers

def read_docx_to_string(file_path):
    """
    读取 docx 文件中的所有文字并拼接成一个字符串

    :param file_path: docx 文件路径
    :return: 包含所有文字的字符串
    """
    doc = Document(file_path)
    full_text = []

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    return '\n'.join(full_text)


def split_text_with_separators(text, separators):
    """
    将字符串根据给定的间隔符列表分割成多个部分。

    :param text: 需要分割的字符串
    :param separators: 间隔符列表
    :return: 分割后的字符串列表
    """
    # 创建正则表达式模式来匹配任意间隔符
    pattern = '|'.join(re.escape(separator) for separator in separators)
    
    # 使用正则表达式分割字符串
    split_text = re.split(pattern, text)
    
    # 去除空字符串并返回结果
    return [part.strip() for part in split_text if part.strip()]

#将整张试卷按照题号进行切割
def produce_split_question_list(input_file, filename):
    timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
    file_name = f"{filename} {timestamp}.docx"
    output_file_path = os.path.join("C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\processed test paper", file_name)
    clean_document(input_file,output_file_path)
    d = read_docx_to_string(output_file_path)
    question_sep = generate_question_separators(100)
    split_result = split_text_with_separators(d, question_sep)

    pattern1 = r'もんだい\d+'
    pattern2 = r'問題\d+'

    split_result_copy = []
    ques_type =" "
    for element in split_result:
        if re.search(pattern1, element):
            test = element.split("もんだい")
            ques_type = test[1]
            split_result_copy.append(test[1]+"\n"+test[0])
          
        elif re.search(pattern2,element):
            test = element.split("問題")
            ques_type = test[1]
            split_result_copy.append(test[1]+"\n"+test[0])

        else:
            split_result_copy.append(ques_type+"\n"+element)
    
    return split_result_copy

"""
    生成两个列表，一个存着所有题目的批改结果，
    示范元素：
    1　＿＿＿の　ことばは　ひらがなで　どう　かきますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。
    このいすに　上着を　かけてください。
    1　うえき		2　うえぎ		3　うわき		4　うわぎ
    the right option is: 4
    the student choose: 3

    另一个存着每道题真确与否，正确为0，错误为1
"""
def return_revised_result(question_path,right_answer_path,wrong_answer_path, filename):
    #
    d = produce_split_question_list(question_path, filename)
    right_answer = read_answers_from_docx(right_answer_path)
    student_answer = read_answers_from_docx(wrong_answer_path)
    differences = read_list_difference(right_answer,student_answer)
    if len(d)-1!=len(differences):
        raise ValueError("#problems not equal to #answers")

    revise_result_all = []
    right_or_wrong = [] #right 0, wrong1
    for difference in differences:
        problem_number = difference[0]
        right_option = difference[1]
        wrong_option = difference[2]
        if right_option!=wrong_option:
            right_or_wrong.append(1)
        else:
            right_or_wrong.append(0)
        revise_result = d[problem_number+1] + "\n" + "the right option is: " +right_option +"\n"+"the student choose: "+wrong_option
        
        revise_result_all.append(revise_result)

    return revise_result_all,right_or_wrong