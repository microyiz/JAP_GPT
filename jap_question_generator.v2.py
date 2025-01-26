'''Question_Generator'''
import re
import os
import glob
import time
import string
import warnings
import docx
import mysql.connector
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate  
from langchain.chains import LLMChain   
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from typing import Any

from jap_paper_revise import produce_new_question_list
from jap_paper_revise import read_docx_to_string_with_format

def split_into_sentences(text):
    sentence_endings = re.compile(r'(?<=[。！？])\s*')
    sentences = sentence_endings.split(text)
    return sentences

def extract_numbered_content(file_path, start_number, end_number):
    """
    Extracts numbered content (e.g., 1-8) from a Word document.

    Args:
        file_path (str): The path to the Word document.
        start_number (int): The starting number of the range to extract.
        end_number (int): The ending number of the range to extract.

    Returns:
        list: A list of tuples where each tuple contains the number and the corresponding text.
    """
    # Load the Word document
    doc = Document(file_path)
    
    # Compile a regex pattern to match the numbered entries
    pattern = re.compile(rf"^\s*(\d+)\.\s*(.*)")
    
    # List to store the extracted content
    content_list = []

    # Iterate through all paragraphs in the document
    for para in doc.paragraphs:
        match = pattern.match(para.text)
        if match:
            number = int(match.group(1))
            text = match.group(2).strip()

            # Check if the number is within the specified range
            if start_number <= number <= end_number:
                content_list.append((number, text))

    return content_list



def extract_grammar_points(docx_file):
    # 读取docx文件
    doc = docx.Document(docx_file)
    
    knowledge_points = []
    
    # 遍历文档的所有段落
    for para in doc.paragraphs:
        print(f"读取的段落内容: {para.text}")
        # 使用正则表达式匹配类似 "1.～あいだ（間）" 的格式
        match = re.match(r'(\d+)\.([^\n]+)', para.text.strip())
        if match:
            # 提取编号和短语部分
            number = match.group(1)
            phrase = match.group(2).strip()
            knowledge_points.append({
                'number': number,
                'phrase': phrase
            })
    
    return knowledge_points


def extract_vocabulary(docx_path):
    doc = Document(docx_path)
    vocabulary_list = []
    number_pattern = r'^\d+\. '  # 匹配以数字和点号开头的行
    
    # 遍历每个段落
    for para in doc.paragraphs:
        # 将每个段落按行分割
        for line in para.text.splitlines():
            # 如果行包含带序号的内容
            if re.match(number_pattern, line.strip()):
                vocabulary_list.append(line.strip())
    
    return vocabulary_list




'''

def grammar_points_revise(rows, output, filepath):
    filename = os.path.splitext(os.path.basename(filepath))[0]
    llm = ChatOpenAI(
        temperature=0.6,
        model='gpt-4o'
    )

    for i in range(100):
        print(f'Processing file {i}...')
        prompt = ChatPromptTemplate.from_template(
        'You are an experienced Japanese examiner who is very familiar with the Japanese N4 and N5 exams.\n'
        'Here is a list of Japanese N4 and N5 knowledge points, including 文法(grammar): {knowledge_points}.\n'
        'Based on the knowledge point and your database related to this knowledge point, generate {num_questions} new questions.\n'
        'Please try to cover all knowledge points and reduce duplication.\n'
        'You should attach the knowledge points after each questions.\n'
        "Finally, all the answers will be attached at the end. Do not attach the answer after each question.\n"
        'The format of these new questions should be as follows:\n'

        もんだい1　（  　　　　　 ）に　何を　入れますか。　1・2・3・4から　いちばん　いい　ものを　一つ　えらんで　ください。

        ①　　かれが　手伝って　（  　　　　　 ）　宿題 (しゅくだい) が　終わらなっかった。
        1　もらったから			2　くれなかったから		
        3　ほしいから				4　ほしかったから
        -Knowledge Points: - Grammar: 10. ～ことにする, 29. ～てあげる・～てもらう・～てくれる, 11. ～ことになる

        ②　宿題 (しゅくだい) を　したのに、　先生が　（  　　　　　 ）。
        1　来なかった				2　してしまった		
        3　会わなかった			4　するつもりだった
        -Knowledge Points:   - Grammar: 12. ～し, 22. ～たら〈その(後)で〉, 7. ～くれる

        ③　うちの　子どもは　勉強 (べんきょう) しないで　（  　　　　　 ）　ばかりいる。
        1　あそび		2　あそぶ		3　あそばない		4　あそんで
        -Knowledge Points:   - Grammar: 47. ～ばかり・～てばかりいる, 58. ～ないで, 48. ～てほしい

        )
        chain_one = LLMChain(llm=llm, prompt=prompt)

        input = {
            'knowledge_points': rows,
            'num_questions': 10
        }
        revise_result = chain_one.run(input)
        output_doc = Document()
        sentences = split_into_sentences(revise_result)
        for sentence in sentences:
            sentence.replace("＿＿＿", "[ ]")
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(output, f"{filename}_new{i}.docx")
        output_doc.save(output_path)
'''

def grammar_points_revise(vocabulary_list, output, filepath):
    filename = os.path.splitext(os.path.basename(filepath))[0]
    llm = ChatOpenAI(
        temperature=0.6,
        model='gpt-4o'
    )

    i = 1
    # 遍历每个知识点（从vocabulary_list提取的知识点）
    for knowledge_point in vocabulary_list:
        print(f'Processing knowledge point: {knowledge_point}...')

        # 动态生成 prompt，要求 GPT 依次为每个词汇出题
        prompt_grammar = ChatPromptTemplate.from_template(
            f'''
    You are an experienced Japanese examiner, well-versed in the N4 and N5 levels of the Japanese Language Proficiency Test (JLPT). Your task is to create 10 multiple-choice questions based on the following grammar knowledge point: **{knowledge_point}**.

    ### Requirements:
    1. Each question should have **4 options**, with only **one correct answer**.
    2. The correct answer must strictly align with the logic of the question stem.
    3. The incorrect options should be reasonably close to the correct answer but clearly wrong when examined in context.
    4. Add logical constraints to the question to ensure that only one answer can be correct (e.g., a specific grammatical structure, verb tense, or meaning).
    5. Ensure that the correct option fits naturally in the sentence context and reflects the meaning of the grammar point.
    6. The incorrect options should deviate from the correct one in a way that seems plausible to confuse students, but they should not be grammatically correct or suitable in context.
    7. After generating the question, apply a check to ensure that the correct answer is unique, and the other options are incorrect and clearly unsuitable.
    8. Prepend the knowledge point you used to generate the question to the front of all questions.
    9. Attach all your answers at the end of the 10 questions.

            ### Question Formats:

            （  　　　　　 ）に　何を　入れますか。　1・2・3・4から　いちばん　いい　ものを　一つ　えらんで　ください。 
                **Example:**  
                Q1 かれが　手伝って　（  　　　　　 ）　宿題 (しゅくだい) が　終わらなっかった。  
                1　もらったから			2　くれなかったから		
                3　ほしいから				4　ほしかったから

                Q2 うちの　子どもは　勉強 (べんきょう) しないで　（  　　　　　 ）　ばかりいる。
                1　あそび		2　あそぶ		3　あそばない		4　あそんで

                Q3  A　「田中さんは　かのじょが　いますか。」
	                B　「いいえ、田中さんは　前の　かのじょと　別れてから、人を好き　（  　　　　　 ）。」
                1　ではありませんでした		    2　にならなくなりました		
                3　でもよくなりました			4　にしなくなりました


            **Answer:** 
            Q1. 2. くれなかったから	
            Q2. 4. あそんで
            Q3. 2. にならなくなりました 

    ### Additional Notes:
    - The generated questions must maintain high linguistic and contextual accuracy.
    - Avoid using cultural or subjective biases that could confuse learners.
    '''
        )

        
        # 创建链条来运行LLM
        chain_one = LLMChain(llm=llm, prompt=prompt_grammar)

        # 输入数据
        input_data = {
            'knowledge_point': knowledge_point,  # 使用当前知识点
            'num_questions': 10
        }

        # 获取生成的题目
        revise_result = chain_one.run(input_data)

        # 处理生成的题目并保存
        output_doc = Document()
        sentences = split_into_sentences(revise_result)
        
        # 将生成的题目逐句添加到文档中
        for sentence in sentences:
            sentence.replace("**Answers:**", "**Answers**")
            sentence = sentence.replace("＿＿＿", "[ ]")  # 替换空格部分
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(output, f"{filename}_new{i}.docx")
        output_doc.save(output_path)
        # print(f'Generated questions for {knowledge_point} saved to {output_path}')
        i += 1



def vocabulary_points_revise(vocabulary_list, output, filepath):
    filename = os.path.splitext(os.path.basename(filepath))[0]
    llm = ChatOpenAI(
        temperature=0.6,
        model='gpt-4o'
    )

    i = 1
    # 遍历每个知识点（从vocabulary_list提取的知识点）
    for knowledge_point in vocabulary_list:
        print(f'Processing knowledge point: {knowledge_point}...')

        # 动态生成 prompt，要求 GPT 依次为每个词汇出题
        prompt_vocabulary = ChatPromptTemplate.from_template(
    f'''
    You are an experienced Japanese examiner, well-versed in the N4 and N5 levels of the Japanese Language Proficiency Test (JLPT). Your task is to create 10 multiple-choice questions based on the following vocabulary knowledge point: **{knowledge_point}**.

    ### Requirements:
    1. Each question should have **4 options**, with only **one correct answer**.
    2. The correct answer must strictly match the meaning or usage of the vocabulary word, ensuring no ambiguity.
    3. The incorrect options should deviate from the correct one but be contextually plausible.
    4. Introduce logical conditions or context into the stem that restricts the possible answers, ensuring only one option fits.
    5. The options should be meaningful, with the incorrect ones being close to the correct answer, yet still clearly wrong in context.
    6. After generating the question, apply a check to ensure that only one answer is clearly correct, and the other options are incorrect.
    7. Prepend the vocabulary knowledge point to the front of all questions.
    8. Attach all your answers at the end of the 10 questions.
    9. If the question is asking how to write a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\
        If the question is asking how to write a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question.\

    ### Question Formats:
            1. **How to write in hiragana:**
                Q1: ＿＿＿の　ことばは　ひらがなで　どう　かきますか。  
                **Example:**  
                あそこに　かわいい　[鳥]が　います。  
                1. いぬ  2. とり  3. ねこ  4. むし  

            2. **Kanji recognition:**  
                Q2: ＿＿＿の　ことばは　どう　かきますか。  
                **Example:**  
                [おっと]は　今、出かけています。  
                1. 大  2. 犬  3. 太  4. 夫  

            3. **Filling in the blanks:**  
                Q3: (   　  ) に　なにを　いれますか。  
                **Example:**  
                これから　ひこうきに　（  　　　　　 ）。  
                1. おります  2. のります  3. あがります  4. のぼります  

            4. **Sentence meaning comparison:**  
                Q4: ＿＿＿の　ぶんと　だいたい　おなじ　いみの　ぶんが　あります。  
                **Example:**  
                ギターは　ちちに　ならいました。  
                1. ギターは　ちちに　もらいました。  
                2. ギターは　ちちに　えらんでもらいました。  
                3. ギターは　ちちに　おしえてもらいました。  
                4. ギターは　ちちに　かってもらいました。  

            5. **Usage of vocabulary:**  
                Q5: つぎの　ことばの　つかいかたで　いちばん　いい　ものを　1・2・3・4から　ひとつ　えらんで　ください。  
                **Example:**  
                わる  
                1. おさらを　[わって]　母に　おこられました。  
                2. おさらを　[わって]　へやに　かざりました。  
                3. おさらを　[わって]　りょうりを　つくりました。  
                4. おさらを　[わって]　コーヒーを　のみました。  

            **Answer:** 
            Q1. 2. とり  
            Q2. 4. 夫 
            Q3. 2. のります 
            Q4. 3. ギターは　ちちに　おしえてもらいました。
            Q5. 1. おさらを　[わって]　母に　おこられました。

    ### Additional Notes:
    - The generated questions must maintain high linguistic and contextual accuracy.
    - Avoid using cultural or subjective biases that could confuse learners.
    '''
)

        
        # 创建链条来运行LLM
        chain_one = LLMChain(llm=llm, prompt=prompt_vocabulary)

        # 输入数据
        input_data = {
            'knowledge_point': knowledge_point,  # 使用当前知识点
            'num_questions': 10
        }

        # 获取生成的题目
        revise_result = chain_one.run(input_data)

        # 处理生成的题目并保存
        output_doc = Document()
        sentences = split_into_sentences(revise_result)
        
        # 将生成的题目逐句添加到文档中
        for sentence in sentences:
            sentence.replace("**Answers:**", "**Answers**")
            sentence = sentence.replace("＿＿＿", "[ ]")  # 替换空格部分
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(output, f"{filename}_new{i}.docx")
        output_doc.save(output_path)
        # print(f'Generated questions for {knowledge_point} saved to {output_path}')
        i += 1



# def question_revise_simple(rows, filename, revised_newpaper_folder, max_iterations=5):
#     llm = ChatOpenAI(
#         temperature=0.8,
#         model='gpt-4o'
#     )
    
#     prompt_revise = ChatPromptTemplate.from_template(
#     "Here are the new generated Japanese practice questions: {new_paper} \
#     You are an experienced Japanese N3 examiner tasked with reviewing and ensuring that all multiple-choice test questions meet the following criteria:\n\n\
#     1. **No duplicate questions**: Ensure that all questions are unique. If a question is repeated or too similar to another, please revise it to create a new question with a distinct structure or context.\n\n\
#     2. **No duplicate options**: All options within a question should be unique, contextually meaningful, and grammatically correct. Avoid options that are too similar to each other, and ensure they are relevant to the question.\n\n\
#     3. **No duplicate correct answers**: The question should have only one correct answer. Do not allow multiple options to seem correct; ensure only one answer is suitable and unambiguous. If necessary, add specific conditions in the question stem or adjust the options to make one answer clearly correct.\n\n\
#     4. **Grammatical correctness**: The title and stem of each question must be grammatically correct. If there is an issue with grammar, revise the question stem and the options to ensure clarity and correctness.\n\n\
#     5. **Relevance of options**: The stem should clearly indicate what cannot be chosen. Ensure that one option is inappropriate or clearly wrong in context, while all other options are suitable. Avoid subjective or culturally biased content. Options should reflect the context logically.\n\n\
#     6. **Pronunciation and Word Usage**: If the question involves pronunciation, katakana, or hiragana forms, the Japanese word should be enclosed in brackets for clarity. If asking about hiragana or katakana conversion, ensure that the word is written in the appropriate form (katakana for hiragana conversion, and vice versa) and that the correct answer is not shown in the question stem.\n\n\
#     If the question is ask a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\
#     If the question is ask a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question.\
#     7. **General guidance**: Eliminate any ambiguity, revise unclear options, and avoid subjective or culturally biased phrasing. After revision, attach the correct answers separately at the very end of the document. Do not display answers within the question itself.\n\n\
#     At the end of the document, provide a summary of the changes you made (e.g., revision of question, grammatical fixes, or option rephrasing)."
#     )



    
#     chain = LLMChain(llm=llm, prompt=prompt_revise)
#     input_data = {'new_paper': rows}
    
#     for _ in range(max_iterations):
#         revised_result = chain.run(input_data)
#         if "No issues found" in revised_result or not check_for_error(revised_result):  # 假设 GPT 会返回类似 "No issues found" 的提示
#             break
#         input_data['new_paper'] = revised_result

#     # 保存最终修订结果
#     output_doc = Document()
#     sentences = split_into_sentences(revised_result)
#     for sentence in sentences:
#         output_doc.add_paragraph(sentence)

#     output_path = os.path.join(revised_newpaper_folder, f"{filename}_revised.docx")
#     output_doc.save(output_path)

def question_revise_simple(rows, filename, revised_newpaper_folder, max_iterations=5):
    llm = ChatOpenAI(
        temperature=0.8,
        model='gpt-4o'
    )
    
    prompt_revise = ChatPromptTemplate.from_template(
    "Here are the new generated Japanese practice questions: {new_paper} \
    You are an experienced Japanese N3 examiner tasked with reviewing and ensuring that all multiple-choice test questions meet the following criteria:\n\n\
    1. **No duplicate questions**: Ensure that all questions are unique. If a question is repeated or too similar to another, please revise it to create a new question with a distinct structure or context.\n\n\
    2. **No duplicate options**: All options within a question should be unique, contextually meaningful, and grammatically correct. Avoid options that are too similar to each other, and ensure they are relevant to the question.\n\n\
    3. **No duplicate correct answers**: The question should have only one correct answer. Do not allow multiple options to seem correct; ensure only one answer is suitable and unambiguous. If necessary, add specific conditions in the question stem or adjust the options to make one answer clearly correct.\n\n\
    4. **Grammatical correctness**: The title and stem of each question must be grammatically correct. If there is an issue with grammar, revise the question stem and the options to ensure clarity and correctness.\n\n\
    5. **Relevance of options**: The stem should clearly indicate what cannot be chosen. Ensure that one option is inappropriate or clearly wrong in context, while all other options are suitable. Avoid subjective or culturally biased content. Options should reflect the context logically.\n\n\
    6. **Pronunciation and Word Usage**: If the question involves pronunciation, katakana, or hiragana forms, the Japanese word should be enclosed in brackets for clarity. If asking about hiragana or katakana conversion, ensure that the word is written in the appropriate form (katakana for hiragana conversion, and vice versa) and that the correct answer is not shown in the question stem.\n\n\
    If the question is ask a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\
    If the question is ask a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question.\
    7. **General guidance**: Eliminate any ambiguity, revise unclear options, and avoid subjective or culturally biased phrasing. After revision, attach the correct answers separately at the very end of the document. Do not display answers within the question itself.\n\n\
    At the end of the document, provide a summary of the changes you made (e.g., revision of question, grammatical fixes, or option rephrasing)."
    )
    
    chain = LLMChain(llm=llm, prompt=prompt_revise)
    input_data = {'new_paper': rows}
    
    for iteration in range(max_iterations):
        revised_result = chain.run(input_data)
        errors = check_for_error(revised_result)
        
        if not errors:
            print(f"No issues found after {iteration + 1} iterations.")
            break
        
        print(f"Iteration {iteration + 1}: Detected errors - {', '.join(errors)}")
        input_data['new_paper'] = revised_result
        
        # 保存中间修订结果
        intermediate_path = os.path.join(revised_newpaper_folder, f"{filename}_iteration_{iteration + 1}.docx")
        output_doc = Document()
        sentences = split_into_sentences(revised_result)
        for sentence in sentences:
            output_doc.add_paragraph(sentence)
        output_doc.save(intermediate_path)
        
        # 保存错误日志
        log_path = os.path.join(revised_newpaper_folder, f"{filename}_error_log.txt")
        with open(log_path, 'a', encoding='utf-8') as log_file:
            log_file.write(f"Iteration {iteration + 1} Errors: {', '.join(errors)}\n")
    else:
        print(f"Maximum iterations ({max_iterations}) reached. Errors may still exist.")
    
    # 保存最终修订结果
    output_path = os.path.join(revised_newpaper_folder, f"{filename}_revised.docx")
    output_doc = Document()
    sentences = split_into_sentences(revised_result)
    for sentence in sentences:
        output_doc.add_paragraph(sentence)
    output_doc.save(output_path)



def check_for_error(revised_text):
    """
    Check for errors in the revised question set, including:
    - Multiple correct answers
    - Duplicate questions
    - Errors in the question stem
    - Duplicate options
    
    :param revised_text: The revised text output from GPT.
    :return: True if errors are found, False otherwise.
    """
    try:
        if has_multiple_correct_answers(revised_text):
            print("Error detected: Multiple correct answers")
            return True

        if has_duplicate_questions(revised_text):
            print("Error detected: Duplicate questions")
            return True

        if has_stem_errors(revised_text):
            print("Error detected: Stem errors")
            return True

        if has_duplicate_options(revised_text):
            print("Error detected: Duplicate options")
            return True

        return False  # No errors detected
    except Exception as e:
        print(f"Error in check_for_error: {e}")
        return True  # Treat as error if an exception occurs

def has_multiple_correct_answers(text):
        llm = ChatOpenAI(
            temperature=0.6,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now check the new generated Japanese practice questions: {new_paper} \
            Please revise are there multiple correct answers for the question options? \
            If so, just output True, otherwise output False"
        )

        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'new_paper': text}

        result = chain.run(input)
        bool_dict = {"True": True, "False": False, "true": True, "false":False}

        return bool_dict[result]

# def has_duplicate_questions(text):
#     """
#     Check if any questions are duplicated.
    
#     :param text: The text to check.
#     :return: True if duplicate questions are detected.
#     """
#     questions = re.findall(
#         r'\d+\.\s*(.*?)\n(?:1\.\s*.*?\n){4}',  # Regex to capture the question text only
#         text, re.DOTALL
#     )
#     seen_questions = set()
#     for question in questions:
#         question_text = question.strip()  # Normalize the question text
#         if question_text in seen_questions:
#             print(f"Duplicate question detected: {question_text}")
#             return True  # Duplicate found
#         seen_questions.add(question_text)
#     return False


def has_stem_errors(text):
    """
    Check for errors in the question stem using GPT-4.
    
    :param text: The text to check.
    :return: True if errors in the stem are detected, False otherwise.
    """
    llm = ChatOpenAI(
        temperature=0.6,  # Adjusted for more deterministic behavior
        model="gpt-4o"
    )
    prompt = ChatPromptTemplate.from_template(
        "Now check the new generated Japanese practice questions: {new_paper} \
        Are there errors in the questions (e.g., grammatical errors, ambiguous stems)? \
        If errors exist, just output True. If no errors, output False."
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    input = {'new_paper': text}

    try:
        result = chain.run(input).strip()  # Remove unnecessary whitespace or newlines
        # Normalize result to ensure consistent matching
        result = result.lower()  # Convert to lowercase for case-insensitive matching
        bool_dict = {"true": True, "false": False}
        return bool_dict.get(result, False)  # Return False if result is not in bool_dict
    except Exception as e:
        print(f"Error processing has_stem_errors: {e}")
        return False  # Default to no errors if an exception occurs

def has_duplicate_options(text):
        """
        Check for duplicate options in the questions.
        
        :param text: The text to check.
        :return: True if - options are found within a question.
        """
        # Example: Detect duplicate options for a given question.
        questions_with_options = re.findall(
            r'(\d+)\.\s*(.*?)\n(1\.\s*(.*?)\n)(2\.\s*(.*?)\n)(3\.\s*(.*?)\n)(4\.\s*(.*?)\n)',
            text, re.DOTALL
        )

        for question, _, opt1, _, opt2, _, opt3, _, opt4, _ in questions_with_options:
            options = {opt1.strip(), opt2.strip(), opt3.strip(), opt4.strip()}
            if len(options) < 4:  # If any options are duplicates
                print(f"Duplicate options detected in question {question}: {opt1.strip()}, {opt2.strip()}, {opt3.strip()}, {opt4.strip()}")
                return True
        return False




def normalize_text(text):
    """
    Normalize text by:
    1. Converting to lowercase.
    2. Removing non-essential characters such as punctuation and extra spaces.
    
    :param text: The text to normalize.
    :return: Normalized text.
    """
    # Convert to lowercase
    text = text.lower()
    
    # Remove punctuation and extra spaces
    text = text.translate(str.maketrans('', '', string.punctuation))
    text = re.sub(r'\s+', ' ', text).strip()  # Remove extra spaces
    
    return text

def has_duplicate_questions(text):
    """
    Check if any questions are duplicated, considering both the question text and options,
    while ignoring case and non-key characters like spaces and punctuation.
    
    :param text: The text to check.
    :return: True if duplicate questions are detected.
    """
    questions = re.findall(
        r'(\d+)\.\s*(.*?)\n(1\.\s*(.*?)\n)(2\.\s*(.*?)\n)(3\.\s*(.*?)\n)(4\.\s*(.*?)\n)',  # Capture question text and options
        text, re.DOTALL
    )
    
    seen_questions = set()
    
    for question, _, opt1, _, opt2, _, opt3, _, opt4, _ in questions:
        # Normalize question text and options
        question_text = normalize_text(question.strip())
        options = {normalize_text(opt1.strip()), normalize_text(opt2.strip()), 
                   normalize_text(opt3.strip()), normalize_text(opt4.strip())}
        
        # Create a normalized string for comparison: question + sorted options
        normalized_question = f"{question_text} - {', '.join(sorted(options))}"
        
        if normalized_question in seen_questions:
            print(f"Duplicate question detected: {question_text} with options {options}")
            return True  # Duplicate found
        seen_questions.add(normalized_question)
    
    return False


'''


def question_revise(rows, filename, revised_newpaper_folder, max_iterations = 5):
    llm=ChatOpenAI(
        temperature=0.8,
        model='gpt-4o'
    )
    prompt_revise = ChatPromptTemplate.from_template(
        "Now these are the new generated Japanese practice questions: {new_paper} \
            You are an excellent Japanese N3 examiner and provide students with appropriate multiple-choice test questions. All provided questions should meet the following criteria:\
            1.No duplicate questions. All the questions should be unique. Delete any repeated questions and replace them with new ones.\
            2.No duplicate options. All options should be unique and meaningful within the context of the question.\
            3.No duplicate correct answers. The answer to the question should be unique in the context of the exam. Please not have two or more than two suitable answer to choose the most suitable one, make sure it has only one suitable answer that is absolutely correct, you can add specific condition in the question stem or change the options.\
            4.Grammatical correctness. The title and stem of the question should be grammatically correct.You can put back the correct option to the question stem, if there is a grammar issue, please revise the question stem and the options.\
            5.Relevance of options. \
            Another modification idea is that the question should clearly indicate what cannot be chosen. The stem must specify the context in which one option is clearly inappropriate, while all other options are suitable.\
            So in these options, ignore the culture backgroud and avoid subjective consciousness questions and options.\
            6. If the question is about the pronunciation of a word or how a particular word is used or its katakana, hiragana, use the brackets to emphasize the Japanese words. Do not have any underline in the questions. Do not show the right answer in the question stem.\
            If the question is ask a katakana word's hiragana, make sure the word in the question is katakana and all the options are hiragana, and do not show right answer in the question.\
            If the question is ask a hiragana word's katakana, make sure the word in the question is hiragana and all the options are katakana, and do not show right answer in the question.\
            If any of the above problems occur, please modify the questions to eliminate these issues. Ensure that the structure remains the same as the original questions, and all answers should be attached at the end. Do not attach the answer after each question. Do not add questions. \
            Report the changes made at last of the file."
    )
    # rows, answer_list = loop_each_question(rows)
    # final_answer_list = []
    # for i in range(10):
    #     question_number = str(i+1) + '.'
    #     final_answer_list.append(question_number + answer_list[i])
    # answers = "**Answers**"+ "\n"+'\n'.join(final_answer_list)
    # rows = rows + answers
    rows, answer_list = loop_each_question(rows)

    if len(answer_list) < 10:
        print(f"Warning: Only {len(answer_list)} answers generated, filling with defaults.")
        while len(answer_list) < 10:
            answer_list.append("No answer provided.")

    final_answer_list = []
    for i in range(len(answer_list)):
        question_number = str(i+1) + '.'
        final_answer_list.append(question_number + answer_list[i])

    answers = "**Answers**\n" + '\n'.join(final_answer_list)
    rows = rows + answers


    chain_two = LLMChain(
        llm=llm,
        prompt=prompt_revise
    )
    input_two = {'new_paper': rows}

    for i in range(max_iterations):
        revised_result = chain_two.run(input_two)
        if check_for_error(revised_result):
            input_two['new_paper']=revised_result
        else:
            break

    output_doc = Document()
    sentences = split_into_sentences(revised_result)
        
    for sentence in sentences:
        sentence.replace("＿＿＿", "[ ]")
        output_doc.add_paragraph(sentence)

    # 路径修改
    output_path = os.path.join(revised_newpaper_folder, f"{filename}_revised.docx")
    output_doc.save(output_path)


# check each question
def loop_question(question, max_iterations):
        llm = ChatOpenAI(
            temperature= 0.6,
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now here is a Japanese practice question:{question}, please revise it to make sure that the question stem and the options have no grammarly error.\
            You should put the options back to the blank in the stem one by one, then check if the sentence has any grammar error, if so, modify it.\
            Then you need to check if there are multiple correct options for this question, if so, please revise the question stem or the options to make sure this question has only one correct answer. \
            Please not have two or more than two suitable answers to choose the most suitable one, make sure it has only one suitable answer that is absolutely correct, you can add specific condition in the question stem or change the options.\
            Only output the modified question and a single number which is the correct option of this question at last. Do not include the analysis."
        )
        # 目前问题，还是有答案错误的
        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'question':question}

        for i in range(max_iterations):
            revise_result = chain.run(input)
            if question_check(revise_result):  # Implement this method to validate output
                input['question'] = revise_result
            else:
                break  # Exit loop if no errors are found
        return revise_result

def question_check(question):
        llm = ChatOpenAI(
            temperature=0.6,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now here is a Japanese practice question:{question}, each one has four options. You need to put back each options to the question stem to check if there is any grammar issue or just check if the stem has any grammar issue.\
            If there are grammar issues, only output a single word: True, otherwise just output a single word: False. Do not output the analysis."
        )

        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'question': question}

        result = chain.run(input)
        for i in range(50):
            if(result == 'True' or result == 'true' or result == 'False' or result == 'False'):
                break
            else:
                result = chain.run(input)

        bool_dict = {"True": True, "False": False, "true": True, "false":False}

        return bool_dict[result]

def get_answer(question):
        return question[len(question)-1]

def loop_each_question(questions):
    # questions 是从gpt直接生成的result
        question_list = produce_new_question_list(questions, 10)
        question_number = 1
        new_question_list = []
        answer_list = []
        for question in question_list:
            revised_question = loop_question(question, 10)
            answer = get_answer(revised_question)
            answer_list.append(answer)
            revised_question = revised_question[:-1]

            new_question_list.append( "**"+ str(question_number)+"**" +revised_question)

            print(revised_question)
            print("end of question:",question_number)
            question_number += 1

        return '\n'.join(new_question_list), answer_list

def check_for_error(revised_text):
    """
        Check for errors in the revised question set, including:
        - Multiple correct answers
        - Duplicate questions
        - Errors in the question stem
        - Duplicate options
        
        :param revised_text: The revised text output from GPT.
        :return: True if errors are found, False otherwise.
        """

        # Error check 1: Multiple correct answers
    if has_multiple_correct_answers(revised_text):
        return True

        # Error check 2: Duplicate questions
    if has_duplicate_questions(revised_text):
        return True

        # Error check 3: Errors in the question stem
    if has_stem_errors(revised_text):
        return True

        # Error check 4: Duplicate options
    if has_duplicate_options(revised_text):
        return True

    return False  # No errors detected


def has_multiple_correct_answers(text):
        llm = ChatOpenAI(
            temperature=0.6,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now check the new generated Japanese practice questions: {new_paper} \
            Please revise are there multiple correct answers for the question options? \
            If so, just output True, otherwise output False"
        )

        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'new_paper': text}

        result = chain.run(input)
        bool_dict = {"True": True, "False": False, "true": True, "false":False}

        return bool_dict[result]

def has_duplicate_questions(text):
        """
        Check if any questions are duplicated.
        
        :param text: The text to check.
        :return: True if duplicate questions are detected.
        """
        questions = re.findall(r'\d+\.\s*(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n\d+\s(.*?)\n', 
                               text, re.DOTALL
                            )
        seen_questions = set()
        for question in questions:
            question_text = question[0].strip()  # Get the question part only
            if question_text in seen_questions:
                return True  # Duplicate found
            seen_questions.add(question)
        return False

def has_stem_errors(text):
        """
        Check for errors in the question stem.
        
        :param text: The text to check.
        :return: True if errors in the stem are detected.
        """
        # # 只是一个最简单的检查是否有标点符号的检查，需要修改为检查题目内容？
        # # Example: Detect missing or malformed question stems.
        # questions = re.findall(r'\d+\.\s*(.*?)\n', text)
        # for question in questions:
        #     if not question.strip() or len(question.strip().split()) < 5:  # Example check
        #         return True  # Detected a malformed stem
        # return False
    
        llm = ChatOpenAI(
            temperature=0.6,  # Adjusted for more deterministic behavior
            model="gpt-4o"
        )
        prompt = ChatPromptTemplate.from_template(
            "Now check the new generated Japanese practice questions: {new_paper} \
            Please revise are there some errors in the questions? \
            If so, just output True, otherwise output False"
        )

        chain = LLMChain(llm=llm, prompt = prompt)
        input = {'new_paper': text}

        result = chain.run(input)
        bool_dict = {"True": True, "False": False, "true": True, "false":False}
        return bool_dict[result]

def has_duplicate_options(text):
        """
        Check for duplicate options in the questions.
        
        :param text: The text to check.
        :return: True if - options are found within a question.
        """
        # Example: Detect duplicate options for a given question.
        questions_with_options = re.findall(
            r'(\d+)\.\s*(.*?)\n(1\.\s*(.*?)\n)(2\.\s*(.*?)\n)(3\.\s*(.*?)\n)(4\.\s*(.*?)\n)',
            text, re.DOTALL
        )

        for question, _, opt1, _, opt2, _, opt3, _, opt4, _ in questions_with_options:
            options = {opt1.strip(), opt2.strip(), opt3.strip(), opt4.strip()}
            if len(options) < 4:  # If any options are duplicates
                print(f"Duplicate options detected in question {question}: {opt1.strip()}, {opt2.strip()}, {opt3.strip()}, {opt4.strip()}")
                return True
        return False

'''




# 测试
docx_file_grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\N4N5 material\\N4 Notes 文法_numbered.docx"
docx_file_vocabulary = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\N4N5 material\\N4 Notes 語彙_numbered.docx"
test_knowledge_points = "C:\\Users\\刘宇\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\test_knowledge_points.docx"
test_grammar = "C:\\Users\\刘宇\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\test_grammar.docx"
test_vocabulary = "C:\\Users\\刘宇\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\test_vocabulary.docx"

output_grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\N4 grammar"
output_vocabulary = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\N4 vocabulary"
revised_output_vocabualry = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\revised_vocabulary"
revised_output_grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\revised_grammar"
# #grammar_points = extract_grammar_points(docx_file_grammar)
# #grammar_points_revise(grammar_points, output_grammar, docx_file_grammar)

# #rows = extract_numbered_content(test_knowledge_points, 1, 8)
# vocabulary = extract_numbered_content(test_vocabulary, 1, 6)
# grammar = extract_numbered_content(test_grammar, 1, 4)

# print(vocabulary)
# print(grammar)

# vocabulary_points_revise(vocabulary, revised_output_vocabualry, test_vocabulary)
# grammar_points_revise(grammar, revised_output_grammar, test_grammar)

# # Iterate over all the new question files and fix them
# for filepath in glob.glob(os.path.join(revised_output_vocabualry, "*.docx")):

#     filename = os.path.splitext(os.path.basename(filepath))[0]
            
#     start_time = time.time()

#     new_que = read_docx_to_string_with_format(filepath)
#     #question_revise(new_que, filename, revised_output)
#     question_revise_simple(new_que, filename, revised_output_vocabualry)

#     end_time = time.time()
#     print(f"Completed revising new questions {filename} in: {end_time - start_time:.2f} seconds")


# # Iterate over all the new question files and fix them
# for filepath in glob.glob(os.path.join(revised_output_grammar, "*.docx")):

#     filename = os.path.splitext(os.path.basename(filepath))[0]
            
#     start_time = time.time()

#     new_que = read_docx_to_string_with_format(filepath)
#     #question_revise(new_que, filename, revised_output)
#     question_revise_simple(new_que, filename, revised_output_grammar)

#     end_time = time.time()
#     print(f"Completed revising new questions {filename} in: {end_time - start_time:.2f} seconds")


test_duplicate_stem = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\AI Questions\\test_vocabulary_new1_revised.docx"
filename = os.path.splitext(os.path.basename(test_duplicate_stem))[0]
new_que = read_docx_to_string_with_format(test_duplicate_stem)
question_revise_simple(new_que, filename, revised_output_vocabualry)