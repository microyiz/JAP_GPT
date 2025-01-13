from jap_paper_revise import return_revised_result
from jap_paper_revise import return_paper
from jap_paper_revise import extract_student_id
from jap_paper_revise import read_name_from_docx
from jap_paper_revise import read_answers_from_docx
import mysql.connector
import os
import glob
from docx import Document
from db_util import drop_table_query ,create_table_query ,insert_query,show_fiverows_query,select_mistake_query,db
from db_question_students_results import insert_questions_query, insert_students_query, insert_exam_results_query,select_mistake_query,select_all_query,db
from join_search import join_search

'''
# 连接到 MySQL 数据库
question_path = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\sample paper and answer\\Test 1 Question Paper.docx"
right_answer_path = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\sample paper and answer\\Test 1 Model Answer.docx"
wrong_answer_path = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\基于大模型的学习平台开发\\template paper from CUHK\\student paper\\1155142665 Test 1.docx"
filename = os.path.splitext(os.path.basename(question_path))[0]
student_id = extract_student_id(wrong_answer_path)

answer = return_revised_result(question_path, right_answer_path, wrong_answer_path, filename)


cursor = db.cursor()
#删除已经有的table
cursor.execute(drop_table_query)
#创建新的jap_table
cursor.execute(create_table_query)
revised_problem_answer_list = answer[0]
right_or_not = answer[1]
#insert all the data into the table
for i in range(len(right_or_not)):
    cursor.execute(insert_query, (student_id, revised_problem_answer_list[i], right_or_not[i]))
#show the first five rows of the table
# cursor.execute(show_fiverows_query)
# rows = cursor.fetchall()
# for row in rows:
#     print(row)

cursor.execute(select_mistake_query, (student_id,))
rows = cursor.fetchall()
for row in rows:
    print(row)



table_creation_query = """
CREATE TABLE jap_table (
    id INT AUTO_INCREMENT PRIMARY KEY,
    student_id VARCHAR(255) NOT NULL,
    question_text TEXT NOT NULL,
    is_correct BOOLEAN NOT NULL
)
"""
cursor.execute(table_creation_query)
drop_jap_table(db.cursor)
create_jap_table(db.cursor)

def process_students_results(question_path, wrong_answer_path):
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_id = extract_student_id(wrong_answer_path)
    student_name = read_name_from_docx(wrong_answer_path)
    student_email = student_id + "@link.cuhk.edu.cn"

    cursor = db.cursor()
    cursor.execute(drop_students_table_query)
    cursor.execute(create_students_table_query)

    # Insert student information into the database
    cursor.execute(insert_students_query, (student_id, student_name, student_email))
    db.commit()
    rows = cursor.fetchall()

    return rows

def process_students_results(question_path, wrong_answer_path):
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_id = extract_student_id(wrong_answer_path)
    student_name = read_name_from_docx(wrong_answer_path)
    student_email = student_id + "@link.cuhk.edu.cn"

    cursor = db.cursor()

    # 先删除与该学生相关的 exam_results 记录
    cursor.execute("DELETE FROM exam_results WHERE student_id = %s", (student_id,))
    db.commit()

    # 删除 `students` 表
    cursor.execute(drop_students_table_query)
    cursor.execute(create_students_table_query)

    # 插入学生信息到数据库
    cursor.execute(insert_students_query, (student_id, student_name, student_email))
    db.commit()

    # 查询结果（根据需要修改）
    cursor.execute("SELECT * FROM students WHERE student_id = %s", (student_id,))
    rows = cursor.fetchall()

    return rows


生成学生表 students
def process_students_results(question_path, wrong_answer_path):
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_id = extract_student_id(wrong_answer_path)
    student_name = read_name_from_docx(wrong_answer_path)
    student_email = student_id + "@link.cuhk.edu.cn"

    cursor = db.cursor()

    # 禁用外键约束
    cursor.execute("SET foreign_key_checks = 0")

    # 先删除与该学生相关的 exam_results 记录
    cursor.execute("DELETE FROM exam_results WHERE student_id = %s", (student_id,))
    db.commit()

    # 删除 `students` 表
    cursor.execute(drop_students_table_query)
    cursor.execute(create_students_table_query)

    # 插入学生信息到数据库
    cursor.execute(insert_students_query, (student_id, student_name, student_email))
    db.commit()

    # 启用外键约束
    cursor.execute("SET foreign_key_checks = 1")

    # 查询学生信息，打印查询 SQL 和结果
    cursor.execute("SELECT * FROM students WHERE student_no = %s", (student_id,))
    rows = cursor.fetchall()

    # 打印查询的 SQL 和结果
    print("SELECT * FROM students WHERE student_no = ", student_id)
    print("Query result:", rows)

    return rows


生成学生表 students
def process_students_results(question_path, wrong_answer_path):
    # 获取文件名和学生信息
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_id = extract_student_id(wrong_answer_path)
    student_name = read_name_from_docx(wrong_answer_path)
    student_email = student_id + "@link.cuhk.edu.cn"

    cursor = db.cursor()

    # 禁用外键约束（防止删除操作失败）
    cursor.execute("SET FOREIGN_KEY_CHECKS=0;")

    # 删除与该学生相关的 exam_results 表记录
    cursor.execute("DELETE FROM exam_results WHERE student_id = %s", (student_id,))
    db.commit()

    # 插入或更新学生信息
    cursor.execute("""
        INSERT INTO students (student_no, name, email)
        VALUES (%s, %s, %s)
        ON DUPLICATE KEY UPDATE
            name = VALUES(name),
            email = VALUES(email)
    """, (student_id, student_name, student_email))
    db.commit()

    # 启用外键约束
    cursor.execute("SET FOREIGN_KEY_CHECKS=1;")

    # 查询学生信息，打印查询 SQL 和结果
    cursor.execute("SELECT * FROM students WHERE student_no = %s", (student_id,))
    rows = cursor.fetchall()

    # 打印查询的 SQL 和结果
    print("SELECT * FROM students WHERE student_no = ", student_id)
    print("Query result:", rows)

    return rows

'''


'''

# 生成学生表 students
def process_students_results(wrong_answer_path):
    """
    处理学生数据，将学生信息插入或更新到 students 表中。
    使用 student_id作为主键。
    """
    # 从文件路径中提取学生信息
    student_no = extract_student_id(wrong_answer_path)  # 提取学号
    student_name = read_name_from_docx(wrong_answer_path)  # 提取学生姓名
    student_email = f"{student_no}@link.cuhk.edu.cn"  # 自动生成学生邮箱

    cursor = db.cursor()

    # 检查当前使用的数据库
    cursor.execute("SELECT DATABASE();")
    current_db = cursor.fetchone()
    print(f"Current database: {current_db[0]}")  # 输出当前使用的数据库名称

    # 插入或更新学生信息
    try:
        # cursor.execute("""
        #     INSERT INTO students (student_no, name, email)
        #     VALUES (%s, %s, %s)
        #     ON DUPLICATE KEY UPDATE
        #         name = VALUES(name),
        #         email = VALUES(email)
        # """, (student_no, student_name, student_email))
        cursor.execute(insert_students_query, (student_no, student_name, student_email))
        db.commit()
        print(f"Student {student_no} inserted/updated successfully.")
    except Exception as e:
        db.rollback()
        print(f"Error inserting/updating student {student_no}: {e}")
        return None

    # 查询学生是否成功插入/更新
    try:
        cursor.execute("SELECT student_no FROM students WHERE student_no = %s", (student_no,))
        result = cursor.fetchone()
        if result:
            queried_student_no = result[0]
            print(f"Queried student_no: {queried_student_no}")
        else:
            print(f"Student No {student_no} not found in students table!")
            return None
    except Exception as e:
        print(f"Error querying student_no {student_no}: {e}")
        return None
    
    print("Checking students table structure...")
    cursor.execute("DESCRIBE students;")
    table_structure = cursor.fetchall()
    print("Students table structure:", table_structure)
    print("Executing SQL query:")  # 打印执行的 SQL 语句
    print(cursor.statement)  # 打印 MySQL Connector 的最后执行语句

    return student_no



# 生成题目表 questions
def process_paper(question_path, right_answer_path):
    # 获取文件名和正确答案
    filename = os.path.splitext(os.path.basename(question_path))[0]
    right_answer = read_answers_from_docx(right_answer_path)
    answer = return_paper(question_path, right_answer_path, filename)
    revised_problem_answer_list = answer[0]
    revised_knowledge_points = answer[1]
    revised_difficulty_level = answer[2]

    print(f"Right answers: {right_answer}")  # 打印正确答案，检查是否正确读取

    cursor = db.cursor()

    # 检查并清空题目表
    cursor.execute("SET FOREIGN_KEY_CHECKS=0;")
    cursor.execute("TRUNCATE TABLE questions;")
    cursor.execute("SET FOREIGN_KEY_CHECKS=1;")

    # 插入新数据至 questions 表
    for i in range(len(revised_problem_answer_list)):
        question_index = f"{filename}{i}"
        content = revised_problem_answer_list[i]
        correct_answer = right_answer[i]
        knowledge_points = revised_knowledge_points[i]
        difficulty_level = revised_difficulty_level[i]


        print(f"Inserting question_index: {question_index}, content: {content}, correct_answer: {correct_answer}")

        try:
            cursor.execute(
                insert_questions_query,
                (question_index, content, correct_answer, knowledge_points, difficulty_level, 0)
            )
        except Exception as e:
            print(f"Error inserting question_index {question_index}: {e}")
            db.rollback()  # 如果插入失败，回滚事务
            raise e

    db.commit()

    # 查询插入的题目
    query = "SELECT * FROM questions WHERE question_index LIKE %s"
    cursor.execute(query, (f"{filename}%",))
    rows = cursor.fetchall()

    # 打印查询结果
    print(f"SELECT * FROM questions WHERE question_index LIKE '{filename}%'")
    print("Query result:", rows)

    return rows




# 考试结果表
def process_exam_results(question_path, right_answer_path, wrong_answer_path):
    """
    处理学生考试结果，确保学生、题目、以及考试记录正确插入/更新到数据库中。
    使用 student_id 作为学生的唯一标识。
    """
    # 获取文件名和学生学号
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_no = extract_student_id(wrong_answer_path)  # 提取学生学号
    student_answer = read_answers_from_docx(wrong_answer_path)

    print(f"Processing paper for student_no: {student_no}, filename: {filename}")

    # 获取修正后的结果
    answer = return_revised_result(question_path, right_answer_path, wrong_answer_path, filename)
    revised_problem_answer_list = answer[0]
    right_or_not = answer[1]

    print(f"Revised problem and answers list: {revised_problem_answer_list}")
    print(f"Right or not list: {right_or_not}")

    cursor = db.cursor()

    # 确保学生信息已经存在
    student_name = read_name_from_docx(wrong_answer_path)  # 从文档中提取学生姓名
    student_email = f"{student_no}@link.cuhk.edu.cn"  # 自动生成学生邮箱
    print(f"Inserting or updating student: No={student_no}, Name={student_name}, Email={student_email}")

    # 插入或更新学生信息
    # cursor.execute(""" 
    #     INSERT INTO students (student_no, name, email)
    #     VALUES (%s, %s, %s)
    #     ON DUPLICATE KEY UPDATE
    #         name = VALUES(name),
    #         email = VALUES(email)
    # """, (student_no, student_name, student_email))
    cursor.execute(insert_students_query, (student_no, student_name, student_email))
    db.commit()

    # 获取 student_id
    cursor.execute("SELECT student_id FROM students WHERE student_no = %s", (student_no,))
    student_row = cursor.fetchone()
    if not student_row:
        raise ValueError(f"Student with student_no={student_no} not found in database.")
    student_id = student_row[0]

    print(f"Student_id for {student_no} retrieved: {student_id}")


    answer = return_paper(question_path, right_answer_path, filename)
    revised_problem_list = answer[0]
    revised_knowledge_points = answer[1]
    revised_difficulty_level = answer[2]

    # 确保题目已经存在
    for i in range(len(revised_problem_list)):
        question_index = f"{filename}{i}"  # 生成题目索引
        content = revised_problem_list[i]
        correct_answer = right_answer_path[i]
        knowledge_points = revised_knowledge_points[i]
        difficulty_level = revised_difficulty_level[i]

        print(f"Inserting or updating question: Index={question_index}, Content={content}, Correct Answer={correct_answer}")

        # 插入或更新题目
        # cursor.execute("""
        #     INSERT INTO questions (question_index, content, correct_answer, type, level, is_gpt)
        #     VALUES (%s, %s, %s, %s, %s, %s)
        #     ON DUPLICATE KEY UPDATE
        #         content = VALUES(content),
        #         correct_answer = VALUES(correct_answer)
        # """, (question_index, content, correct_answer, "VOCABULARY", "N4", False))
        cursor.execute(insert_questions_query, (question_index, content, correct_answer, knowledge_points, difficulty_level, 0))
    db.commit()

    print(f"Questions for filename '{filename}' processed successfully.")

    # 清理 exam_results 表中与当前学生相关的记录
    cursor.execute("""
    DELETE FROM exam_results
    WHERE student_id = %s
    """, (student_id,))
    db.commit()

    print(f"Cleared previous exam results for student_id {student_id}.")

    # 插入考试结果到 exam_results 表
    for i in range(len(right_or_not)):
        question_index = f"{filename}{i}"

        # 查询 question_id
        cursor.execute("SELECT question_id FROM questions WHERE question_index = %s", (question_index,))
        question_row = cursor.fetchone()
        if not question_row:
            raise ValueError(f"Question not found for index: {question_index}")
        question_id = question_row[0]

        print(f"Found question_id={question_id} for question_index={question_index}")

        # 插入考试结果，使用 student_id 作为外键
        print(f"Inserting exam result: Question ID={question_id}, Student ID={student_id}, Answer={student_answer[i]}, Is Correct={right_or_not[i]}")

        # cursor.execute("""
        #     INSERT INTO exam_results (question_id, student_id, student_answer, is_correct)
        #     VALUES (%s, %s, %s, %s)
        # """, (question_id, student_id, student_answer[i], right_or_not[i]))
        cursor.execute(insert_exam_results_query, (question_id, student_id, student_answer[i], right_or_not[i]))
    db.commit()

    print(f"Exam results for student_id {student_id} processed successfully.")

    # 查询学生的错误
    print(f"Fetching incorrect answers for student_id: {student_id}")
    # cursor.execute("""
    #     SELECT * FROM exam_results
    #     WHERE student_id = %s AND is_correct = 0
    # """, (student_id,))
    cursor.execute(select_mistake_query, (student_id,))
    rows = cursor.fetchall()

    print(f"Incorrect answers for student_id {student_id}: {rows}")

    return rows
'''


def insert_or_update_student(student_no, student_name, student_email):
    cursor = db.cursor()
    
    # 检查学生是否已经存在
    cursor.execute("SELECT student_id FROM students WHERE student_no = %s", (student_no,))
    existing_student = cursor.fetchone()

    if existing_student:
        # 学生已经存在，更新学生信息
        cursor.execute("UPDATE students SET name = %s WHERE student_no = %s", (student_name, student_no, student_email))
        db.commit()
        print(f"Student {student_no} updated")
        return existing_student[0]  # 返回已有的 student_id
    else:
        # 学生不存在，插入新记录
        cursor.execute(insert_students_query, (student_no, student_name, student_email))
        db.commit()
        student_id = cursor.lastrowid
        print(f"Student {student_no} inserted")
        return student_id



def insert_or_update_question(question_index, content, correct_answer, knowledge_points, difficulty_level):
    """插入或更新题目信息"""
    cursor = db.cursor()
    try:
        cursor.execute(insert_questions_query, (question_index, content, correct_answer, knowledge_points, difficulty_level, 0))
        db.commit()
        print(f"Question {question_index} inserted/updated successfully.")
    except Exception as e:
        db.rollback()
        print(f"Error inserting/updating question {question_index}: {e}")
        return None

'''
def process_exam_results(question_path, right_answer_path, wrong_answer_path):
    """处理学生考试结果"""
    # 获取文件名和学生学号
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_no = extract_student_id(wrong_answer_path)  # 提取学生学号
    student_name = read_name_from_docx(wrong_answer_path)  # 提取学生姓名
    student_answer = read_answers_from_docx(wrong_answer_path)

    # 插入/更新学生信息
    student_id = insert_or_update_student(student_no, student_name)
    if student_id is None:
        return None

    print(f"Processing paper for student_no: {student_no}, filename: {filename}")

    # 获取修正后的结果
    paper_content = return_paper(question_path, right_answer_path, filename)
    revised_problem_list = paper_content[0]
    revised_knowledge_points = paper_content[1]
    revised_difficulty_level = paper_content[2]

    answer = return_revised_result(question_path, right_answer_path, wrong_answer_path, filename)
    right_or_not = answer[1]

    # 插入/更新题目数据
    for i in range(len(revised_problem_list)):
        question_index = f"{filename}{i}"
        content = revised_problem_list[i]
        correct_answer = right_answer_path[i]
        knowledge_points = revised_knowledge_points[i]
        difficulty_level = revised_difficulty_level[i]
        insert_or_update_question(question_index, content, correct_answer, knowledge_points, difficulty_level)

    # 清理 exam_results 表中与当前学生相关的记录
    cursor = db.cursor()
    cursor.execute("DELETE FROM exam_results WHERE student_id = %s", (student_id,))
    db.commit()

    # 插入考试结果
    for i in range(len(right_or_not)):
        question_index = f"{filename}{i}"
        # 查询 question_id
        cursor.execute("SELECT question_id FROM questions WHERE question_index = %s", (question_index,))
        question_row = cursor.fetchone()
        if not question_row:
            print(f"Question not found for index: {question_index}")
            continue
        question_id = question_row[0]
        cursor.execute(insert_exam_results_query, (question_id, student_id, student_answer[i], right_or_not[i]))
    db.commit()

    # 查询学生的错误
    cursor.execute(select_mistake_query)
    rows = cursor.fetchall()
    print(f"Incorrect answers for student_id {student_id}: {rows}")

    return rows
'''

'''
def process_paper_and_store_results(question_path, right_answer_path, wrong_answer_path):
    # 获取文件名和学生ID
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_id = extract_student_id(wrong_answer_path)

    # 获取修正后的结果
    answer = return_revised_result(question_path, right_answer_path, wrong_answer_path, filename)
    revised_problem_answer_list = answer[0]
    right_or_not = answer[1]

    cursor = db.cursor()
    
    # 删除已经存在的表并创建新表
    cursor.execute(drop_table_query)
    cursor.execute(create_table_query)
    
    # 插入数据到数据库
    for i in range(len(right_or_not)):
        cursor.execute(insert_query, (student_id, revised_problem_answer_list[i], right_or_not[i]))
    db.commit()

    # 查询学生的错误
    cursor.execute(select_mistake_query, (student_id,))
    rows = cursor.fetchall()

    return rows
'''

def process_student_results(question_path, right_answer_path, wrong_answer_path):
    """处理学生的考试结果并插入数据库"""
    # 获取文件名 学生学号/邮箱 学生答案
    filename = os.path.splitext(os.path.basename(question_path))[0]
    student_no = extract_student_id(wrong_answer_path)  # 提取学生学号
    student_name = read_name_from_docx(wrong_answer_path)  # 提取学生姓名
    student_email = f"{student_no}@link.cuhk.edu.cn"  # 自动生成学生邮箱
    student_answer = read_answers_from_docx(wrong_answer_path)

    # 插入/更新学生信息
    student_id = insert_or_update_student(student_no, student_name, student_email)
    if student_id is None:
        return None

    print(f"Processing paper for student_no: {student_no}, filename: {filename}")

    # 获取修正后的结果
    paper_content = return_paper(question_path, right_answer_path, filename)
    revised_problem_list = paper_content[0]
    revised_knowledge_points = paper_content[1]
    revised_difficulty_level = paper_content[2]

    answer = return_revised_result(question_path, right_answer_path, wrong_answer_path, filename)
    right_or_not = answer[1]
    # 获取正确答案
    right_option = read_answers_from_docx(right_answer_path)

    # 插入/更新题目数据
    for i in range(len(revised_problem_list)):
        question_index = f"{filename}{i}"
        content = revised_problem_list[i]
        correct_answer = right_option[i]
        knowledge_points = revised_knowledge_points[i]
        difficulty_level = revised_difficulty_level[i]
        insert_or_update_question(question_index, content, correct_answer, knowledge_points, difficulty_level)

    # 清理 exam_results 表中与当前学生相关的记录
    cursor = db.cursor()
    cursor.execute("DELETE FROM exam_results WHERE student_id = %s", (student_id,))
    db.commit()

    # 插入考试结果
    for i in range(len(right_or_not)):
        question_index = f"{filename}{i}"
        cursor.execute("SELECT question_id FROM questions WHERE question_index = %s", (question_index,))
        question_row = cursor.fetchone()
        if not question_row:
            print(f"Question not found for index: {question_index}")
            continue
        question_id = question_row[0]
        cursor.execute(insert_exam_results_query, (question_id, student_id, student_answer[i], right_or_not[i]))
    db.commit()

    # 查询学生的错误
    cursor.execute(select_mistake_query)
    rows = cursor.fetchall()
    print(f"Incorrect answers for student_id {student_id}: {rows}")

    return rows

def process_and_save_to_word(question_path, right_answer_path, wrong_answer_path, sample_output):
    """处理考试结果并保存到 Word 文档"""
    # 处理学生考试结果
    process_student_results(question_path, right_answer_path, wrong_answer_path)

    # 查询数据库中的数据
    cursor = db.cursor()

    # 查询 students 表数据
    cursor.execute("SELECT student_no, name, email FROM students")
    students = cursor.fetchall()
    students_data = [{"student_no": row[0], "name": row[1], "email": row[2]} for row in students]

    # 查询 questions 表数据
    cursor.execute("SELECT question_index, content, correct_answer, type, level, is_gpt FROM questions")
    questions = cursor.fetchall()
    questions_data = [{"question_index": row[0], "content": row[1], "correct_answer": row[2], "type": row[3], "level": row[4], "is_gpt": row[5]} for row in questions]

    # 查询 exam_results 表数据
    cursor.execute("SELECT DISTINCT result_id, question_id, student_id, student_answer, is_correct FROM exam_results")
    exam_results = cursor.fetchall()
    exam_results_data = [{"result_id": row[0], "question_id": row[1], "student_id": row[2], "student_answer": row[3], "is_correct": row[4]} for row in exam_results]

    # 保存到 Word 文档
    save_to_word(students_data, questions_data, exam_results_data, sample_output)


# 保存数据到 Word 文档
def save_to_word(students, questions, exam_results, sample_output):
    doc = Document()
    
    # 添加标题
    doc.add_heading('Student Results Report', 0)
    
    # 保存学生表内容
    doc.add_heading('Students Table', level=1)
    table = doc.add_table(rows=1, cols=3)  # 3列: student_no, name, email
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Student No'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Email'
    
    for student in students:
        row_cells = table.add_row().cells
        row_cells[0].text = str(student['student_no'])
        row_cells[1].text = student['name']
        row_cells[2].text = student['email']
    
    doc.add_paragraph()  # Add a space between sections

    # 保存题目表内容
    doc.add_heading('Questions Table', level=1)
    table = doc.add_table(rows=1, cols=6)  # 6列: question_index, content, correct_answer, type, level, is_gpt
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question Index'
    hdr_cells[1].text = 'Content'
    hdr_cells[2].text = 'Correct Answer'
    hdr_cells[3].text = 'Type'
    hdr_cells[4].text = 'Level'
    hdr_cells[5].text = 'Is GPT'
    
    for question in questions:
        row_cells = table.add_row().cells
        row_cells[0].text = question['question_index']
        row_cells[1].text = question['content']
        row_cells[2].text = question['correct_answer']
        row_cells[3].text = question['type']
        row_cells[4].text = question['level']
        row_cells[5].text = str(question['is_gpt'])
    
    doc.add_paragraph()  # Add a space between sections

    # 保存考试结果表内容
    doc.add_heading('Exam Results Table', level=1)
    table = doc.add_table(rows=1, cols=5)  # 5列: result_id, question_id, student_id, student_answer, is_correct
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Result ID'
    hdr_cells[1].text = 'Question ID'
    hdr_cells[2].text = 'Student ID'
    hdr_cells[3].text = 'Student Answer'
    hdr_cells[4].text = 'Is Correct'
    
    for result in exam_results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(result['result_id'])
        row_cells[1].text = str(result['question_id'])
        row_cells[2].text = str(result['student_id'])
        row_cells[3].text = str(result['student_answer'])
        row_cells[4].text = str(result['is_correct'])
    
    # 保存文件
    doc.save(sample_output)
    print(f"Report saved")

# # 处理考试结果并保存数据到 Word 文档
# def process_and_save_to_word(question_path, right_answer_path, wrong_answer_path, sample_output):
#     # 假设 students, questions 和 exam_results 数据已经处理好并存在
#     cursor = db.cursor()

#     # 查询 students 表数据
#     cursor.execute("SELECT student_no, name, email FROM students")
#     students = cursor.fetchall()
#     students_data = [{"student_no": row[0], "name": row[1], "email": row[2]} for row in students]

#     # 查询 questions 表数据
#     cursor.execute("SELECT question_index, content, correct_answer, type, level, is_gpt FROM questions")
#     questions = cursor.fetchall()
#     questions_data = [{"question_index": row[0], "content": row[1], "correct_answer": row[2], "type": row[3], "level": row[4], "is_gpt": row[5]} for row in questions]

#     # 查询 exam_results 表数据并去重
#     cursor.execute("SELECT DISTINCT result_id, question_id, student_id, student_answer, is_correct FROM exam_results")
#     exam_results = cursor.fetchall()
#     exam_results_data = [{"result_id": row[0], "question_id": row[1], "student_id": row[2], "student_answer": row[3], "is_correct": row[4]} for row in exam_results]

#     # 去重：确保没有重复的记录（如果查询时未完全去重，可以在此处处理）
#     exam_results_data = list({v['result_id']: v for v in exam_results_data}.values())

#     # 处理学生考试结果
#     student_no = extract_student_id(wrong_answer_path)  # 提取学生学号
#     student_name = read_name_from_docx(wrong_answer_path)  # 提取学生姓名
#     student_answer = read_answers_from_docx(wrong_answer_path)

#     # 插入/更新学生信息
#     student_id = insert_or_update_student(student_no, student_name)
#     if student_id is None:
#         return None

#     print(f"Processing paper for student_no: {student_no}, filename: {os.path.basename(question_path)}")

#     # 获取修正后的结果
#     paper_content = return_paper(question_path, right_answer_path, os.path.basename(question_path))
#     revised_problem_list = paper_content[0]
#     revised_knowledge_points = paper_content[1]
#     revised_difficulty_level = paper_content[2]

#     answer = return_revised_result(question_path, right_answer_path, wrong_answer_path, os.path.basename(question_path))
#     right_or_not = answer[1]

#     # 插入/更新题目数据
#     for i in range(len(revised_problem_list)):
#         question_index = f"{os.path.basename(question_path)}{i}"
#         content = revised_problem_list[i]
#         correct_answer = right_answer_path[i]
#         knowledge_points = revised_knowledge_points[i]
#         difficulty_level = revised_difficulty_level[i]
#         insert_or_update_question(question_index, content, correct_answer, knowledge_points, difficulty_level)

#     # 清理 exam_results 表中与当前学生相关的记录
#     cursor.execute("DELETE FROM exam_results WHERE student_id = %s", (student_id,))
#     db.commit()

#     # 插入考试结果
#     for i in range(len(right_or_not)):
#         question_index = f"{os.path.basename(question_path)}{i}"
#         # 查询 question_id
#         cursor.execute("SELECT question_id FROM questions WHERE question_index = %s", (question_index,))
#         question_row = cursor.fetchone()
#         if not question_row:
#             print(f"Question not found for index: {question_index}")
#             continue
#         question_id = question_row[0]
#         cursor.execute(insert_exam_results_query, (question_id, student_id, student_answer[i], right_or_not[i]))
#     db.commit()

#     # 保存到 Word 文档
#     save_to_word(students_data, questions_data, exam_results_data, sample_output)




def main():
    question_path = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\processed test paper with knowledge points\\Test 1 Question Paper.docx"
    right_answer_path = "C:\\Users\\刘宇\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\Test1_new\\test 1 paper\\Test 1 Model Answer.docx"
    #wrong_answer_path = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\Test1_new\\student paper_test\\1155193734 Test 1.docx"
    input_test_path = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\Test1_new\\student paper_test"
    sample_output = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_db_test.docx"
    
    for filepath in glob.glob(os.path.join(input_test_path, "*.docx")):
        process_and_save_to_word(question_path, right_answer_path, filepath, sample_output)

    join_search()
    


if __name__ == "__main__":
    try:
       cursor = db.cursor()
    except mysql.connector.errors.OperationalError as e:
        print(f"Connection error: {e}")
        db.reconnect()  # 重新连接
        cursor = db.cursor()
    main()





'''

# 保存数据到 Word 文档
def save_to_word(students, questions, exam_results, sample_output):
    doc = Document()
    
    # 添加标题
    doc.add_heading('Student Results Report', 0)
    
    # 保存学生表内容
    doc.add_heading('Students Table', level=1)
    table = doc.add_table(rows=1, cols=3)  # 3列: student_no, name, email
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Student No'
    hdr_cells[1].text = 'Name'
    hdr_cells[2].text = 'Email'
    
    for student in students:
        row_cells = table.add_row().cells
        row_cells[0].text = str(student['student_no'])
        row_cells[1].text = student['name']
        row_cells[2].text = student['email']
    
    doc.add_paragraph()  # Add a space between sections

    # 保存题目表内容
    doc.add_heading('Questions Table', level=1)
    table = doc.add_table(rows=1, cols=6)  # 6列: question_index, content, correct_answer, type, level, is_gpt
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Question Index'
    hdr_cells[1].text = 'Content'
    hdr_cells[2].text = 'Correct Answer'
    hdr_cells[3].text = 'Type'
    hdr_cells[4].text = 'Level'
    hdr_cells[5].text = 'Is GPT'
    
    for question in questions:
        row_cells = table.add_row().cells
        row_cells[0].text = question['question_index']
        row_cells[1].text = question['content']
        row_cells[2].text = question['correct_answer']
        row_cells[3].text = question['type']
        row_cells[4].text = question['level']
        row_cells[5].text = str(question['is_gpt'])
    
    doc.add_paragraph()  # Add a space between sections

    # 保存考试结果表内容
    doc.add_heading('Exam Results Table', level=1)
    table = doc.add_table(rows=1, cols=5)  # 5列: result_id, question_id, student_id, student_answer, is_correct
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Result ID'
    hdr_cells[1].text = 'Question ID'
    hdr_cells[2].text = 'Student ID'
    hdr_cells[3].text = 'Student Answer'
    hdr_cells[4].text = 'Is Correct'
    
    for result in exam_results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(result['result_id'])
        row_cells[1].text = str(result['question_id'])
        row_cells[2].text = str(result['student_id'])
        row_cells[3].text = str(result['student_answer'])
        row_cells[4].text = str(result['is_correct'])
    
    # 保存文件
    doc.save(sample_output)
    print(f"Report saved")

# 读取表格并调用 save_to_word 方法
def process_and_save_to_word():
    # 假设 students, questions 和 exam_results 数据已经处理好并存在
    cursor = db.cursor()

    # 查询 students 表数据
    cursor.execute("SELECT student_no, name, email FROM students")
    students = cursor.fetchall()
    students_data = [{"student_no": row[0], "name": row[1], "email": row[2]} for row in students]

    # 查询 questions 表数据
    cursor.execute("SELECT question_index, content, correct_answer, type, level, is_gpt FROM questions")
    questions = cursor.fetchall()
    questions_data = [{"question_index": row[0], "content": row[1], "correct_answer": row[2], "type": row[3], "level": row[4], "is_gpt": row[5]} for row in questions]

    # 查询 exam_results 表数据并去重
    cursor.execute("SELECT DISTINCT result_id, question_id, student_id, student_answer, is_correct FROM exam_results")
    exam_results = cursor.fetchall()
    exam_results_data = [{"result_id": row[0], "question_id": row[1], "student_id": row[2], "student_answer": row[3], "is_correct": row[4]} for row in exam_results]

    # 去重：确保没有重复的记录（如果查询时未完全去重，可以在此处处理）
    exam_results_data = list({v['result_id']: v for v in exam_results_data}.values())

    # 保存到 Word 文档
    save_to_word(students_data, questions_data, exam_results_data, sample_output)



# 学生表生成检测
# for filepath in glob.glob(os.path.join(input_test_path, "*.docx")):
#     result_rows = process_students_results(filepath)

# 题目表生成检测
#paper_rows = process_paper(question_path, right_answer_path)

# 考试结果表生成检测
# for filepath in glob.glob(os.path.join(input_test_path, "*.docx")):
#     exam_results = process_exam_results(question_path, right_answer_path, filepath)
# 保存数据到 Word
process_and_save_to_word()


'''

'''

def check_table_structure():
    cursor = db.cursor()
    
    # 检查 students 表结构
    cursor.execute("DESCRIBE students;")
    students_table = cursor.fetchall()
    print("Students table structure:")
    for column in students_table:
        print(column)
    
    # 检查 questions 表结构
    cursor.execute("DESCRIBE questions;")
    questions_table = cursor.fetchall()
    print("Questions table structure:")
    for column in questions_table:
        print(column)
    
    # 检查 exam_results 表结构
    cursor.execute("DESCRIBE exam_results;")
    exam_results_table = cursor.fetchall()
    print("Exam_results table structure:")
    for column in exam_results_table:
        print(column)

# 执行检查
check_table_structure()



def check_primary_key():
    cursor = db.cursor()
    
    # 查询 students 表的索引
    cursor.execute("SHOW INDEXES FROM students;")
    indexes = cursor.fetchall()
    print("Indexes in students table:")
    for index in indexes:
        print(index)
    
    # 检查是否有 student_no 作为主键或唯一索引
    for index in indexes:
        if 'student_no' in index:
            print(f"Found index for student_no: {index}")
            break
    else:
        print("No index found for student_no in students table!")

# 执行检查
check_primary_key()

'''