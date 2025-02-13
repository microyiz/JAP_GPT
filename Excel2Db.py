import os
import pandas as pd
import mysql.connector
from docx import Document
import re
import shutil

db = mysql.connector.connect(
    host="localhost",
    user="root",
    password="123",
    database="japgpt"
)
cursor = db.cursor()

def read_N4_Vocabulary_knowledge_point():
    doc = Document(r"D:\JAP_GPT\JAP_GPT\N4N5 material\N4 Notes 語彙_numbered.docx")
    read_N4_Vocabulary_knowledge_point = []
    for para in doc.paragraphs:
        matches = re.findall(r"\.(.*?)\n", para.text + "\n", re.DOTALL)
        for m in matches:
            read_N4_Vocabulary_knowledge_point.append("- Vocabulary: ・語彙 " + m.strip())
    return read_N4_Vocabulary_knowledge_point

def read_N4_Grammar_knowledge_point():
    doc = Document(r"D:\JAP_GPT\JAP_GPT\N4N5 material\N4 Notes 文法_numbered.docx")
    read_N4_Grammar_knowledge_point = []
    for para in doc.paragraphs:
        if re.match(r"^\d+\.", para.text):
            read_N4_Grammar_knowledge_point.append("- Grammar: "+para.text.split(".")[1])
    return read_N4_Grammar_knowledge_point

def get_last_index(knowledge_point):   #knowledge_point = "N4_grammar_1" // "N4_vocabulary_1"
    query = "SELECT MAX(CAST(SUBSTRING(question_index, LOCATE(%s, question_index) + %s) AS UNSIGNED)) AS max_number FROM questions WHERE question_index Like %s;"
    cursor.execute(query, (knowledge_point, len(knowledge_point)+1, knowledge_point+"%"))
    result = cursor.fetchall()
    return result[0][0]

read_N4_Vocabulary_knowledge_point = read_N4_Vocabulary_knowledge_point()
read_N4_Grammar_knowledge_point = read_N4_Grammar_knowledge_point()

folder_path = r"D:\JAP_GPT\JAP_GPT\Unprocessed_Excel_Files"
processed_folder_path = r"D:\JAP_GPT\JAP_GPT\Processed_Excel_Files"

for file_name in os.listdir(folder_path):
    if file_name.endswith(('.xlsx',)):
        file_path = os.path.join(folder_path, file_name)
        df = pd.read_excel(file_path)
        #knowledge_point = file_name.split(".")[0]   #名字不一定对
        knowledge_point = os.path.splitext(file_name)[0]

        level = knowledge_point[0:2]
        type = knowledge_point[3:]

        if "Vocabulary" in knowledge_point:
            type = type[0:12] + ":" + type[12:]
            i = read_N4_Vocabulary_knowledge_point.index(type)
            question_index_front = level+"_"+'vocabulary'+"_"+str(i+1)
        elif "Grammar" in knowledge_point:
            type = type[0:9] + ":" + type[10:]
            i = read_N4_Grammar_knowledge_point.index(type)
            question_index_front = level+"_"+'grammar'+"_"+str(i+1)

        n = get_last_index(question_index_front)
        if n == None:
            n = 0
        is_gpt = 1
        '''
        if "Vocabulary" in knowledge_point:
            if knowledge_point in read_N4_Vocabulary_knowledge_point:
                level = "N4"
            else:
                level = "N5"
        elif "Grammar" in knowledge_point:
            if knowledge_point in read_N4_Grammar_knowledge_point:
                level = "N4"
            else:
                level = "N5"
        '''
        for index, row in df.iterrows():
            if row[4] == 'OK' or row[4] == 'Minor changes':
                content = row[1]+'\n'+row[2]  #缺个题型信息
                correct_answer = row[3]
                question_index = question_index_front+"_"+str(n+1)
                n += 1
                query = "INSERT INTO questions (question_index, type, level, content, is_gpt, correct_answer) VALUES (%s, %s, %s, %s, %s, %s);"
                params = (question_index, type, level, content, is_gpt, correct_answer)
                cursor.execute(query, params)
                db.commit()
            

    dst_path = os.path.join(processed_folder_path, file_name)
    shutil.move(file_path, dst_path)

db.close()