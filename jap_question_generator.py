'''Question_Generator'''
import re
import os
import glob
import time
import warnings
import docx
import mysql.connector
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate  
from langchain.chains import LLMChain   
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from typing import Any

def split_into_sentences(text):
    sentence_endings = re.compile(r'(?<=[。！？])\s*')
    sentences = sentence_endings.split(text)
    return sentences


def extract_grammar_points(docx_file):
    # 读取docx文件
    doc = docx.Document(docx_file)
    
    grammar_points = []
    
    # 遍历文档的所有段落
    for para in doc.paragraphs:
        print(f"读取的段落内容: {para.text}")
        # 使用正则表达式匹配类似 "1.～あいだ（間）" 的格式
        match = re.match(r'(\d+)\.([^\n]+)', para.text.strip())
        if match:
            # 提取编号和短语部分
            number = match.group(1)
            phrase = match.group(2).strip()
            grammar_points.append({
                'number': number,
                'phrase': phrase
            })
    
    return grammar_points


def extract_vocabulary(docx_path):
    doc = Document(docx_path)
    vocabulary_list = []
    number_pattern = r'^\d+\. '  # 匹配以数字和点号开头的行
    
    # 遍历每个段落
    for para in doc.paragraphs:
        # 将每个段落按行分割
        for line in para.text.splitlines():
            # 如果行包含带序号的内容
            if re.match(number_pattern, line.strip()):
                vocabulary_list.append(line.strip())
    
    return vocabulary_list





def grammar_points_revise(rows, output, filepath):
    filename = os.path.splitext(os.path.basename(filepath))[0]
    llm = ChatOpenAI(
        temperature=0.6,
        model='gpt-4o'
    )

    for i in range(100):
        print(f'Processing file {i}...')
        prompt = ChatPromptTemplate.from_template(
        'You are an experienced Japanese examiner who is very familiar with the Japanese N4 and N5 exams.\n'
        'Here is a list of Japanese N4 and N5 knowledge points, including 文法(grammar): {knowledge_points}.\n'
        'Based on the knowledge point and your database related to this knowledge point, generate {num_questions} new questions.\n'
        'Please try to cover all knowledge points and reduce duplication.\n'
        'You should attach the knowledge points after each questions.\n'
        "Finally, all the answers will be attached at the end. Do not attach the answer after each question.\n"
        'The format of these new questions should be as follows:\n'
        '''
        もんだい1　（  　　　　　 ）に　何を　入れますか。　1・2・3・4から　いちばん　いい　ものを　一つ　えらんで　ください。

        ①　　かれが　手伝って　（  　　　　　 ）　宿題 (しゅくだい) が　終わらなっかった。
        1　もらったから			2　くれなかったから		
        3　ほしいから				4　ほしかったから
        -Knowledge Points: - Grammar: 10. ～ことにする, 29. ～てあげる・～てもらう・～てくれる, 11. ～ことになる

        ②　宿題 (しゅくだい) を　したのに、　先生が　（  　　　　　 ）。
        1　来なかった				2　してしまった		
        3　会わなかった			4　するつもりだった
        -Knowledge Points:   - Grammar: 12. ～し, 22. ～たら〈その(後)で〉, 7. ～くれる

        ③　うちの　子どもは　勉強 (べんきょう) しないで　（  　　　　　 ）　ばかりいる。
        1　あそび		2　あそぶ		3　あそばない		4　あそんで
        -Knowledge Points:   - Grammar: 47. ～ばかり・～てばかりいる, 58. ～ないで, 48. ～てほしい
        '''
        )
        chain_one = LLMChain(llm=llm, prompt=prompt)

        input = {
            'knowledge_points': rows,
            'num_questions': 10
        }
        revise_result = chain_one.run(input)
        output_doc = Document()
        sentences = split_into_sentences(revise_result)
        for sentence in sentences:
            sentence.replace("＿＿＿", "[ ]")
            output_doc.add_paragraph(sentence)

        output_path = os.path.join(output, f"{filename}_new{i}.docx")
        output_doc.save(output_path)


def vocabulary_points_revise(vocabulary_list, output, filepath):
    filename = os.path.splitext(os.path.basename(filepath))[0]
    llm = ChatOpenAI(
        temperature=0.6,
        model='gpt-4o'
    )

    i = 1
    # 遍历每个知识点（从vocabulary_list提取的知识点）
    for knowledge_point in vocabulary_list:
        print(f'Processing vocabulary knowledge point: {knowledge_point}...')

        # 动态生成 prompt，要求 GPT 依次为每个词汇出题
        prompt = ChatPromptTemplate.from_template(
            f'''
            You are an experienced Japanese examiner who is very familiar with the Japanese N4 and N5 exams.
            Here is a Japanese vocabulary knowledge point: {knowledge_point}.
            Based on this knowledge point, generate 3 multiple-choice questions.
            Each question should have 4 options and only one correct answer. Make sure that the options are appropriate and unambiguous.
            The format of the questions should be as follows:

            もんだい1　＿＿＿の　ことばは　ひらがなで　どう　かきますか。　1・2・3・4から　いちばん　いいものを　ひとつ　えらんで　ください。

            ①　このいすに　上着を　かけてください。
            1　うえき		2　うえぎ		3　うわき		4　うわぎ
            -Knowledge Points: -Vocabulary: N4・語彙 うわぎ【上着】

            ②　山田さんは　赤い　ぼうしを　かぶって　います。
            1　あおい		2　あかい		3　くらい		4　くろい
            -Knowledge Points: - Vocabulary: N5・語彙 あかい【赤い】

            ③　紙に　名前と　住所を　書いて　ください。
            1　じゅしょう		2　じゅうしょう		3　じゅうしょ		4　じゅしょ
            -Knowledge Points: - Vocabulary: N4・語彙 じゅうしょ【住所】

            Each question should include the corresponding vocabulary knowledge point at the end.
            '''
        )
        
        # 创建链条来运行LLM
        chain_one = LLMChain(llm=llm, prompt=prompt)

        # 输入数据
        input_data = {
            'knowledge_point': knowledge_point,  # 使用当前知识点
            'num_questions': 3
        }

        # 获取生成的题目
        revise_result = chain_one.run(input_data)

        # 处理生成的题目并保存
        output_doc = Document()
        sentences = split_into_sentences(revise_result)
        
        # 将生成的题目逐句添加到文档中
        for sentence in sentences:
            sentence = sentence.replace("＿＿＿", "[ ]")  # 替换空格部分
            output_doc.add_paragraph(sentence)

        # 确保文件名不会为空
        if knowledge_point.strip():
            output_path = os.path.join(output, f"{filename}_new{i}.docx")
        else:
            output_path = os.path.join(output, f"{filename}_new_{i}.docx")
        
        # 保存文件
        output_doc.save(output_path)
        # print(f'Generated questions for {knowledge_point} saved to {output_path}')
        i += 1




# 测试
docx_file_grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\N4N5 material\\N4 Notes 文法_numbered.docx"
docx_file_vocabulary = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\N4N5 material\\N4 Notes 語彙_numbered.docx"
output_grammar = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\N4 grammar"
output_vocabulary = "C:\\Users\\刘宇\\OneDrive - CUHK-Shenzhen\\桌面\\JAP_GPT\\2025_new_db\\new_questions\\N4 vocabulary"
grammar_points = extract_grammar_points(docx_file_grammar)
vocabulary = extract_vocabulary(docx_file_vocabulary)
grammar_points_revise(grammar_points, output_grammar, docx_file_grammar)
vocabulary_points_revise(vocabulary, output_vocabulary, docx_file_vocabulary)


# # 输出结果
# for point in grammar_points:
#     print(f"Number: {point['number']}, Phrase: {point['phrase']}")


# 打印提取出的词汇列表
# for word in vocabulary:
#     print(word)

# for word in vocabulary[:10]:
#     print(word)